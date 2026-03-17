from docx import Document
from types import SimpleNamespace

from app.services.docx_builder import _markdown_heading_info, _parse_and_render_markdown
from app.services.one_click_generator.format_driven_sections.common import (
    _build_affiliated_units_statement_template,
    _build_service_after_sales_points,
    _build_service_supply_points,
    _build_service_training_points,
    _build_service_fee_commitment_template,
    _build_small_enterprise_declaration_template,
    _normalize_hlj_supplier_qualification_template,
)
from app.services.one_click_generator.format_driven_sections.cs import _build_cs_sections
from app.services.one_click_generator.format_driven_sections.tp import _build_tp_sections, _build_tp_service_plan_section
from app.services.tender_workflow.materialization import _build_materialized_service_section
from app.schemas import ProcurementPackage, ProductSpecification, TenderDocument
from scripts.repair_sample_drafts import TP_SECTION_GROUPS, _reorder_document_sections


def test_hlj_qualification_template_fallback_is_complete():
    content = _normalize_hlj_supplier_qualification_template("")

    assert "黑龙江省政府采购供应商资格承诺函" in content
    assert "承诺具有独立承担民事责任的能力" in content
    assert "执业状态" in content
    assert "附件：缴纳社会保障资金的证明材料清单" in content


def test_markdown_heading_info_keeps_long_numbered_clause_as_body():
    clause = "（三）供应商类型为非企业专业服务机构的，承诺通过合法渠道可查证“执业状态”为“正常”。"
    assert _markdown_heading_info(clause) is None

    short_heading = "一、承诺具有独立承担民事责任的能力"
    assert _markdown_heading_info(short_heading) == (2, short_heading, "outline")


def test_parse_markdown_skips_fallback_for_outline_clauses():
    doc = Document()
    content = """
一、承诺具有独立承担民事责任的能力
一、承诺具有独立承担民事责任的能力（一）供应商类型为企业的，承诺通过合法渠道可查证的信息为：
1. “类型”为“有限责任公司”“股份有限公司”等法人企业或合伙企业。
（三）供应商类型为非企业专业服务机构的，承诺通过合法渠道可查证“执业状态”为“正常”。
""".strip()

    rendered = _parse_and_render_markdown(doc, content)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    assert rendered is True
    assert "请按招标文件原格式填写本节内容。" not in texts
    assert any("执业状态" in text for text in texts)


def test_standard_declaration_templates_are_expanded():
    tender = SimpleNamespace(
        project_name="示例项目",
        project_number="ABC-001",
        purchaser="示例采购人",
        agency="示例代理机构",
    )
    packages = [SimpleNamespace(item_name="流式细胞仪")]

    sme = _build_small_enterprise_declaration_template(tender, packages)
    relation = _build_affiliated_units_statement_template(tender)
    fee = _build_service_fee_commitment_template(tender)

    assert "中小企业声明函（货物）" in sme
    assert "示例采购人" in sme
    assert "流式细胞仪" in sme
    assert "投标人关联单位的说明" in relation
    assert "如经核查不存在上述情形，请直接填写“无”" in relation
    assert "招标代理服务费承诺" in fee
    assert "二选一保留" in fee


def test_service_plan_helpers_are_detailed():
    supply = _build_service_supply_points(
        "流式细胞仪",
        "合同签订后30日内",
        "医院指定地点",
        product_identity="示例厂家 流式细胞仪（型号：FCM-1）",
    )
    training = _build_service_training_points("流式细胞仪")
    after_sales = _build_service_after_sales_points("流式细胞仪")

    assert len(supply) >= 4
    assert any("风险预警" in point or "补救措施" in point for point in supply)
    assert any("培训记录" in point for point in training)
    assert any("质保期外延续服务" in point for point in after_sales)


def test_tp_service_plan_section_is_expanded():
    pkg = SimpleNamespace(package_id="1", item_name="流式细胞仪")
    content = _build_tp_service_plan_section([pkg], "")

    assert "（一）技术服务" in content
    assert "（二）售后服务" in content
    assert "项目组织与职责分工" in content
    assert "到货保护与异常处理" in content
    assert "场地勘查与条件确认" in content
    assert "质保期外延续服务" in content


def test_materialized_service_section_is_expanded():
    tender = TenderDocument(
        project_name="示例项目",
        project_number="ABC-001",
        budget=0,
        purchaser="示例采购人",
        packages=[
            ProcurementPackage(
                package_id="1",
                item_name="流式细胞仪",
                quantity=1,
                budget=0,
                delivery_time="合同签订后30日内",
                delivery_place="医院指定地点",
            )
        ],
    )
    product = ProductSpecification(
        product_name="流式细胞仪",
        manufacturer="示例厂家",
        model="FCM-1",
        price=0,
        specifications={"检测通道": "8色", "进样方式": "自动"},
        training_notes="提供分层培训和实操演练。",
        acceptance_notes="按技术参数逐项验收。",
    )

    content = _build_materialized_service_section(tender, {"1": product})

    assert "（一）技术服务" in content
    assert "（二）售后服务" in content
    assert "项目组织与职责分工" in content
    assert "场地勘查与条件确认" in content
    assert "培训记录" in content
    assert "质保期外延续服务" in content


def _build_sample_tender_for_section_order() -> TenderDocument:
    return TenderDocument(
        project_name="示例项目",
        project_number="ABC-001",
        budget=0,
        purchaser="示例采购人",
        packages=[
            ProcurementPackage(
                package_id="1",
                item_name="流式细胞仪",
                quantity=1,
                budget=0,
                delivery_time="合同签订后30日内",
                delivery_place="医院指定地点",
            )
        ],
    )


def test_tp_section_order_follows_manual_like_chapters():
    tender = _build_sample_tender_for_section_order()
    titles = [section.section_title for section in _build_tp_sections(tender, "")]

    assert titles == [
        "一、响应文件封面格式",
        "第一章、资格性证明文件",
        "四、资格承诺函",
        "七、法定代表人/单位负责人授权书",
        "八、法定代表人/单位负责人和授权代表身份证明",
        "第二章、符合性承诺",
        "十一、投标人关联单位的说明",
        "九、小微企业声明函",
        "十、残疾人福利性单位声明函",
        "第三章、商务及技术部分",
        "二、报价书",
        "三、报价一览表",
        "六、技术服务和售后服务的内容及措施",
        "第四章、报价书附件",
        "五、技术偏离及详细配置明细表",
        "附一、资格性审查响应对照表",
        "附二、符合性审查响应对照表",
        "附三、详细评审响应对照表",
        "附四、投标无效情形汇总及自检表",
    ]


def test_cs_section_order_follows_manual_like_chapters():
    tender = _build_sample_tender_for_section_order()
    titles = [section.section_title for section in _build_cs_sections(tender, "")]

    assert titles == [
        "一、响应文件封面格式",
        "第一章、资格性证明文件",
        "十一、资格承诺函",
        "六、法定代表人/单位负责人授权书",
        "七、法定代表人/单位负责人和授权代表身份证明",
        "第二章、符合性承诺",
        "十、投标人关联单位的说明",
        "八、小微企业声明函",
        "九、残疾人福利性单位声明函",
        "第三章、商务及技术部分",
        "二、首轮报价表",
        "三、分项报价表",
        "五、技术服务和售后服务的内容及措施",
        "第四章、报价书附件",
        "四、技术偏离及详细配置明细表",
        "附一、资格性审查响应对照表",
        "附二、符合性审查响应对照表",
        "附三、详细评审响应对照表",
        "附四、投标无效情形汇总及自检表",
    ]


def test_docx_reorder_groups_existing_sections_into_chapters():
    doc = Document()
    doc.add_paragraph("封面")
    for title in [
        "二、报价书",
        "三、报价一览表",
        "四、资格承诺函",
        "五、技术偏离及详细配置明细表",
        "六、技术服务和售后服务的内容及措施",
        "七、法定代表人/单位负责人授权书",
        "八、法定代表人/单位负责人和授权代表身份证明",
        "九、小微企业声明函",
        "十、残疾人福利性单位声明函",
        "十一、投标人关联单位的说明",
        "附一、资格性审查响应对照表",
        "附二、符合性审查响应对照表",
        "附三、详细评审响应对照表",
        "附四、投标无效情形汇总及自检表",
    ]:
        doc.add_paragraph(title)
        doc.add_paragraph(f"{title}正文")

    changed = _reorder_document_sections(doc, TP_SECTION_GROUPS)
    texts = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]

    assert changed is True
    assert texts[:6] == [
        "封面",
        "第一章、资格性证明文件",
        "四、资格承诺函",
        "四、资格承诺函正文",
        "七、法定代表人/单位负责人授权书",
        "七、法定代表人/单位负责人授权书正文",
    ]
    assert "第四章、报价书附件" in texts
    assert texts.index("第三章、商务及技术部分") < texts.index("二、报价书")
    assert texts.index("第四章、报价书附件") < texts.index("五、技术偏离及详细配置明细表")
