"""招标文件解析服务"""
import json
import logging
from pathlib import Path
import time
from typing import Any, Sequence

import pypdf
from langchain_core.messages import AIMessage
from langchain_core.prompt_values import PromptValue
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import Runnable
from langchain_openai import ChatOpenAI

from app.config import get_settings
from app.schemas import (
    TenderDocument,
    ProcurementPackage,
    ResponseSectionTemplate,
    TenderTableColumn,
    TenderTableLayoutHint,
    TenderTableRowTemplate,
    TenderTableTemplate,
)
from app.services.requirement_processor import (
    _find_requirement_pair_position,
    _is_bad_requirement_name,
    _is_bad_requirement_value,
    _package_forbidden_terms, _extract_match_tokens, _safe_text,
)

try:
    from docx import Document as _DocxDocument
    from docx.document import Document as _DocxDocumentType
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table as _DocxTable
    from docx.text.paragraph import Paragraph as _DocxParagraph
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    _DocxDocumentType = None
    CT_Tbl = None
    CT_P = None
    _DocxTable = None
    _DocxParagraph = None

logger = logging.getLogger(__name__)
_FALLBACK_PARSE_CHAR_LIMITS = (24000, 16000, 10000)
_PARSE_NETWORK_RETRY_ATTEMPTS = 3
_PARSE_NETWORK_RETRY_DELAY_SECONDS = 1.0
DEFAULT_CS_SECTION_TITLES = [
    "一、响应文件封面格式",
    "二、首轮报价表",
    "三、分项报价表",
    "四、技术偏离及详细配置明细表",
    "五、技术服务和售后服务的内容及措施",
    "六、法定代表人/单位负责人授权书",
    "七、法定代表人/单位负责人和授权代表身份证明",
    "八、小微企业声明函",
    "九、残疾人福利性单位声明函",
    "十、投标人关联单位的说明",
    "十一、资格承诺函",
]

DEFAULT_ZB_SECTION_TITLES = [
    "一、投标函",
    "二、开标一览表",
    "三、投标分项报价表",
    "四、法定代表人授权书",
    "五、资格证明文件",
    "六、商务条款响应及偏离表",
    "七、技术要求响应及偏离表",
    "八、供货、安装调试、质量保障及售后服务方案",
    "九、资格性审查响应对照表",
    "十、符合性审查响应对照表",
    "十一、详细评审响应对照表",
    "十二、无效投标情形自检表",
]

DEFAULT_TP_SECTION_TITLES = [
    "一、响应文件封面格式",
    "二、报价书",
    "三、报价一览表",
    "四、资格承诺函",
    "五、技术偏离及详细配置明细表",
    "六、技术服务和售后服务的内容及措施",
    "七、法定代表人/单位负责人授权书",
    "八、法定代表人/单位负责人和授权代表身份证明",
    "九、小微企业声明函",
    "十、残疾人福利性单位声明函",
    "十一、投标人关联单位的说明",
]

import re
from typing import List


_INVALID_NOISE_PATTERNS = [
    r"否则将视为.*?废标",
    r"否则评标时不予认可",
    r"未提供.*?将导致.*?无效投标",
    r"注[:：]",
    r"说明[:：]",
]


def _clean_invalid_line(line: str) -> str:
    """清理无效行。"""
    s = re.sub(r"\s+", " ", (line or "")).strip(" \t\r\n|：:;；，,")
    s = re.sub(r"^[（(]?\d+[)）]\s*", "", s)
    s = re.sub(r"^\d+[.、]\s*", "", s)
    s = re.sub(r"^[一二三四五六七八九十]+\s*[、.]\s*", "", s)
    return s.strip()


def _looks_like_invalid_item(line: str) -> bool:
    """判断like无效项。"""
    s = _clean_invalid_line(line)
    if not s or len(s) < 10:
        return False
    for pat in _INVALID_NOISE_PATTERNS:
        if re.search(pat, s):
            return False
    keywords = [
        "无效投标", "投标无效", "被拒绝", "视为无效", "不予受理",
        "不予认可", "不得参加", "未按", "不符合", "拒绝其投标",
    ]
    return any(k in s for k in keywords)


def _dedupe_preserve_order(items: List[str]) -> List[str]:
    """去重preserveorder。"""
    seen = set()
    out = []
    for x in items:
        key = re.sub(r"\s+", "", x)
        if not key or key in seen:
            continue
        seen.add(key)
        out.append(x)
    return out


def _extract_invalid_bid_items_strict(text: str) -> List[str]:
    """提取废标项严格。"""
    text = text or ""
    items: List[str] = []

    # 一、优先抓“其他无效投标情况”和“26.5 无效投标处理”后的编号条款
    block_patterns = [
        r"26\.5[\s\S]{0,3000}",
        r"本项目规定的其他无效投标情况[:：]?[\s\S]{0,2000}",
        r"23\.2[\s\S]{0,800}",
        r"26\.6[\s\S]{0,800}",
        r"3\.5[\s\S]{0,500}",
        r"15\.3[\s\S]{0,500}",
        r"16\.1[\s\S]{0,500}",
    ]

    blocks = []
    for pat in block_patterns:
        m = re.search(pat, text)
        if m:
            blocks.append(m.group(0))

    # 二、从 block 中只提枚举条款
    enum_pat = re.compile(r"(?:^|\n)\s*[（(]?\d+[)）]\s*(.+?)(?=(?:\n\s*[（(]?\d+[)）]\s*)|\Z)", re.S)
    for blk in blocks:
        for m in enum_pat.finditer(blk):
            s = _clean_invalid_line(m.group(1))
            if _looks_like_invalid_item(s):
                items.append(s)

    # 三、补抓直接表述句
    direct_patterns = [
        r"未按上述要求提供进口产品逐级授权的投标视为未响应招标文件实质性要求，其投标无效",
        r"凡没有根据投标人须知第 15\.1 和 15\.2 条的规定随附投标保证金的投标，将按投标人须知第 23 条的规定视为无效投标予以拒绝",
        r"投标有效期不满足要求的投标将被视为无效投标而予以拒绝",
        r"投标人存在下列情况之一的，投标无效",
        r"投标人不能证明其报价合理性的，评标委员会应当将其作为无效投标处理",
    ]
    for pat in direct_patterns:
        for m in re.finditer(pat, text):
            s = _clean_invalid_line(m.group(0))
            if _looks_like_invalid_item(s):
                items.append(s)

    return _dedupe_preserve_order(items)

def _normalize_table_key(title: str, idx: int) -> str:
    """归一化表格键。"""
    title = re.sub(r"\s+", "", title or "")
    mapping = {
        "序号": "seq",
        "评审项目": "review_item",
        "审查项": "review_item",
        "条款名称": "review_item",
        "采购文件要求": "tender_requirement",
        "磋商文件要求": "tender_requirement",
        "招标文件要求": "tender_requirement",
        "响应文件内容": "response_content",
        "响应文件对应内容": "response_content",
        "投标文件内容": "response_content",
        "是否满足": "is_match",
        "是否响应": "is_match",
        "备注": "remark",
        "页码": "page_no",
        "分值": "score",
        "标准分": "score",
        "评分标准": "score_rule",
        "评审标准": "score_rule",
        "无效情形": "invalid_reason",
        "证明材料": "evidence",
        "证明材料页码": "evidence",
        "自评说明": "self_note",
    }
    return mapping.get(title, f"col_{idx}")


def _default_columns(table_kind: str) -> list[TenderTableColumn]:
    """返回默认columns。"""
    if table_kind == "qualification":
        titles = ["序号", "审查项", "采购文件要求", "响应文件对应内容", "是否满足", "备注"]
    elif table_kind == "compliance":
        titles = ["序号", "审查项", "采购文件要求", "响应文件对应内容", "是否满足", "备注"]
    elif table_kind == "detailed":
        titles = ["序号", "评审项", "采购文件评分要求", "响应文件对应内容", "自评说明", "证明材料/页码"]
    elif table_kind == "invalid":
        titles = ["序号", "无效情形", "自检结果", "备注"]
    else:
        titles = ["序号", "内容"]

    return [
        TenderTableColumn(
            key=_normalize_table_key(t, i),
            title=t,
            required=(i <= 3),
        )
        for i, t in enumerate(titles, start=1)
    ]


def _split_pipe_row(line: str) -> list[str]:
    """切分pipe行。"""
    parts = [part.strip() for part in line.strip().strip("|").split("|")]
    return [p for p in parts if p]


def _is_separator_row(line: str) -> bool:
    """判断separator行。"""
    sample = line.strip().replace("|", "").replace(":", "").replace("-", "")
    return not sample


def _make_row_from_text(text: str, columns: list[TenderTableColumn], seq: int) -> TenderTableRowTemplate:
    """返回文本中的行。"""
    cells: dict[str, str] = {}
    if columns:
        cells[columns[0].key] = str(seq)
    if len(columns) >= 2:
        cells[columns[1].key] = text
    for col in columns[2:]:
        cells[col.key] = ""

    return TenderTableRowTemplate(
        seq=str(seq),
        cells=cells,
        source_text=text,
        is_material=("★" in text or "※" in text or "实质性" in text),
    )


def _parse_table_template_from_block(block: str, table_name: str, table_kind: str) -> TenderTableTemplate | None:
    """解析文本块中的表格模板。"""
    block = (block or "").strip()
    if not block:
        return None

    lines = [line.strip() for line in block.splitlines() if line.strip()]
    pipe_lines = [line for line in lines if "|" in line and len(_split_pipe_row(line)) >= 2]

    columns: list[TenderTableColumn] = []
    rows: list[TenderTableRowTemplate] = []

    # 情况1：原文里能识别出“|”表格
    if pipe_lines:
        header_cells = _split_pipe_row(pipe_lines[0])
        columns = [
            TenderTableColumn(
                key=_normalize_table_key(title, idx),
                title=title,
                required=(idx <= 3),
            )
            for idx, title in enumerate(header_cells, start=1)
        ]

        seq_no = 1
        for line in pipe_lines[1:]:
            if _is_separator_row(line):
                continue
            cells_raw = _split_pipe_row(line)
            if not cells_raw:
                continue

            cell_map: dict[str, str] = {}
            for idx, col in enumerate(columns):
                cell_map[col.key] = cells_raw[idx] if idx < len(cells_raw) else ""

            rows.append(
                TenderTableRowTemplate(
                    seq=cell_map.get(columns[0].key, str(seq_no)),
                    cells=cell_map,
                    source_text=" | ".join(cells_raw),
                    is_material=("★" in " ".join(cells_raw) or "※" in " ".join(cells_raw)),
                )
            )
            seq_no += 1

    # 情况2：不是标准表格，而是编号条款，退化成“行模板”
    if not columns:
        columns = _default_columns(table_kind)
        merged: list[str] = []

        for line in lines:
            if re.match(r"^(?:\d+[、.）)]|[（(]?\d+[）)]|[一二三四五六七八九十]+[、.]|★|※)", line):
                merged.append(line)
            else:
                if merged:
                    merged[-1] += " " + line

        for idx, item in enumerate(merged, start=1):
            if len(item.strip()) < 2:
                continue
            rows.append(_make_row_from_text(item.strip(), columns, idx))

    if not rows:
        return None

    return TenderTableTemplate(
        table_name=table_name,
        section_title=table_name,
        source_title=table_name,
        columns=columns,
        rows=rows,
        raw_block=block,
    )


def _extract_heading_block(text: str, heading_keywords: list[str], stop_keywords: list[str] | None = None) -> str:
    """提取heading文本块。"""
    stop_keywords = stop_keywords or [
        "资格性审查", "符合性审查", "详细评审", "评分标准", "响应文件格式",
        "采购需求", "商务要求", "合同草案", "合同包", "采购包", "投标无效", "响应无效",
    ]

    all_keys = heading_keywords + stop_keywords
    pattern = re.compile(
        rf"(?:{'|'.join(map(re.escape, heading_keywords))})[：:]?(.*?)(?=(?:{'|'.join(map(re.escape, all_keys))})[：:]?|$)",
        re.S,
    )
    match = pattern.search(text or "")
    return match.group(1).strip() if match else ""


def _extract_zb_format_block(text: str) -> str:
    """
    公开招标项目：精确截取“第六章 投标文件格式/响应文件格式”正文块。
    目录里也会出现“第六章 投标文件格式”，因此不能简单取最后一次出现的位置，
    需要验证命中点后面是否真的进入了模板正文。
    """
    text = text or ""
    if not text:
        return ""

    chapter_pat = re.compile(r"(?:^|\n)\s*第六章\s*(?:投标文件格式|响应文件格式)", re.M)
    expected_markers = (
        "格式 1",
        "格式1",
        "格式 2",
        "格式2",
        "格式 3",
        "格式3",
        "格式 8",
        "格式8",
        "格式 9",
        "格式9",
        "中小企业声明函",
        "残疾人福利性单位声明函",
        "类似项目业绩表",
        "制造商授权书",
        "采购需求响应及偏离表",
        "招标代理服务费承诺",
    )

    best_start = None
    for m in chapter_pat.finditer(text):
        tail = text[m.start():]
        window = tail[:2500]
        if any(marker in window for marker in expected_markers):
            best_start = m.start()
            break

    if best_start is None:
        fallback_positions = [m.start() for m in chapter_pat.finditer(text)]
        if not fallback_positions:
            return ""
        best_start = fallback_positions[-1]

    tail = text[best_start:]
    # 公开招标文件的第六章里经常自带“附件”“附录”“申请表”等子块，
    # 这些内容同样属于投标文件格式，不能在这里提前截断。
    stop_patterns = [
        re.compile(r"(?:^|\n)\s*第七章\b", re.M),
    ]

    stop_pos = None
    for pat in stop_patterns:
        m = pat.search(tail)
        if m:
            if stop_pos is None or m.start() < stop_pos:
                stop_pos = m.start()

    return tail[:stop_pos].strip() if stop_pos is not None else tail.strip()

_ZB_FORMAT_TITLE_PATTERNS = (
    re.compile(r"^格式\s*\d+(?:-\d+)?(?:\.\d+)?(?:[.．、]?\s*.*)?$"),
    re.compile(r"^(?:7\.(?:7(?:\.\d+)?|11|12)|8\.\d+|9\.\d+)\s*(?:中小企业声明函|残疾人福利性单位声明函|节能.?环保材料|类似项目业绩表|制造商授权书|采购需求响应及偏离表|.*技术.*方案|.*证明文件.*|招标代理服务费承诺).*$"),
    re.compile(r"^售后服务承诺书$"),
    re.compile(r"^招标代理服务费承诺$"),
    re.compile(r"^[一二三四五六七八九十]+、\s*招标代理服务费承诺$"),
)

def _normalize_zb_format_title(line: str) -> str:
    """归一化ZB 格式格式标题。"""
    raw = (line or "").strip()
    if not raw:
        return ""
    compact = re.sub(r"\s+", "", raw).replace("．", ".")
    compact = compact.replace("（", "(").replace("）", ")")
    for pat in _ZB_FORMAT_TITLE_PATTERNS:
        if pat.match(compact):
            return compact
    return ""

def _zb_order_no_from_title(title: str, fallback_idx: int) -> str:
    """返回标题中的orderno。"""
    title = title or ""
    m = re.search(r"格式\s*(\d+(?:-\d+)?(?:\.\d+)?)", title)
    if m:
        return m.group(1)
    m = re.search(r"^(\d+\.\d+)", title)
    if m:
        return m.group(1)
    m = re.search(r"^([一二三四五六七八九十]+)、", title)
    if m:
        return m.group(1)
    return str(fallback_idx)


def _cleanup_zb_template_raw_block(title: str, raw_block: str) -> str:
    """返回ZB 格式模板raw文本块。"""
    lines = [line.rstrip() for line in (raw_block or "").splitlines()]
    if not lines:
        return ""

    removable_tail_markers = {"技术文件部分", "商务文件部分"}
    keep_group_markers = any(token in (title or "") for token in ("投标函", "投标书"))

    while len(lines) > 1 and lines[-1].strip() in removable_tail_markers and not keep_group_markers:
        lines.pop()

    return "\n".join(lines).strip()


def _extract_zb_response_section_templates(tender_text: str) -> list[ResponseSectionTemplate]:
    """
    从公开招标文件第六章中抽取原始“格式模板”。
    只识别真正的模板标题，不再把 3.1 / 8.1 / 17.2 这类正文条款当成模板标题。
    """
    block = _extract_zb_format_block(tender_text)
    if not block:
        return []

    lines = [line.rstrip() for line in block.splitlines()]
    title_hits: list[tuple[int, str]] = []

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        title = _normalize_zb_format_title(line)

        if title and re.fullmatch(r"格式\s*\d+(?:-\d+)+", title):
            next_title = ""
            j = i + 1
            while j < len(lines):
                nxt = lines[j].strip()
                if nxt:
                    next_title = nxt
                    break
                j += 1
            if next_title:
                merged_title = _normalize_zb_format_title(f"{title} {next_title}") or f"{title} {next_title}"
                title_hits.append((i, merged_title))
                i += 1
                continue

        if title:
            title_hits.append((i, title))

        i += 1

    if len(title_hits) < 3:
        return []

    templates: list[ResponseSectionTemplate] = []
    seen: set[str] = set()
    for idx, (start_idx, title) in enumerate(title_hits):
        end_idx = title_hits[idx + 1][0] if idx + 1 < len(title_hits) else len(lines)
        raw_block = _cleanup_zb_template_raw_block(title, "\n".join(lines[start_idx:end_idx]).strip())
        key = re.sub(r"\s+", "", title)
        if not raw_block or key in seen:
            continue
        seen.add(key)
        templates.append(
            ResponseSectionTemplate(
                order_no=_zb_order_no_from_title(title, idx + 1),
                title=title,
                required=True,
                raw_block=raw_block,
            )
        )
    return templates

def _iter_docx_blocks(doc: _DocxDocumentType):
    """
    按 Word 正文真实顺序遍历段落/表格。
    这是关键，否则会把所有段落先取完、再把所有表格追加到最后，导致第六章模板块被打散。
    """
    if not _DOCX_AVAILABLE:
        return
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield _DocxParagraph
        elif isinstance(child, CT_Tbl):
            yield _DocxTable


def _default_section_titles(procurement_type: str) -> list[str]:
    """返回默认章节标题。"""
    mode = (procurement_type or "").strip()

    if "公开招标" in mode or ("招标" in mode and "谈判" not in mode and "磋商" not in mode):
        return DEFAULT_ZB_SECTION_TITLES.copy()
    if "磋商" in mode:
        return DEFAULT_CS_SECTION_TITLES.copy()
    if "谈判" in mode:
        return DEFAULT_TP_SECTION_TITLES.copy()

    return []


def _default_response_section_templates(procurement_type: str) -> list[ResponseSectionTemplate]:
    """返回默认响应章节模板。"""
    templates: list[ResponseSectionTemplate] = []
    for title in _default_section_titles(procurement_type):
        match = re.match(r"^([一二三四五六七八九十]+)、", title)
        templates.append(
            ResponseSectionTemplate(
                order_no=match.group(1) if match else "",
                title=title,
                required=True,
                raw_block="",
            )
        )
    return templates


def _normalize_layout_header(text: str) -> str:
    """归一化布局表头。"""
    return re.sub(r"[\s\r\n\t]+", "", text or "")


def _header_signature(headers: Sequence[str]) -> tuple[str, ...]:
    """生成可比较的表头签名。"""
    return tuple(_normalize_layout_header(item) for item in headers if _normalize_layout_header(item))

def _fails_package_domain_guard(item_name: str, key: str, req: str) -> bool:
    """
    防止不同包的技术条款串包。
    例如：化学发光设备不应出现“琼脂凝胶电泳法”，
    电泳设备不应出现“流式细胞分析仪”“化学发光法”。
    """
    item = _safe_text(item_name, "")
    text = f"{_safe_text(key, '')} {_safe_text(req, '')}"

    guards: list[tuple[tuple[str, ...], tuple[str, ...]]] = [
        (("电泳",), ("化学发光", "流式细胞")),
        (("化学发光",), ("电泳", "流式细胞")),
        (("流式细胞",), ("电泳", "化学发光")),
        (("荧光显微镜",), ("流式细胞", "化学发光", "电泳法")),
        (("荧光操作仪",), ("流式细胞", "化学发光", "电泳法")),
    ]

    for item_tokens, bad_tokens in guards:
        if any(tok in item for tok in item_tokens):
            if any(bad in text for bad in bad_tokens):
                return True
    return False


class TenderParser:
    """招标文件解析器"""

    def __init__(self, llm: ChatOpenAI | None = None):
        """
        初始化解析器

        Args:
            llm: 语言模型实例，如果为None则使用默认配置
        """
        self.llm = llm or ChatOpenAI(model="gpt-4o-mini", temperature=0)
        settings = get_settings()
        self.max_parse_chars = max(0, settings.tender_parse_char_limit)

        self.extraction_prompt = ChatPromptTemplate.from_messages([
            ("system", """你是一个专业的政府采购招标文件分析专家。你的任务是从招标文件中准确提取关键信息。

要求：
1. 确保所有数字信息准确无误（金额、数量等）
2. 完整提取技术参数，不要遗漏
3. 准确识别商务条款
4. 如果某些信息缺失，对应字段设为空字符串或空列表
5. 输出必须是有效的JSON格式

输出JSON格式示例：
{{
    "project_name": "项目名称",
    "project_number": "项目编号",
    "budget": 1000000.00,
    "purchaser": "采购单位",
    "agency": "代理机构",
    "procurement_type": "竞争性谈判/竞争性磋商/公开招标",
    "packages": [
        {{
            "package_id": "1",
            "item_name": "货物名称",
            "quantity": 1,
            "budget": 100000.00,
            "technical_requirements": {{"参数1": "要求1", "参数2": "要求2"}},
            "delivery_time": "签订合同后XX个工作日",
            "delivery_place": "采购人指定地点"
        }}
    ],
    "commercial_terms": {{
        "payment_method": "验收完成后支付",
        "validity_period": "90日历天",
        "warranty_period": "1年",
        "performance_bond": "不收取"
    }},
    "evaluation_criteria": {{"技术部分": 60, "商务部分": 20, "价格部分": 20}},
    "special_requirements": "特殊要求说明"
}}"""),
            ("user", "招标文件内容：\n\n{tender_text}\n\n请提取上述信息并返回JSON格式的结果。")
        ])

    @staticmethod
    def _response_text(content: Any) -> str:
        """兼容不同模型返回结构，统一抽取为纯文本。"""
        if isinstance(content, str):
            return content.strip()
        if isinstance(content, list):
            parts: list[str] = []
            for item in content:
                if isinstance(item, dict):
                    text = item.get("text")
                    if isinstance(text, str) and text.strip():
                        parts.append(text.strip())
            return "\n".join(parts).strip()
        return str(content).strip()

    @staticmethod
    def _extract_json_dict(raw_text: str) -> dict[str, Any]:
        """提取JSONdict。"""
        text = raw_text.strip()
        if not text:
            raise ValueError("LLM返回为空")

        if text.startswith("```json"):
            text = text[7:]
        if text.startswith("```"):
            text = text[3:]
        if text.endswith("```"):
            text = text[:-3]
        text = text.strip()

        # 优先直接按完整JSON解析
        try:
            data = json.loads(text)
            if isinstance(data, dict):
                return data
        except json.JSONDecodeError:
            pass

        # 兜底：抽取第一个可能的JSON对象片段
        match = re.search(r"\{[\s\S]*}", text)
        if match:
            data = json.loads(match.group(0))
            if isinstance(data, dict):
                return data

        raise ValueError("未识别到有效JSON对象")

    # 解析招标文件需要返回完整 JSON，必须保证足够的输出 token
    _MIN_PARSE_MAX_TOKENS = 4096

    def _parse_with_llm(self, tender_text: str) -> dict[str, Any]:
        # 确保 max_tokens 足够容纳完整 JSON 响应
        """解析withLLM。"""
        llm = self._ensure_parse_tokens(self.llm)
        chain = self.extraction_prompt | llm
        response = chain.invoke({"tender_text": tender_text})
        self._check_finish_reason(response, context="parse_tender")
        response_text = self._response_text(response.content)
        return self._extract_json_dict(response_text)

    @staticmethod
    def _is_retryable_parse_error(exc: Exception) -> bool:
        """判断当前解析异常是否适合重试。"""
        text = str(exc or "").strip().lower()
        if not text:
            return False
        retryable_markers = (
            "connection error",
            "timed out",
            "timeout",
            "temporarily unavailable",
            "rate limit",
            "502",
            "503",
            "504",
            "apiconnectionerror",
        )
        return any(marker in text for marker in retryable_markers)

    def _ensure_parse_tokens(self, llm: ChatOpenAI) -> Runnable[
                                                           PromptValue | str | Sequence[Any], AIMessage] | ChatOpenAI:
        """如果当前 max_tokens 过小，返回一个放宽限制的副本。"""
        current = getattr(llm, "max_tokens", None) or 0
        if 0 < current < self._MIN_PARSE_MAX_TOKENS:
            logger.info(
                "解析调用 max_tokens=%d 过小，临时提升至 %d",
                current, self._MIN_PARSE_MAX_TOKENS,
            )
            return llm.bind(max_tokens=self._MIN_PARSE_MAX_TOKENS)
        return llm

    @staticmethod
    def _check_finish_reason(response: Any, context: str = "") -> None:
        """检测 LLM 响应是否被截断（finish_reason != 'stop'）。"""
        finish = getattr(response, "response_metadata", {}).get("finish_reason", "")
        if finish and finish != "stop":
            logger.warning(
                "LLM 响应被截断（finish_reason=%s, context=%s）。"
                "这通常是 max_tokens 不足导致的，请在 .env 中增大 LLM_MAX_TOKENS 或设为 0。",
                finish, context,
            )

    def _apply_parse_length_limit(self, tender_text: str) -> str:
        """解析applylengthlimit。"""
        if 0 < self.max_parse_chars < len(tender_text):
            logger.warning(
                "招标文件过长 (%d 字符)，将截取前%d字符",
                len(tender_text),
                self.max_parse_chars,
            )
            return tender_text[:self.max_parse_chars]
        return tender_text

    @staticmethod
    def _needs_technical_enrichment(tender_doc: TenderDocument) -> bool:
        """判断是否需要技术enrichment。"""
        for pkg in tender_doc.packages:
            tech = pkg.technical_requirements or {}
            if len(tech) < 2:
                return True
        return False

    def _enrich_package_requirements(self, tender_doc: TenderDocument, tender_text: str) -> TenderDocument:
        """为技术参数不足的包件补提取需求项。"""
        if not tender_doc.packages:
            return tender_doc

        updated_packages: list[ProcurementPackage] = []
        changed = False
        for pkg in tender_doc.packages:
            tech = pkg.technical_requirements or {}
            if len(tech) >= 2:
                updated_packages.append(pkg)
                continue

            try:
                extracted = self.extract_technical_requirements(tender_text, pkg.package_id)
            except Exception as exc:  # noqa: BLE001
                logger.warning("采购包%s技术参数补充提取失败：%s", pkg.package_id, exc)
                extracted = {}

            if extracted:
                updated_packages.append(pkg.model_copy(update={"technical_requirements": extracted}))
                changed = True
            else:
                updated_packages.append(pkg)

        if changed:
            return tender_doc.model_copy(update={"packages": updated_packages})
        return tender_doc

    @staticmethod
    def _find_next_package_start(lines: list[str], start_idx: int, current_package_id: str) -> int | None:
        """从 start_idx 之后找到下一个不同采购包的起始行索引。"""
        same_markers = (f"包{current_package_id}", f"第{current_package_id}包", f"{current_package_id}包")
        for i in range(start_idx + 1, len(lines)):
            line = lines[i].strip()
            if not line:
                continue
            if re.search(r"包\s*\d+|第\s*\d+\s*包|\d+\s*包", line) and not any(
                marker in line for marker in same_markers
            ):
                return i
        return None

    @staticmethod
    def _extract_package_scope(tender_text: str, package_id: str, item_name: str) -> str:
        """按「当前包开始 → 下一包开始」切范围，找不到则返回空串。"""
        lines = tender_text.splitlines()
        if not lines:
            return ""

        item_tokens = [token for token in re.split(r"[，,、；;（）()\\s/]+", item_name or "") if len(token) >= 2]
        markers = (f"包{package_id}", f"第{package_id}包", f"{package_id}包")
        candidate_indexes: list[int] = []

        for idx, raw_line in enumerate(lines):
            line = raw_line.strip()
            if not line:
                continue
            if any(marker in line for marker in markers) or any(token in line for token in item_tokens):
                candidate_indexes.append(idx)

        if not candidate_indexes:
            return ""

        scopes: list[str] = []
        same_package_markers = (f"包{package_id}", f"第{package_id}包", f"{package_id}包")
        for idx in candidate_indexes[:3]:
            current_line = lines[idx].strip()
            start = idx
            if not any(marker in current_line for marker in same_package_markers):
                while start > 0 and idx - start < 6:
                    previous = lines[start - 1].strip()
                    if previous and re.search(r"包\s*\d+|第\s*\d+\s*包|\d+\s*包", previous) and not any(
                        marker in previous for marker in same_package_markers
                    ):
                        break
                    start -= 1

            next_start = TenderParser._find_next_package_start(lines, idx, package_id)
            end = next_start if next_start is not None else min(len(lines), idx + 120)
            scope = "\n".join(lines[start:end]).strip()
            if scope:
                scopes.append(scope)
        return "\n".join(scope for scope in scopes if scope)

    @classmethod
    def _infer_package_quantity_from_text(cls, tender_text: str, package_id: str, item_name: str, current_quantity: int) -> int:
        """数量只在当前包 scope 里找，找不到保留原值。"""
        scope = cls._extract_package_scope(tender_text, package_id, item_name)
        if not scope.strip():
            return max(1, current_quantity)

        patterns = (
            r"设备总台数\s*[:：;；]?\s*(\d+)\s*台",
            r"采购数量\s*[:：;；]?\s*(\d+)\s*(?:台|套|个|把|件|组|副|本)?",
            r"数量\s*[:：;；]?\s*(\d+)\s*(?:台|套|个|把|件|组|副|本)",
        )

        for raw_line in scope.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            for pattern in patterns:
                match = re.search(pattern, line)
                if not match:
                    continue
                quantity = int(match.group(1))
                if quantity > 0:
                    return quantity

        return max(1, current_quantity)

    def _enrich_package_quantities(self, tender_doc: TenderDocument, tender_text: str) -> TenderDocument:
        """为数量缺失的包件补提取采购数量。"""
        if not tender_doc.packages:
            return tender_doc

        updated_packages: list[ProcurementPackage] = []
        changed = False
        for pkg in tender_doc.packages:
            inferred_quantity = self._infer_package_quantity_from_text(
                tender_text=tender_text,
                package_id=pkg.package_id,
                item_name=pkg.item_name,
                current_quantity=pkg.quantity,
            )
            if inferred_quantity != pkg.quantity:
                updated_packages.append(pkg.model_copy(update={"quantity": inferred_quantity}))
                changed = True
            else:
                updated_packages.append(pkg)

        if changed:
            return tender_doc.model_copy(update={"packages": updated_packages})
        return tender_doc

    @staticmethod
    def _extract_response_section_titles(tender_text: str, procurement_type: str) -> list[str]:
        """提取响应章节标题。"""
        mode = (procurement_type or "").strip()
        is_zb = "公开招标" in mode or ("招标" in mode and "谈判" not in mode and "磋商" not in mode)

        # 公开招标：优先提取第六章原始格式标题；失败时再回退默认骨架
        if is_zb:
            templates = _extract_zb_response_section_templates(tender_text)
            if templates:
                return [tpl.title for tpl in templates]
            return DEFAULT_ZB_SECTION_TITLES.copy()

        # 竞争性磋商/谈判当前走系统固定章节骨架，不直接复刻原文件第六章。
        # PDF 中前文常会引用“第六章 响应文件格式”，继续向后盲提 `一、/二、...`
        # 容易把“由于厂家赠送（奖励）因素需提供以下材料”这类说明项误收进必需章节。
        return _default_section_titles(procurement_type)

    def _extract_response_section_templates(self, tender_text: str, procurement_type: str):
        """提取响应章节模板。"""
        mode = (procurement_type or "").strip()
        is_zb = "公开招标" in mode or ("招标" in mode and "谈判" not in mode and "磋商" not in mode)

        # 公开招标：优先抽第六章原格式；失败时才回退默认 12 章节
        if is_zb:
            templates = _extract_zb_response_section_templates(tender_text)
            if templates:
                return templates

            return _default_response_section_templates(procurement_type)

        return _default_response_section_templates(procurement_type)

    def _extract_review_tables(self, tender_text: str, procurement_type: str) -> dict[str, TenderTableTemplate | None]:
        """提取评审表格。"""
        def _extract_precise_block(text: str, anchor_patterns: list[str], stop_patterns: list[str]) -> str:
            """提取precise文本块。"""
            text = text or ""
            if not text:
                return ""

            start = None
            for pat in anchor_patterns:
                m = re.search(pat, text, re.S | re.M)
                if m:
                    start = m.start()
                    break

            if start is None:
                return ""

            tail = text[start:]
            stop_pos = None
            for pat in stop_patterns:
                m = re.search(pat, tail, re.S | re.M)
                if m and m.start() > 0:
                    if stop_pos is None or m.start() < stop_pos:
                        stop_pos = m.start()

            return tail[:stop_pos].strip() if stop_pos is not None else tail.strip()

        def _table_headers(tpl: TenderTableTemplate | None) -> list[str]:
            """提取表模板中的表头标题。"""
            return [str(getattr(col, "title", "") or "").strip() for col in (getattr(tpl, "columns", None) or [])]

        def _header_matches(tpl: TenderTableTemplate | None, expected: list[str]) -> bool:
            """判断表模板表头是否与目标表头一致。"""
            headers = [re.sub(r"\s+", "", item) for item in _table_headers(tpl)]
            target = [re.sub(r"\s+", "", item) for item in expected]
            return headers == target

        is_zb = (
            "公开招标" in (procurement_type or "")
            or ("招标" in (procurement_type or "") and "谈判" not in (procurement_type or "") and "磋商" not in (procurement_type or ""))
        )
        zb_format_block = _extract_zb_format_block(tender_text) if is_zb else ""

        qualification = None
        compliance = None
        detailed = None

        if zb_format_block:
            qualification_block = _extract_precise_block(
                zb_format_block,
                [
                    r"资格性检查索引",
                    r"资格性审查响应对照表",
                    r"资格性审查表",
                    r"资格性检查",
                    r"资格性审查",
                ],
                [
                    r"符合性检查索引",
                    r"符合性检查",
                    r"符合性审查",
                    r"评分办法索引",
                    r"一、商务部分",
                    r"格式\s*1",
                ],
            )
            qualification = _parse_table_template_from_block(
                qualification_block,
                "资格性审查响应对照表",
                "qualification",
            )
            if qualification and not _header_matches(
                qualification,
                ["序号", "审查内容", "合格条件", "投标文件对应页码"],
            ):
                qualification = None

            compliance_block = _extract_precise_block(
                zb_format_block,
                [
                    r"符合性检查索引",
                    r"符合性审查响应对照表",
                    r"符合性审查表",
                    r"符合性检查",
                    r"符合性审查",
                ],
                [
                    r"评分办法索引",
                    r"一、商务部分",
                    r"格式\s*1",
                ],
            )
            compliance = _parse_table_template_from_block(
                compliance_block,
                "符合性审查响应对照表",
                "compliance",
            )
            if compliance and not _header_matches(
                compliance,
                ["序号", "审查内容", "合格条件", "投标文件所在页码"],
            ):
                compliance = None

            detailed_block = _extract_precise_block(
                zb_format_block,
                [
                    r"评分办法索引",
                    r"详细评审响应对照表",
                    r"详细评审",
                ],
                [
                    r"一、商务部分",
                    r"格式\s*1",
                ],
            )
            detailed = _parse_table_template_from_block(
                detailed_block,
                "详细评审响应对照表",
                "detailed",
            )
            if detailed and not _header_matches(
                detailed,
                ["序号", "内容", "评分因素分项", "评审标准", "投标文件对应页码"],
            ):
                detailed = None

        qualification_block = ""
        if qualification is None:
            qualification_block = _extract_precise_block(
                tender_text,
                [
                    r"评审方法前附表[（(]一[）)]\s*资格性检查",
                    r"资格性检查",
                    r"资格审查表",
                    r"资格性审查表",
                ],
                [
                    r"评审方法前附表[（(]二[）)]\s*符合性检查",
                    r"符合性检查",
                    r"符合性审查",
                    r"评分办法索引",
                    r"评分标准",
                ],
            )

        compliance_block = ""
        if compliance is None:
            compliance_block = _extract_precise_block(
                tender_text,
                [
                    r"评审方法前附表[（(]二[）)]\s*符合性检查",
                    r"符合性检查索引",
                    r"符合性检查",
                    r"符合性审查表",
                    r"符合性审查",
                ],
                [
                    r"评分办法索引",
                    r"7\.\s*评分因素和评分标准",
                    r"评分标准",
                    r"第六章\s*投标文件格式",
                    r"第六章\s*响应文件格式",
                ],
            )

        detailed_block = ""
        if detailed is None:
            detailed_block = _extract_precise_block(
                tender_text,
                [
                    r"7\.\s*评分因素和评分标准",
                    r"评分标准",
                    r"评分办法索引",
                    r"详细评审",
                ],
                [
                    r"第六章\s*投标文件格式",
                    r"第六章\s*响应文件格式",
                    r"格式\s*1",
                ],
            )

        invalid_block = _extract_precise_block(
            tender_text,
            [
                r"26\.5\s*.*?其他无效投标情况",
                r"其他无效投标情况",
                r"投标无效",
                r"无效投标",
                r"响应无效",
            ],
            [
                r"第六章\s*投标文件格式",
                r"第六章\s*响应文件格式",
                r"格式\s*1",
            ],
        )

        if qualification is None:
            qualification = _parse_table_template_from_block(
                qualification_block,
                "资格性审查响应对照表",
                "qualification",
            )
        if compliance is None:
            compliance = _parse_table_template_from_block(
                compliance_block,
                "符合性审查响应对照表",
                "compliance",
            )
        if detailed is None:
            detailed = _parse_table_template_from_block(
                detailed_block,
                "详细评审响应对照表",
                "detailed",
            )
        invalid_table = _parse_table_template_from_block(invalid_block, "投标无效情形汇总及自检表", "invalid")

        if qualification is None:
            qualification = TenderTableTemplate(
                table_name="资格性审查响应对照表",
                section_title="资格性审查响应对照表",
                source_title="资格性审查响应对照表",
                columns=_default_columns("qualification"),
                rows=[],
                raw_block=qualification_block or "",
            )

        if compliance is None:
            compliance = TenderTableTemplate(
                table_name="符合性审查响应对照表",
                section_title="符合性审查响应对照表",
                source_title="符合性审查响应对照表",
                columns=_default_columns("compliance"),
                rows=[],
                raw_block=compliance_block or "",
            )

        if detailed is None:
            detailed = TenderTableTemplate(
                table_name="详细评审响应对照表",
                section_title="详细评审响应对照表",
                source_title="详细评审响应对照表",
                columns=_default_columns("detailed"),
                rows=[],
                raw_block=detailed_block or "",
            )

        if invalid_table is None:
            invalid_table = TenderTableTemplate(
                table_name="投标无效情形汇总及自检表",
                section_title="投标无效情形汇总及自检表",
                source_title="投标无效情形汇总及自检表",
                columns=_default_columns("invalid"),
                rows=[],
                raw_block=invalid_block or "",
            )

        return {
            "qualification_review_table": qualification,
            "compliance_review_table": compliance,
            "detailed_review_table": detailed,
            "invalid_bid_table": invalid_table,
        }

    def _normalize_procurement_type(self, tender_doc: TenderDocument, tender_text: str) -> TenderDocument:
        """归一化采购type。"""
        text = " ".join(
            [
                tender_text or "",
                str(getattr(tender_doc, "project_name", "") or ""),
                str(getattr(tender_doc, "project_number", "") or ""),
                str(getattr(tender_doc, "procurement_type", "") or ""),
            ]
        )

        detected = getattr(tender_doc, "procurement_type", "") or ""

        if "竞争性谈判文件" in text or "[TP]" in text or "竞争性谈判" in text:
            detected = "竞争性谈判"
        elif "竞争性磋商文件" in text or "[CS]" in text or "竞争性磋商" in text:
            detected = "竞争性磋商"
        elif (
                "公开招标" in text
                or ("招标文件" in text and "投标人须知" in text and ("评标办法" in text or "综合评分法" in text))
        ):
            detected = "公开招标"

        if detected != (getattr(tender_doc, "procurement_type", "") or ""):
            return tender_doc.model_copy(update={"procurement_type": detected})
        return tender_doc

    def _enrich_format_templates(self, tender_doc: TenderDocument, tender_text: str) -> TenderDocument:
        """格式化enrich模板。"""
        section_templates = self._extract_response_section_templates(tender_text, tender_doc.procurement_type)
        review_tables = self._extract_review_tables(tender_text, tender_doc.procurement_type)
        section_titles = [item.title for item in section_templates]

        update_payload = {
            "response_section_titles": section_titles,
            "response_section_templates": section_templates,
            **review_tables,
        }
        return tender_doc.model_copy(update=update_payload)

    def _extract_docx_table_layout_hints(self, docx_path: str | Path) -> list[TenderTableLayoutHint]:
        """提取DOCX表格layouthints。"""
        if not _DOCX_AVAILABLE:
            return []

        try:
            doc = _DocxDocument(str(docx_path))
        except Exception as exc:  # noqa: BLE001
            logger.warning("读取 DOCX 表格版式失败：%s", exc)
            return []

        hints: list[TenderTableLayoutHint] = []
        last_title = ""

        for block in _iter_docx_blocks(doc):
            if isinstance(block, _DocxParagraph):
                text = re.sub(r"\s+", " ", (block.text or "")).strip()
                if text:
                    last_title = text
                continue

            if not isinstance(block, _DocxTable):
                continue

            header_titles: list[str] = []
            for row in block.rows:
                cells = [re.sub(r"\s+", " ", (cell.text or "")).strip() for cell in row.cells]
                if any(cells):
                    header_titles = cells
                    break

            signature = _header_signature(header_titles)
            if len(signature) < 2:
                continue

            widths_cm: list[float] = []
            grid = getattr(block._tbl, "tblGrid", None)
            if grid is not None:
                for col in getattr(grid, "gridCol_lst", []):
                    raw_width = int(getattr(col, "w", 0) or 0)
                    if raw_width > 0:
                        widths_cm.append(round(raw_width / 360000.0, 3))

            if len(widths_cm) != len(header_titles):
                widths_cm = []

            align = ""
            if block.alignment is not None:
                align = str(block.alignment).split(".")[-1].lower()

            style_name = ""
            try:
                style_name = block.style.name if block.style else ""
            except Exception:  # noqa: BLE001
                style_name = ""

            hints.append(
                TenderTableLayoutHint(
                    header_titles=header_titles,
                    column_width_cm=widths_cm,
                    section_title=last_title,
                    source_path=str(docx_path),
                    table_style=style_name,
                    alignment=align,
                )
            )

        deduped: list[TenderTableLayoutHint] = []
        seen: set[tuple[tuple[str, ...], tuple[int, ...], str]] = set()
        for hint in hints:
            key = (
                _header_signature(hint.header_titles),
                tuple(int(round(w * 1000)) for w in hint.column_width_cm),
                re.sub(r"\s+", "", hint.section_title or ""),
            )
            if key in seen:
                continue
            seen.add(key)
            deduped.append(hint)

        return deduped

    def _enrich_docx_table_layouts(self, tender_doc: TenderDocument, file_path: str | Path) -> TenderDocument:
        """返回DOCX表格layouts。"""
        suffix = Path(file_path).suffix.lower()
        if suffix not in (".docx", ".doc"):
            return tender_doc

        hints = self._extract_docx_table_layout_hints(file_path)
        if not hints:
            return tender_doc

        return tender_doc.model_copy(update={"table_layout_hints": hints})



    @staticmethod
    def extract_text_from_pdf(pdf_path: str | Path) -> str:
        """从PDF文件中提取文本，并对碎片化换行进行归一化拼接。"""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = pypdf.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"

                logger.info(f"成功从PDF提取文本，共 {len(text)} 字符")
                text = TenderParser._normalize_pdf_text(text)
                return text
        except Exception as e:
            logger.error(f"PDF文本提取失败: {str(e)}")
            raise ValueError(f"无法读取PDF文件: {str(e)}")

    @staticmethod
    def _normalize_pdf_text(text: str) -> str:
        """归一化 PDF 提取文本：合并被碎片化换行拆断的行。

        PDF 排版经常将一句话拆成多行短片段，例如：
            □\n企业\n□\n事业单位
        需要将它们重新拼接为连续文本。
        """
        if not text:
            return text
        # 去除页码标记
        text = re.sub(r"-\s*第?\s*\d+\s*页?\s*-", "\n", text)
        # 去除零宽字符
        text = re.sub(r"[\u200b\ufeff\u00a0]+", " ", text)

        lines = text.split("\n")
        merged: list[str] = []
        buf = ""

        for line in lines:
            stripped = line.strip()
            if not stripped:
                # 空行：刷出缓冲区
                if buf:
                    merged.append(buf)
                    buf = ""
                merged.append("")
                continue

            # 判断当前行是否为"结构性标题"——不应与前行合并
            is_structural = bool(re.match(
                r"^(?:"
                r"第[一二三四五六七八九十百]+(?:章|节|部分|条)"
                r"|[一二三四五六七八九十百]+\s*[、.]"
                r"|(?:（|[(])\s*[一二三四五六七八九十\d]+\s*(?:）|[)])"
                r"|\d+(?:\.\d+)*\s*[、.]"
                r"|[★▲■●※]"
                r"|附[一二三四五六七八九十\d]"
                r"|#+\s"
                r"|\|"   # markdown table row
                r")",
                stripped,
            ))

            # 如果当前行是结构标题，先刷出缓冲
            if is_structural:
                if buf:
                    merged.append(buf)
                    buf = ""
                buf = stripped
                continue

            # 短片段（≤6字符且不以句末标点结尾）大概率是 PDF 碎片，拼接到缓冲
            is_fragment = (
                len(stripped) <= 6
                and not stripped.endswith(("。", "；", "：", ".", ";", ":", "）", ")"))
                and not re.fullmatch(r"\d+", stripped)
            )

            if buf:
                # 如果上一行末尾是中文且当前行开头也是中文，直接拼接无空格
                if buf and stripped:
                    last_char = buf[-1]
                    first_char = stripped[0]
                    is_cjk_last = "\u4e00" <= last_char <= "\u9fff" or last_char in "\uff0c\u3002\uff1b\uff1a\u3001\uff09\u3011\u300d\u300b\u2019\u201d\u25a1\u2611\u2610"
                    is_cjk_first = "\u4e00" <= first_char <= "\u9fff" or first_char in "\uff08\u3010\u300c\u300a\u2018\u201c\u25a1\u2611\u2610"
                    if is_fragment or (is_cjk_last and is_cjk_first):
                        buf += stripped
                    elif is_cjk_last or is_cjk_first:
                        buf += stripped
                    else:
                        buf += " " + stripped
                else:
                    buf += stripped
            else:
                buf = stripped

            # 如果行以句末标点结束，刷出缓冲
            if buf.endswith(("。", "；", ".", ";")) and not is_fragment:
                merged.append(buf)
                buf = ""

        if buf:
            merged.append(buf)

        # 去除连续空行
        result: list[str] = []
        prev_empty = False
        for line in merged:
            if not line.strip():
                if prev_empty:
                    continue
                prev_empty = True
            else:
                prev_empty = False
            result.append(line)

        return "\n".join(result)

    @staticmethod
    def extract_text_from_docx(docx_path: str | Path) -> str:
        """从Word文档(.docx)中提取文本，并尽量保留段落/表格原始顺序。"""
        if not _DOCX_AVAILABLE:
            raise ValueError("python-docx 未安装，无法读取Word文件")
        try:
            doc = _DocxDocument(str(docx_path))
            blocks: list[str] = []

            for block in _iter_docx_blocks(doc):
                if isinstance(block, _DocxParagraph):
                    line = (block.text or "").strip()
                    if line:
                        blocks.append(line)
                    continue

                if isinstance(block, _DocxTable):
                    for row in block.rows:
                        cells = [" ".join((cell.text or "").split()) for cell in row.cells]
                        cells = [c for c in cells if c]
                        if cells:
                            # 用 markdown table 风格保留表格行，便于后续模板解析
                            blocks.append("| " + " | ".join(cells) + " |")

            text = "\n".join(blocks)
            logger.info(f"成功从Word文档提取文本，共 {len(text)} 字符")
            return text
        except Exception as e:
            logger.error(f"Word文档文本提取失败: {str(e)}")
            raise ValueError(f"无法读取Word文件: {str(e)}")

    def extract_text(self, file_path: str | Path) -> str:
        """根据文件扩展名自动选择提取方式（支持 .pdf 和 .docx）"""
        path = Path(file_path)
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"不支持的文件格式：{suffix}，仅支持 .pdf 和 .docx")

    def parse_tender_document(self, pdf_path: str | Path) -> TenderDocument:
        """
        解析招标文件（支持PDF和DOCX）

        Args:
            pdf_path: 文件路径（.pdf 或 .docx）

        Returns:
            结构化的招标文件数据
        """
        file_path = Path(pdf_path)

        # 1. 提取文本（自动识别格式）
        tender_text = self.extract_text(file_path)

        # 可选限制长度并增加降级重试，降低空响应概率
        tender_text = self._apply_parse_length_limit(tender_text)

        candidate_lengths = [len(tender_text)] + [x for x in _FALLBACK_PARSE_CHAR_LIMITS if x < len(tender_text)]
        last_error: Exception | None = None

        for limit in candidate_lengths:
            candidate_text = tender_text[:limit]
            for attempt in range(1, _PARSE_NETWORK_RETRY_ATTEMPTS + 1):
                try:
                    parsed_data = self._parse_with_llm(candidate_text)
                    tender_doc = TenderDocument(**parsed_data)
                    tender_doc = self._normalize_procurement_type(tender_doc, candidate_text)

                    if self._needs_technical_enrichment(tender_doc):
                        tender_doc = self._enrich_package_requirements(tender_doc, candidate_text)

                    tender_doc = self._enrich_package_quantities(tender_doc, candidate_text)
                    tender_doc = self._enrich_format_templates(tender_doc, candidate_text)
                    tender_doc = self._enrich_docx_table_layouts(tender_doc, file_path)
                    logger.info(f"成功解析招标文件: {tender_doc.project_name}（输入长度={limit}）")
                    return tender_doc
                except Exception as exc:  # noqa: BLE001
                    last_error = exc
                    logger.warning("解析尝试失败（输入长度=%d，第%d次）：%s", limit, attempt, exc)
                    if not self._is_retryable_parse_error(exc) or attempt >= _PARSE_NETWORK_RETRY_ATTEMPTS:
                        break
                    time.sleep(_PARSE_NETWORK_RETRY_DELAY_SECONDS * attempt)

        settings = get_settings()
        logger.error(
            "招标文件解析失败，所有重试均未返回有效JSON：%s（model=%s, max_tokens=%d）",
            last_error, settings.llm_model, settings.llm_max_tokens,
        )
        raise ValueError(
            f"LLM未返回有效JSON（model={settings.llm_model}）。"
            f"常见原因：1) max_tokens 过小（当前={settings.llm_max_tokens}，建议≥4096）"
            f" 2) API Key 无权限 3) 模型名称不正确。"
            f"请检查 .env 中的 LLM_MODEL/LLM_BASE_URL/LLM_MAX_TOKENS 配置。"
        )

    def parse_tender_text(self, tender_text: str) -> TenderDocument:
        """
        直接从文本解析招标信息（用于测试或已经提取好的文本）

        Args:
            tender_text: 招标文件文本内容

        Returns:
            结构化的招标文件数据
        """
        tender_text = self._apply_parse_length_limit(tender_text)

        try:
            parsed_data = self._parse_with_llm(tender_text)
            tender_doc = TenderDocument(**parsed_data)
            tender_doc = self._normalize_procurement_type(tender_doc, tender_text)

            if self._needs_technical_enrichment(tender_doc):
                tender_doc = self._enrich_package_requirements(tender_doc, tender_text)

            tender_doc = self._enrich_package_quantities(tender_doc, tender_text)
            tender_doc = self._enrich_format_templates(tender_doc, tender_text)
            logger.info(f"成功解析招标文本: {tender_doc.project_name}")
            return tender_doc
        except Exception as e:  # noqa: BLE001
            settings = get_settings()
            logger.error(f"招标文本解析失败: {e}（model={settings.llm_model}, max_tokens={settings.llm_max_tokens}）")
            raise ValueError(
                f"LLM未返回有效JSON（model={settings.llm_model}）。"
                f"常见原因：1) max_tokens 过小（当前={settings.llm_max_tokens}，建议≥4096）"
                f" 2) API Key 无权限 3) 模型名称不正确。"
            ) from e

    def extract_technical_requirements(self, tender_text: str, package_id: str, item_name=None) -> dict[str, Any]:
        """
        针对特定采购包提取详细技术要求

        Args:
            tender_text: 招标文件文本
            package_id: 采购包编号

        Returns:
            技术要求字典
        """
        prompt = ChatPromptTemplate.from_messages([
            ("system", """你是技术参数提取专家。你的任务是从招标文件中提取指定采购包的**每一条**独立技术参数。

**严格要求：**
1. 必须把每个技术参数拆成**独立的键值对**，每个参数一行。
2. **绝对禁止**输出类似 {{"核心技术参数": "详见招标文件"}} 这种笼统概括。
3. 如果招标文件中有表格形式的参数（如"激光器 ≥3个"），必须逐行提取。
4. 参数名称要精确，参数值要包含数量词、比较符号（≥、≤、不低于等）和单位。
5. 即使参数是从长段落中提取的，也必须拆分为原子级条目。

输出JSON格式，示例：
{{
    "产地": "进口",
    "检测方法": "流式细胞术",
    "激光器数量": "≥3个独立激光器",
    "荧光检测通道": "≥11色荧光检测",
    "前向散射光": "具备",
    "侧向散射光": "具备",
    "检测速度": "≥10000个事件/秒",
    "样本获取速率": "3档可调",
    "液流模式": "鞘液聚焦",
    "最小上样体积": "≤10μL",
    "绝对计数": "支持无需外加微球的绝对计数",
    "自动补偿": "全自动荧光补偿",
    "质控功能": "具备每日自动质控功能",
    "软件": "中文操作软件",
    "上样工作站": "配置自动上样工作站"
}}

如果确实没有找到任何技术参数，返回空对象 {{}}"""),
            ("user", f"请从以下招标文件中提取采购包{package_id}的全部技术参数（必须逐条列出，禁止笼统概括）：\n\n{{tender_text}}")
        ])

        llm = self._ensure_parse_tokens(self.llm)
        chain = prompt | llm
        scope_text = self._extract_package_scope(tender_text, package_id, item_name)
        if not scope_text.strip():
            logger.warning("采购包%s 未解析到稳定范围，跳过技术参数补充提取", package_id)
            return {}
        response = chain.invoke({
            "tender_text": scope_text,
            "package_id": package_id,
            "item_name": item_name,
        })
        self._check_finish_reason(response, context=f"extract_tech_pkg{package_id}")

        try:
            content = response.content.strip()
            if content.startswith("```json"):
                content = content[7:]
            if content.startswith("```"):
                content = content[3:]
            if content.endswith("```"):
                content = content[:-3]
            content = content.strip()

            raw = json.loads(content)
            if not isinstance(raw, dict):
                return {}

            forbidden_terms = _package_forbidden_terms(item_name or "")
            cleaned: dict[str, str] = {}

            for k, v in raw.items():
                key = str(k).strip()
                val = str(v).strip()
                if not key or not val:
                    continue
                if _is_bad_requirement_name(key):
                    continue
                if _is_bad_requirement_value(val):
                    continue

                raw_text = f"{key}：{val}"
                if forbidden_terms and any(tok in raw_text for tok in forbidden_terms):
                    continue

                # 回贴验证：必须能在当前包 scope 中找到
                pos, matched = _find_requirement_pair_position(scope_text, key, val)
                if pos < 0:
                    key_pos = scope_text.find(key)
                    if key_pos < 0:
                        continue

                    left = max(0, key_pos - 40)
                    right = min(len(scope_text), key_pos + len(key) + 120)
                    local_excerpt = scope_text[left:right]

                    bad_scope_hints = ("投标报价", "报价书", "预算", "履约保证金", "付款方式", "交货期")
                    if any(tok in local_excerpt for tok in bad_scope_hints):
                        continue

                    val_tokens = [t for t in _extract_match_tokens(val) if len(t) >= 2]
                    if val_tokens and not any(t in local_excerpt for t in val_tokens[:4]):
                        continue

                    cleaned[key] = val or matched

            return cleaned
        except:
            logger.warning(f"无法解析采购包{package_id}的技术要求")
            return {}


def create_tender_parser(llm: ChatOpenAI | None = None) -> TenderParser:
    """
    工厂函数：创建招标文件解析器实例

    Args:
        llm: 可选的LLM实例

    Returns:
        TenderParser实例
    """
    return TenderParser(llm)
