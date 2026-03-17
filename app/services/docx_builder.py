"""Word 文档生成服务（基于 python-docx）"""
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from app.schemas import BidDocumentSection, TenderDocument, CompanyProfile, DraftLevel
from app.services.one_click_generator.format_driven_sections.common import (
    _build_affiliated_units_statement_template,
    _build_disabled_unit_declaration_template,
    _build_hlj_supplier_qualification_commitment_template,
    _build_small_enterprise_declaration_template,
    _build_vendor_qualification_paste_section,
)


def _set_cell_bg(cell, hex_color: str) -> None:
    """设置表格单元格背景色"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _style_table(table) -> None:
    """为表格应用基础样式"""
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def _add_heading(doc: Document, text: str, level: int) -> None:
    """添加标题（1~3级），并避免标题落在页末单独成页/孤行。"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    run = heading.runs[0] if heading.runs else heading.add_run(text)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    run.font.name = "黑体"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.bold = True

    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)

    fmt = heading.paragraph_format
    fmt.keep_with_next = True      # 标题和下一段绑定，避免标题单独留页尾
    fmt.keep_together = True       # 标题自身不拆
    fmt.widow_control = True
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)


def _add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    """添加普通段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    para.paragraph_format.space_after = Pt(4)


def _append_inline_runs(para, text: str, size: Pt) -> None:
    """将含 **粗体** 的文本写入段落"""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            run = para.add_run(part)
        run.font.size = size


def _heading_body_fallback_text(title: str) -> str:
    """为只有标题没有正文的页面补一行可编辑提示，避免整页只剩标题。"""
    s = re.sub(r"\s+", "", title or "")
    if not s:
        return "请按招标文件原格式填写本节内容。"
    if "身份证正反面复印件" in s:
        return "本页用于粘贴对应身份证正反面复印件，并加盖投标人公章。"
    if "制造商授权书" in s:
        return "本页用于放置制造商授权文件原件或复印件，并按招标文件要求签章。"
    if "类似项目业绩表" in s:
        return "请按招标文件要求填写类似项目业绩，并在表后附对应证明材料。"
    if "声明函" in s:
        return "请按招标文件原格式填写声明内容，并完成签字盖章。"
    if "承诺书" in s:
        return "请按招标文件原格式填写承诺内容，并完成签字盖章。"
    if any(token in s for token in ("偏离表", "响应表", "响应及偏离表", "响应对照表")):
        return "请按招标文件要求逐项填写，不得漏项、缺项或仅复制采购要求原文。"
    if any(token in s for token in ("明细表", "报价表", "一览表", "申请表")):
        return "请按招标文件原格式填写本表内容，并保留原有列项。"
    return "请按招标文件原格式填写本节内容。"


def _append_heading_body_fallback(doc: Document, title: str) -> None:
    """在缺少正文时追加标题兜底内容。"""
    _add_paragraph(doc, _heading_body_fallback_text(title))


def _normalize_title(text: str) -> str:
    """标题归一化（用于去重匹配）"""
    return re.sub(r"[\s#`*:：、（）()\-—_]", "", text or "")

def _normalize_cover_placeholder(value: str, label: str) -> str:
    """归一化封面占位符文本。"""
    text = (value or "").strip()
    if not text:
        return f"【待填写：{label}】"
    if text.startswith("[") and text.endswith("]"):
        return f"【待填写：{label}】"
    return text


def _is_internal_draft(draft_level: DraftLevel | str | None) -> bool:
    """判断当前稿件是否为内部底稿。"""
    if isinstance(draft_level, DraftLevel):
        return draft_level == DraftLevel.internal_draft
    return str(draft_level or "").strip() == DraftLevel.internal_draft.value


def _resolve_document_date(company: CompanyProfile | None, label: str = "日期") -> str:
    """生成封面可用的文档日期文本。"""
    text = (getattr(company, "document_date", "") or "").strip() if company is not None else ""
    if not text or "待填写" in text or "待补充" in text:
        return f"【待填写：{label}】"
    return text


def _rewrite_cover_content_for_draft_level(content: str, draft_level: DraftLevel | str | None) -> str:
    """按稿件级别改写封面文本。"""
    if not _is_internal_draft(draft_level):
        return content

    replacements = {
        "投 标 文 件": "投 标 底 稿",
        "投标文件": "投标底稿",
        "响 应 文 件": "响 应 底 稿",
        "响应文件": "响应底稿",
    }
    rewritten = content
    for source, target in replacements.items():
        rewritten = rewritten.replace(source, target)
    return rewritten


def _clean_markdown_content(section_title: str, content: str) -> str:
    """
    清洗章节内容：
    1. 去除 Markdown 代码块围栏
    2. 去除与章节标题重复的首部标题
    """
    lines = content.splitlines()
    cleaned: list[str] = []
    norm_section = _normalize_title(section_title)
    can_skip_heading = True

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            continue
        if stripped.startswith(">"):
            line = re.sub(r"^>\s*", "", stripped)
            stripped = line.strip()
        if stripped.startswith("#### "):
            line = "### " + stripped[5:]
            stripped = line.strip()

        if can_skip_heading and stripped:
            raw_heading = re.sub(r"^#+\s*", "", stripped)
            norm_heading = _normalize_title(raw_heading)
            if (
                norm_heading == norm_section
                or (norm_section and norm_heading.startswith(norm_section))
                or (norm_section and norm_section.startswith(norm_heading))
            ):
                continue
            # 只在章节开头阶段尝试跳过重复标题
            if stripped and not stripped.startswith("#"):
                can_skip_heading = False

        cleaned.append(line.rstrip())

    return "\n".join(cleaned).strip()


def _get_fixed_table_widths(header_cells: list[str]):
    """按表头组合返回固定列宽配置。"""
    def _compact(text: str) -> str:
        """压缩表头文本，便于匹配固定列宽模板。"""
        return re.sub(r"\s+", "", text or "")

    key = tuple(_compact(cell) for cell in header_cells)

    if key in {
        ("序号", "技术参数项", "采购文件技术要求", "响应文件响应情况", "偏离情况"),
        ("序号", "服务名称", "磋商文件的服务需求", "响应文件响应情况", "偏离情况"),
    }:
        return [Cm(1.2), Cm(2.8), Cm(5.8), Cm(5.8), Cm(2.2)]

    if key == ("序号", "审查项", "采购文件要求", "响应文件对应内容", "是否满足", "备注"):
        return [Cm(1.0), Cm(2.2), Cm(6.0), Cm(4.8), Cm(1.8), Cm(2.0)]

    if key == ("序号", "审查内容", "合格条件", "投标文件所在页码"):
        return [Cm(1.0), Cm(3.2), Cm(8.2), Cm(3.2)]

    if key == ("序号", "审查内容", "合格条件", "投标文件对应页码"):
        return [Cm(1.0), Cm(3.2), Cm(8.2), Cm(3.2)]

    if key == ("序号", "内容", "评分因素分项", "评审标准", "投标文件对应页码"):
        return [Cm(1.0), Cm(2.2), Cm(3.0), Cm(8.0), Cm(2.8)]

    if key == ("序号", "评审项", "采购文件评分要求", "响应文件对应内容", "自评说明", "证明材料/页码"):
        return [Cm(1.0), Cm(2.2), Cm(5.4), Cm(4.4), Cm(3.2), Cm(2.2)]

    if key == ("序号", "无效情形", "自检结果", "备注"):
        return [Cm(1.0), Cm(10.5), Cm(2.4), Cm(3.0)]

    if key == ("序号", "无效投标情形", "自检结果", "备注"):
        return [Cm(1.0), Cm(10.5), Cm(2.4), Cm(3.0)]

    if key == ("招标文件条目号", "招标文件采购需求的内容与数值", "投标人的技术响应内容与数值", "技术响应偏差说明",
               "技术支持资料（或证明材料）说明"):
        return [Cm(2.0), Cm(5.2), Cm(5.2), Cm(2.4), Cm(3.2)]

    # 兼容旧模板
    if key == ("序号", "审查项", "招标文件要求", "响应情况", "对应材料/页码"):
        return [Cm(1.2), Cm(3.0), Cm(6.0), Cm(4.5), Cm(3.0)]

    return None


def _normalize_header_signature(cells: list[str]) -> tuple[str, ...]:
    """归一化表头签名。"""
    return tuple(re.sub(r"[\s\r\n\t]+", "", cell or "") for cell in cells)


def _select_table_layout_widths(
    tender: TenderDocument | None,
    header_cells: list[str],
    section_title: str = "",
):
    """选择表格布局使用的列宽。"""
    if tender is None:
        return None

    target_sig = _normalize_header_signature(header_cells)
    if not target_sig:
        return None

    section_key = re.sub(r"[\s\r\n\t]+", "", section_title or "")
    candidates = []

    for hint in getattr(tender, "table_layout_hints", []) or []:
        hint_headers = list(getattr(hint, "header_titles", []) or [])
        widths = list(getattr(hint, "column_width_cm", []) or [])
        if len(hint_headers) != len(header_cells) or len(widths) != len(header_cells):
            continue
        if _normalize_header_signature(hint_headers) != target_sig:
            continue

        source_key = re.sub(r"[\s\r\n\t]+", "", getattr(hint, "section_title", "") or "")
        score = 0
        if section_key and source_key:
            if section_key == source_key:
                score = 3
            elif section_key in source_key or source_key in section_key:
                score = 2
            else:
                shared = sum(1 for token in ("资格", "符合", "评审", "报价", "偏离", "投标", "无效") if token in section_key and token in source_key)
                score = 1 if shared else 0
        candidates.append((score, widths))

    if not candidates:
        return None

    candidates.sort(key=lambda item: item[0], reverse=True)
    return [Cm(width) for width in candidates[0][1]]


def _set_repeat_table_header(row) -> None:
    """设置 Word 表格表头在分页时重复显示。"""
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.xpath("./w:tblHeader")
    if existing:
        return
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _section_content_width_twips(doc: Document) -> int:
    """返回当前 section 可用版心宽度（twips）。"""
    section = doc.sections[-1]
    available = int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    return max(available // 635, 1)


def _length_to_twips(length) -> int:
    """把 python-docx 长度对象转换为 twips。"""
    if hasattr(length, "twips"):
        return int(getattr(length, "twips"))
    return max(int(round(int(length) / 635)), 1)


def _table_weight_for_header(header_text: str) -> float:
    """根据表头语义估算列宽权重。"""
    compact = re.sub(r"[\s\r\n\t]+", "", header_text or "")
    if any(token in compact for token in ("序号", "包号")):
        return 0.9
    if any(token in compact for token in ("页码", "页次", "是否", "结果", "分值", "数量", "单位", "单价", "合计", "报价")):
        return 1.2
    if "备注" in compact:
        return 1.4
    return 3.0


def _normalize_table_widths_twips(widths, available_width_twips: int) -> list[int]:
    """把表格列宽缩放到版心内，避免 Word 中表格横向溢出。"""
    normalized = [max(_length_to_twips(width), 1) for width in widths if width is not None]
    if not normalized:
        return []

    total = sum(normalized)
    if total <= 0:
        return []
    if total > available_width_twips:
        scale = available_width_twips / total
        normalized = [max(int(round(width * scale)), 1) for width in normalized]

    diff = available_width_twips - sum(normalized)
    if normalized:
        normalized[-1] = max(normalized[-1] + diff, 1)
    return normalized


def _fallback_table_widths_twips(header_cells: list[str], available_width_twips: int) -> list[int]:
    """在没有模板列宽时，按表头语义分配版心内宽度。"""
    weights = [_table_weight_for_header(header) for header in header_cells]
    total_weight = sum(weights) or float(len(header_cells) or 1)
    widths = [max(int(round(available_width_twips * weight / total_weight)), 1) for weight in weights]
    if widths:
        widths[-1] = max(widths[-1] + (available_width_twips - sum(widths)), 1)
    return widths


def _set_table_width(table, width_twips: int) -> None:
    """设置表格总宽度。"""
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.xpath("./w:tblW")
    tbl_w = existing[0] if existing else OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(max(width_twips, 1)))
    if not existing:
        tbl_pr.append(tbl_w)


def _set_table_fixed_layout(table) -> None:
    """强制 Word 使用固定表格布局。"""
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.xpath("./w:tblLayout")
    layout = existing[0] if existing else OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    if not existing:
        tbl_pr.append(layout)


def _set_table_grid_widths(table, width_twips: list[int]) -> None:
    """同步设置表格网格列宽，避免 Word 打开后按默认网格重算。"""
    grid = getattr(table._tbl, "tblGrid", None)
    if grid is None:
        return
    grid_cols = list(getattr(grid, "gridCol_lst", []) or [])
    for idx, grid_col in enumerate(grid_cols):
        if idx >= len(width_twips):
            break
        grid_col.w = Cm(width_twips[idx] / 567.0)


def _set_cell_width(cell, width_twips: int) -> None:
    """显式写入单元格宽度。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.xpath("./w:tcW")
    tc_w = existing[0] if existing else OxmlElement("w:tcW")
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(max(width_twips, 1)))
    if not existing:
        tc_pr.append(tc_w)


def _cell_alignment_for_header(header_text: str, *, is_header: bool) -> WD_ALIGN_PARAGRAPH:
    """为表头和正文选择单元格对齐方式。"""
    if is_header:
        return WD_ALIGN_PARAGRAPH.CENTER

    compact = re.sub(r"[\s\r\n\t]+", "", header_text or "")
    centered_tokens = (
        "序号",
        "包号",
        "页码",
        "页次",
        "是否",
        "备注",
        "分值",
        "数量",
        "单位",
        "单价",
        "合计",
        "报价",
        "结果",
    )
    if any(token in compact for token in centered_tokens):
        return WD_ALIGN_PARAGRAPH.CENTER
    return WD_ALIGN_PARAGRAPH.LEFT


def _extract_outline_items(content: str) -> list[str]:
    """从章节 Markdown 中提取目录项（主要提取二级小节）"""
    items: list[str] = []
    seen: set[str] = set()
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        candidate = ""
        if stripped.startswith("## "):
            candidate = stripped[3:].strip()
        elif re.match(r"^[一二三四五六七八九十]+、", stripped):
            candidate = stripped
        elif re.match(r"^[（(][一二三四五六七八九十]+[）)]", stripped):
            candidate = stripped

        if candidate and candidate not in seen:
            seen.add(candidate)
            items.append(candidate)

    return items[:20]


def _is_outline_heading_candidate(stripped: str, *, parenthetical: bool) -> bool:
    """判断自动编号行是否适合作为 Word 子标题，而不是普通正文。"""
    compact = re.sub(r"\s+", "", stripped or "")
    if not compact:
        return False

    if parenthetical:
        compact = re.sub(r"^[（(][一二三四五六七八九十]+[）)]", "", compact)
        max_length = 16
    else:
        compact = re.sub(r"^[一二三四五六七八九十]+、", "", compact)
        max_length = 24

    if not compact:
        return False
    if len(compact) > max_length:
        return False
    if any(token in compact for token in ("。", "；", "：", "，", "？", "！", ":", ";", "?", "!")):
        return False
    if "http" in compact or "https" in compact or "www." in compact:
        return False
    return True


def _markdown_heading_info(stripped: str) -> tuple[int, str, str] | None:
    """解析 Markdown 标题级别、标题文本及来源类型。"""
    if stripped.startswith(">"):
        stripped = re.sub(r"^>\s*", "", stripped)
    if stripped.startswith("### "):
        return 3, stripped[4:].strip(), "markdown"
    if stripped.startswith("#### "):
        return 3, stripped[5:].strip(), "markdown"
    if stripped.startswith("## "):
        return 2, stripped[3:].strip(), "markdown"
    if stripped.startswith("# "):
        return 1, stripped[2:].strip(), "markdown"
    if re.match(r"^第[一二三四五六七八九十]+章[、\s]", stripped):
        return 1, stripped, "chapter"
    if re.match(r"^[一二三四五六七八九十]+、", stripped) and _is_outline_heading_candidate(stripped, parenthetical=False):
        return 2, stripped, "outline"
    if re.match(r"^[（(][一二三四五六七八九十]+[）)]", stripped) and _is_outline_heading_candidate(stripped, parenthetical=True):
        return 3, stripped, "outline"
    return None


def _enable_update_fields_on_open(doc: Document) -> None:
    """要求 Word 打开文档时尝试更新域（含目录）。"""
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _insert_toc_field(paragraph, levels: str = "1-3") -> None:
    """
    在段落中插入 Word 目录域：
    { TOC \\o "1-3" \\h \\z \\u }

    \\o "1-3" = 收录 1~3 级标题
    \\h       = 目录项带超链接，可点击跳转
    \\z       = Web 布局中隐藏页码
    \\u       = 使用段落大纲级别
    """
    # begin
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    run._r.append(fld_begin)

    # instrText
    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f'TOC \\o "{levels}" \\h \\z \\u'
    run._r.append(instr)

    # separate
    run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    # end
    run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _add_toc(doc: Document, sections: Iterable[BidDocumentSection]) -> None:
    """添加可点击跳转的 Word 自动目录。"""
    _add_heading(doc, "目录", 1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _insert_toc_field(p, levels="1-3")

def _render_markdown_table(
    doc: Document,
    lines: list[str],
    tender: TenderDocument | None = None,
    section_title: str = "",
) -> None:
    """将 Markdown 表格渲染为 Word 表格，并优先复用招标模板列宽。"""
    data_rows = [l for l in lines if not re.match(r"^\|[-| :]+\|$", l.strip())]
    if not data_rows:
        return

    rows_cells = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        rows_cells.append(cells)

    if not rows_cells:
        return

    # 关键：以表头列数为准，不再取 max(len(row))
    header_cells = rows_cells[0]
    col_count = len(header_cells)

    normalized_rows = [header_cells]
    for row in rows_cells[1:]:
        if len(row) < col_count:
            row = row + [""] * (col_count - len(row))
        elif len(row) > col_count:
            row = row[:col_count - 1] + [" | ".join(row[col_count - 1:])]
        normalized_rows.append(row)

    table = doc.add_table(rows=len(normalized_rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    content_width_twips = _section_content_width_twips(doc)
    widths = _select_table_layout_widths(tender, header_cells, section_title) or _get_fixed_table_widths(header_cells)
    width_twips = (
        _normalize_table_widths_twips(widths, content_width_twips)
        if widths
        else _fallback_table_widths_twips(header_cells, content_width_twips)
    )
    table.autofit = False
    _style_table(table)
    _set_table_fixed_layout(table)
    _set_table_width(table, sum(width_twips))
    if table.rows:
        _set_repeat_table_header(table.rows[0])

    if width_twips:
        for col_idx, width in enumerate(width_twips):
            if col_idx < len(table.columns):
                table.columns[col_idx].width = Cm(width / 567.0)
        _set_table_grid_widths(table, width_twips)

    for i, row_data in enumerate(normalized_rows):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            clean_text = cell_text.replace("**", "")
            cell.text = clean_text

            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(clean_text)
            run.font.size = Pt(10)

            if i == 0:
                run.font.bold = True
                _set_cell_bg(cell, "D9E1F2")

            p.alignment = _cell_alignment_for_header(header_cells[j], is_header=(i == 0))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if width_twips and j < len(width_twips):
                _set_cell_width(cell, width_twips[j])
                cell.width = Cm(width_twips[j] / 567.0)

    doc.add_paragraph()


def _parse_and_render_markdown(
    doc: Document,
    content: str,
    tender: TenderDocument | None = None,
    section_title: str = "",
) -> bool:
    """
    将 Markdown 文本逐行解析并写入 Word 文档。
    支持：# 标题、**粗体**、- 列表、| 表格、普通段落
    """
    lines = content.splitlines()
    i = 0
    rendered_body = False
    seen_inner_heading = False
    body_since_heading = False
    last_heading_text = ""
    last_heading_source = ""

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行
        if not stripped:
            i += 1
            continue

        # 代码块围栏
        if stripped.startswith("```"):
            i += 1
            continue

        # 标题
        heading_info = _markdown_heading_info(stripped)
        if heading_info:
            level, heading_text, heading_source = heading_info
            if (
                seen_inner_heading
                and not body_since_heading
                and last_heading_text
                and last_heading_source != "outline"
            ):
                _append_heading_body_fallback(doc, last_heading_text)
                rendered_body = True
                body_since_heading = True
            if (
                seen_inner_heading
                and body_since_heading
                and heading_source != "outline"
                and last_heading_source != "outline"
            ):
                doc.add_page_break()
            _add_heading(doc, heading_text, level)
            seen_inner_heading = True
            body_since_heading = False
            last_heading_text = heading_text
            last_heading_source = heading_source

        # 分割线
        elif stripped.startswith("---"):
            doc.add_paragraph("─" * 40)
            rendered_body = True
            body_since_heading = True

        # 无序列表
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _append_inline_runs(para, stripped[2:], Pt(10.5))
            rendered_body = True
            body_since_heading = True

        # 有序列表
        elif re.match(r"^\d+[\.、]\s*", stripped):
            para = doc.add_paragraph(style="List Number")
            _append_inline_runs(para, re.sub(r"^\d+[\.、]\s*", "", stripped), Pt(10.5))
            rendered_body = True
            body_since_heading = True

        elif stripped in {"[PAGE_BREAK]", "[[PAGE_BREAK]]", "<PAGE_BREAK>", "\f"}:
            if (
                seen_inner_heading
                and not body_since_heading
                and last_heading_text
                and last_heading_source != "outline"
            ):
                _append_heading_body_fallback(doc, last_heading_text)
                rendered_body = True
                body_since_heading = True
            doc.add_page_break()
            body_since_heading = False

        elif re.fullmatch(r"\[\[PASTE_AREA:(\d+)]]", stripped):
            match = re.fullmatch(r"\[\[PASTE_AREA:(\d+)]]", stripped)
            blank_lines = max(3, min(int(match.group(1)), 20)) if match else 8
            for _ in range(blank_lines):
                para = doc.add_paragraph()
                run = para.add_run(" ")
                run.font.size = Pt(11)
                para.paragraph_format.space_after = Pt(0)
            rendered_body = True
            body_since_heading = True

        # Markdown 表格：收集连续的表格行一起渲染
        elif stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _render_markdown_table(doc, table_lines, tender=tender, section_title=section_title)
            rendered_body = True
            body_since_heading = True
            continue

        # 普通段落（含粗体处理）
        else:
            para = doc.add_paragraph()
            _append_inline_runs(para, stripped, Pt(11))
            para.paragraph_format.space_after = Pt(4)
            rendered_body = True
            body_since_heading = True

        i += 1

    if (
        seen_inner_heading
        and not body_since_heading
        and last_heading_text
        and last_heading_source != "outline"
    ):
        _append_heading_body_fallback(doc, last_heading_text)
        rendered_body = True

    return rendered_body


def _set_document_style(doc: Document) -> None:
    """设置文档全局样式"""
    style = doc.styles["Normal"]
    style.font.name = "仿宋"
    style.font.size = Pt(11)
    # 兼容中文字体
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")

    # 页边距
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)


def _add_cover(
    doc: Document,
    tender: TenderDocument,
    company: CompanyProfile,
    *,
    draft_level: DraftLevel | str | None = None,
) -> None:
    """生成投标文件封面"""
    def _center_line(text: str, size: int, *, bold: bool = False) -> None:
        """在封面中追加一行居中的文本。"""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "黑体" if bold else "仿宋"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if bold else "仿宋")
        p.paragraph_format.space_after = Pt(6)

    _center_line(tender.project_name or "【待填写：项目名称】", 18, bold=True)
    doc.add_paragraph()
    _center_line("投 标 底 稿" if _is_internal_draft(draft_level) else "投 标 文 件", 24, bold=True)
    doc.add_paragraph()
    _center_line(f"招标编号：{tender.project_number or '【待填写：项目编号】'}", 14)
    doc.add_paragraph()
    doc.add_paragraph()
    _center_line(f"投标单位：{_normalize_cover_placeholder(company.name, '投标人名称')}（盖章）", 14)
    _center_line(f"单位地址：{_normalize_cover_placeholder(company.address, '单位地址')}", 12)
    doc.add_paragraph()
    _center_line(_resolve_document_date(company), 12)


def _is_cover_section(section: BidDocumentSection) -> bool:
    """判断章节是否属于封面内容。"""
    title = (getattr(section, "section_title", "") or "").strip()
    if "封面" in title:
        return True
    content = _clean_markdown_content(title, getattr(section, "content", "") or "")
    compact = re.sub(r"\s+", "", content[:120])
    return "政府采购" in compact and "响应文件" in compact


def _detect_structure_mode_from_tender(tender=None, sections=None) -> str:
    """识别文档应采用的章节结构模式。"""
    titles = [getattr(s, "section_title", "") or "" for s in (sections or [])]
    text = " ".join(
        [
            str(getattr(tender, "project_number", "") or ""),
            str(getattr(tender, "procurement_type", "") or ""),
            str(getattr(tender, "project_name", "") or ""),
            " ".join(getattr(tender, "response_section_titles", []) or []),
            " ".join(titles),
        ]
    )

    if "[TP]" in text or "竞争性谈判" in text:
        return "tp"
    if "[CS]" in text or "竞争性磋商" in text:
        return "cs"
    if "[ZB]" in text or "公开招标" in text or ("招标" in text and "谈判" not in text and "磋商" not in text):
        return "zb"
    return "unknown"

def _default_zb_titles() -> list[str]:
    """返回 ZB 模式下的默认章节标题。"""
    return [
        "一、投标函",
        "二、开标一览表",
        "三、投标分项报价表",
        "四、法定代表人授权书",
        "五、资格证明文件",
        "六、商务条款响应及偏离表",
        "七、技术要求响应及偏离表",
        "八、供货、安装调试、质量保障及售后服务方案",
        "九、资格性审查响应对照表",
        "十、符合性审查响应对照表",
        "十一、详细评审响应对照表",
        "十二、无效投标情形自检表",
    ]


def _is_bad_zb_section_title(title: str) -> bool:
    """判断标题是否属于应过滤的 ZB 异常章节。"""
    s = re.sub(r"\s+", "", title or "")
    if not s:
        return True

    bad_words = (
        "项目基本情况",
        "招标公告",
        "投标人须知",
        "评标办法",
        "合同条款",
        "采购需求",
        "目录",
        "评审索引",
        "资格性检查索引",
        "符合性检查索引",
        "评分办法索引",
    )
    return any(x in s for x in bad_words)

def _is_probable_zb_template_title(title: str) -> bool:
    """判断标题是否像 ZB 原模板章节。"""
    s = re.sub(r"\s+", "", title or "")
    if not s:
        return False

    if re.match(r"^格式\s*\d+(?:-\d+)?(?:\.\d+)?", s):
        return True

    positive_words = (
        "投标函",
        "开标一览表",
        "投标分项报价表",
        "资格证明文件",
        "投标保证金说明函",
        "授权书",
        "投标人一般情况表",
        "类似项目业绩表",
        "中小企业声明函",
        "残疾人福利性单位声明函",
        "节能环保材料",
        "商务条款响应及偏离表",
        "商务条款偏离表",
        "采购需求响应及偏离表",
        "技术要求响应及偏离表",
        "技术支持资料",
        "其他技术方案",
        "供货、安装调试、质量保障及售后服务方案",
        "资格性审查响应对照表",
        "符合性审查响应对照表",
        "详细评审响应对照表",
        "无效投标情形自检表",
        "制造商授权书",
        "售后服务承诺书",
        "招标代理服务费承诺",
    )
    if any(word in s for word in positive_words):
        return True

    if _is_bad_zb_section_title(s):
        return False

    if re.match(r"^(?:7|8|9)\.\d+(?:\.\d+)?", s) and any(
        key in s for key in ("声明函", "业绩表", "授权书", "响应及偏离表", "技术方案", "证明文件", "承诺")
    ):
        return True

    return False


def _usable_exact_titles(tender=None, exact_titles=None) -> bool:
    """判断招标文件自带标题列表是否可直接使用。"""
    exact_titles = [str(x).strip() for x in (exact_titles or []) if str(x).strip()]
    mode = _detect_structure_mode_from_tender(tender=tender)

    if not exact_titles:
        return False

    if mode != "zb":
        return False

    good_titles: list[str] = []
    seen: set[str] = set()
    for title in exact_titles:
        if not _is_probable_zb_template_title(title):
            continue
        key = re.sub(r"\s+", "", title)
        if key in seen:
            continue
        seen.add(key)
        good_titles.append(title)

    return len(good_titles) >= 3


def _required_titles_for_tender(tender=None, sections=None) -> list[str]:
    """计算当前招标文件必须保留的章节标题。"""
    exact_titles = [str(x).strip() for x in (getattr(tender, "response_section_titles", []) or []) if str(x).strip()]
    if _usable_exact_titles(tender=tender, exact_titles=exact_titles):
        seen: set[str] = set()
        filtered: list[str] = []
        for title in exact_titles:
            if not _is_probable_zb_template_title(title):
                continue
            key = re.sub(r"\s+", "", title)
            if key in seen:
                continue
            seen.add(key)
            filtered.append(title)

        mode = _detect_structure_mode_from_tender(tender=tender, sections=sections)
        if mode == "zb":
            existing_titles = [
                (getattr(section, "section_title", "") or "").strip()
                for section in (sections or [])
                if (getattr(section, "section_title", "") or "").strip()
            ]

            def _has_existing(*keywords: str) -> bool:
                """判断当前章节列表中是否已经存在目标标题。"""
                return any(any(keyword in title for keyword in keywords) for title in existing_titles)

            ordered_titles: list[str] = []
            for title in (
                "资格性审查响应对照表",
                "符合性审查响应对照表",
                "详细评审响应对照表",
            ):
                if _has_existing(title):
                    ordered_titles.append(title)

            ordered_titles.extend(filtered)

            for title in (
                "供货、安装调试、质量保障及售后服务方案",
                "无效投标情形自检表",
            ):
                if _has_existing(title):
                    ordered_titles.append(title)

            deduped: list[str] = []
            seen.clear()
            for title in ordered_titles:
                key = re.sub(r"\s+", "", title)
                if key in seen:
                    continue
                seen.add(key)
                deduped.append(title)
            return deduped
        return filtered

    mode = _detect_structure_mode_from_tender(tender=tender)

    if mode == "tp":
        return [
            "一、响应文件封面格式",
            "二、报价书",
            "三、报价一览表",
            "四、资格承诺函",
            "五、技术偏离及详细配置明细表",
            "六、技术服务和售后服务的内容及措施",
            "七、法定代表人/单位负责人授权书",
            "八、法定代表人/单位负责人和授权代表身份证明",
            "九、小微企业声明函",
            "十、残疾人福利性单位声明函",
            "十一、投标人关联单位的说明",
        ]

    if mode == "cs":
        return [
            "一、响应文件封面格式",
            "二、首轮报价表",
            "三、分项报价表",
            "四、技术偏离及详细配置明细表",
            "五、技术服务和售后服务的内容及措施",
            "六、法定代表人/单位负责人授权书",
            "七、法定代表人/单位负责人和授权代表身份证明",
            "八、小微企业声明函",
            "九、残疾人福利性单位声明函",
            "十、投标人关联单位的说明",
            "十一、资格承诺函",
        ]

    if mode == "zb":
        return _default_zb_titles()

    return []

def _backfill_required_sections(sections, tender=None):
    """补齐缺失的必需章节占位内容。"""
    packages = list(getattr(tender, "packages", []) or []) if tender is not None else []
    placeholder_map = {
        "二、报价书": "【待人工补齐：按招标文件原格式填写报价书】",
        "三、报价一览表": "【待人工补齐：按招标文件原格式填写报价一览表】",
        "四、资格承诺函": _build_hlj_supplier_qualification_commitment_template(),
        "五、资格证明文件": _build_vendor_qualification_paste_section(tender, packages, ""),
        "五、技术偏离及详细配置明细表": "【待人工补齐：按采购文件逐条填写技术偏离及详细配置明细】",
        "六、技术服务和售后服务的内容及措施": "【待人工补齐：按采购文件补齐供货、安装调试、验收、售后服务方案】",
        "七、法定代表人/单位负责人授权书": "【待人工补齐：法定代表人/单位负责人授权书】",
        "八、法定代表人/单位负责人和授权代表身份证明": "【待人工补齐：法定代表人/单位负责人和授权代表身份证明】",
        "九、小微企业声明函": _build_small_enterprise_declaration_template(tender, packages),
        "十、残疾人福利性单位声明函": _build_disabled_unit_declaration_template(tender, packages),
        "十一、投标人关联单位的说明": _build_affiliated_units_statement_template(tender),
        "二、首轮报价表": "采用电子招投标的项目无需编制该表格，按投标客户端报价部分填写。",
        "三、分项报价表": "采用电子招投标的项目无需编制该表格，按投标客户端报价部分填写。",
        "四、技术偏离及详细配置明细表": "【待人工补齐：按采购文件逐条填写技术偏离及详细配置明细】",
        "五、技术服务和售后服务的内容及措施": "【待人工补齐：按评分项展开供货、安装调试、质量保证、售后服务方案】",
        "六、法定代表人/单位负责人授权书": "【待人工补齐：法定代表人/单位负责人授权书】",
        "七、法定代表人/单位负责人和授权代表身份证明": "【待人工补齐：法定代表人/单位负责人和授权代表身份证明】",
        "八、小微企业声明函": _build_small_enterprise_declaration_template(tender, packages),
        "九、残疾人福利性单位声明函": _build_disabled_unit_declaration_template(tender, packages),
        "十、投标人关联单位的说明": _build_affiliated_units_statement_template(tender),
        "十一、资格承诺函": _build_hlj_supplier_qualification_commitment_template(),
    }

    existing = {
        (getattr(s, "section_title", "") or "").strip(): s
        for s in (sections or [])
        if (getattr(s, "section_title", "") or "").strip()
    }

    ordered = _required_titles_for_tender(tender, sections=sections)
    if not ordered:
        return list(sections or [])

    ordered_set = set(ordered)
    filled = []

    for title in ordered:
        if title in existing:
            filled.append(existing[title])
        else:
            filled.append(
                BidDocumentSection(
                    section_title=title,
                    content=placeholder_map.get(title, "【待人工补齐本章节内容】"),
                )
            )

    mode = _detect_structure_mode_from_tender(tender=tender, sections=sections)
    allow_extra_keywords = (
        "资格性审查",
        "符合性审查",
        "符合性检查",
        "详细评审",
        "评分因素",
        "评分标准",
        "无效投标",
        "否决投标",
        "废标情形",
        "供货",
        "安装",
        "调试",
        "质量保障",
        "售后服务",
        "服务方案",
    )

    for s in (sections or []):
        title = (getattr(s, "section_title", "") or "").strip()
        if not title:
            continue
        if title in ordered_set:
            continue

        if mode == "zb":
            compact = re.sub(r"\s+", "", title)
            if _is_bad_zb_section_title(compact):
                continue
            if not _is_probable_zb_template_title(title) and not any(k in compact for k in allow_extra_keywords):
                continue

        filled.append(s)

    return filled

def _assert_new_structure_only(sections, tender=None) -> None:
    """校验结果中是否只保留新结构章节。"""
    titles = [
        (getattr(s, "section_title", "") or "").strip()
        for s in (sections or [])
        if (getattr(s, "section_title", "") or "").strip()
    ]

    exact_titles = [str(x).strip() for x in (getattr(tender, "response_section_titles", []) or []) if str(x).strip()]
    mode = _detect_structure_mode_from_tender(tender=tender, sections=sections)

    if _usable_exact_titles(tender=tender, exact_titles=exact_titles):
        missing = [x for x in exact_titles if x not in titles]
        if missing:
            raise RuntimeError(f"检测到招标文件第六章/响应文件格式中的必需章节缺失: {'；'.join(missing)}")
        return

    if mode == "tp":
        required = {
            "一、响应文件封面格式",
            "二、报价书",
            "三、报价一览表",
            "四、资格承诺函",
            "五、技术偏离及详细配置明细表",
            "六、技术服务和售后服务的内容及措施",
            "七、法定代表人/单位负责人授权书",
            "八、法定代表人/单位负责人和授权代表身份证明",
            "九、小微企业声明函",
            "十、残疾人福利性单位声明函",
            "十一、投标人关联单位的说明",
        }
        forbidden = {
            "一、封面格式",
            "二、首轮报价表",
            "三、分项报价表",
            "七、资格性审查响应对照表",
            "八、符合性审查响应对照表",
            "九、投标无效情形汇总及自检表",
            "七、报价书附件",
        }
    elif mode == "cs":
        required = {
            "一、响应文件封面格式",
            "二、首轮报价表",
            "三、分项报价表",
            "四、技术偏离及详细配置明细表",
            "五、技术服务和售后服务的内容及措施",
            "六、法定代表人/单位负责人授权书",
            "七、法定代表人/单位负责人和授权代表身份证明",
            "八、小微企业声明函",
            "九、残疾人福利性单位声明函",
            "十、投标人关联单位的说明",
            "十一、资格承诺函",
        }
        forbidden = {
            "一、封面格式",
            "二、报价书",
            "三、报价一览表",
            "四、资格承诺函",
            "五、详细配置明细",
            "六、技术偏离表",
            "七、报价书附件",
        }
    elif mode == "zb":
        required = set(_default_zb_titles())
        forbidden = {
            "一、项目基本情况",
            "招标公告",
            "投标人须知",
            "评标办法",
            "合同条款",
            "采购需求",
        }
    else:
        return

    missing = [x for x in required if x not in titles]
    hit = [x for x in titles if x in forbidden or _is_bad_zb_section_title(x)]

    if missing:
        raise RuntimeError(f"检测到当前采购方式对应必需章节缺失: {'；'.join(missing)}")
    if hit:
        raise RuntimeError(f"检测到当前采购方式不应出现的章节: {'；'.join(hit)}")

def build_bid_docx(
    sections: list[BidDocumentSection],
    tender: TenderDocument,
    company: CompanyProfile,
    output_path: Path,
    *,
    draft_level: DraftLevel | str | None = None,
) -> Path:
    """
    将投标文件各章节内容写入 Word (.docx) 文件。

    Args:
        sections: 各章节列表（BidDocumentSection）
        tender: 招标文件结构化数据
        company: 企业信息
        output_path: 输出 .docx 文件路径
        draft_level: 稿件等级，internal_draft 时按底稿样式输出封面

    Returns:
        输出文件路径
    """
    sections = _backfill_required_sections(sections, tender=tender)
    _assert_new_structure_only(sections, tender=tender)
    doc = Document()
    _set_document_style(doc)
    _enable_update_fields_on_open(doc)

    # 逐章节渲染（跳过自动生成的封面/目录章节，避免重复）
    skip_titles = {"封面", "目录"}
    render_sections = [section for section in sections if section.section_title not in skip_titles]

    mode = _detect_structure_mode_from_tender(tender=tender, sections=render_sections)

    if render_sections and _is_cover_section(render_sections[0]):
        cover_section = render_sections[0]
        cover_content = _rewrite_cover_content_for_draft_level(
            _clean_markdown_content(cover_section.section_title, cover_section.content),
            draft_level,
        )
        rendered_cover = (
            _parse_and_render_markdown(doc, cover_content, tender=tender, section_title=cover_section.section_title)
            if cover_content else False
        )
        if not rendered_cover:
            _append_heading_body_fallback(doc, cover_section.section_title)
        render_sections = render_sections[1:]
        if render_sections:
            doc.add_page_break()
    elif mode == "zb":
        _add_cover(doc, tender, company, draft_level=draft_level)
        if render_sections:
            doc.add_page_break()

    # 自动目录页
    if render_sections:
        _add_toc(doc, render_sections)
        doc.add_page_break()

    for idx, section in enumerate(render_sections):
        clean_content = _clean_markdown_content(section.section_title, section.content)

        # 章节标题
        _add_heading(doc, section.section_title, 1)
        # doc.add_paragraph()

        # 章节内容（Markdown → Word）
        rendered_body = (
            _parse_and_render_markdown(doc, clean_content, tender=tender, section_title=section.section_title)
            if clean_content else False
        )

        # 附件列表
        if section.attachments:
            _add_paragraph(doc, "附件：", bold=True)
            for att in section.attachments:
                doc.add_paragraph(att, style="List Bullet")
            rendered_body = True

        if not rendered_body:
            _append_heading_body_fallback(doc, section.section_title)

        if idx < len(render_sections) - 1:
            doc.add_page_break()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
