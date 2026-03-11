"""招投标正式工作流服务。"""
from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any

import pypdf
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI

from app.schemas import (
    BidDocumentSection,

    CompanyProfile,

    ProductSpecification,
    TenderDocument,

)
from app.services.one_click_generator import (
    _apply_template_pollution_guard,
    _atomize_requirements,

    _effective_requirements,
    _GENERIC_TECH_KEYS,

    generate_bid_sections,

)
from app.services.retriever import search_knowledge
from app.services.tender_parser import TenderParser

try:
    from docx import Document as _DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

_MAX_RAW_PROMPT_CHARS = 24000
_MAX_REVIEW_SECTION_CHARS = 1800
_MAX_CITATION_QUOTE_CHARS = 220
_DEFAULT_CITATION_TOP_K = 6
_PLACEHOLDER_PATTERNS = (
    "[待填写]",
    "[投标方公司名称]",
    "[法定代表人]",
    "[授权代表]",
    "[联系电话]",
    "[联系地址]",
    "[公司注册地址]",
    "[品牌型号]",
    "[生产厂家]",
    "[品牌]",
    "[待补充]",
    "（此处留空",
    "(此处留空",
)
_MEDICAL_KEYWORDS = ("医疗", "器械", "检验", "试剂", "诊断", "流式", "医院")
_SME_KEYWORDS = ("中小企业", "小微", "监狱企业", "残疾人福利性单位", "价格扣除", "声明函")
_IMPORTED_KEYWORDS = ("进口", "原装", "境外", "国外")
_CONSORTIUM_KEYWORDS = ("联合体", "共同投标")
_STAGE_STATUS_COMPLETED = "completed"
_STAGE_STATUS_WARNING = "warning"
_STAGE_STATUS_BLOCKED = "blocked"
_STAGE_STATUS_SKIPPED = "skipped"
_TEXT_SECTION_MAX_LINES = 14
_TEXT_SECTION_MAX_CHARS = 320

# ── 详细度目标（Detail Targets）──
_DETAIL_TARGETS = {
    "technical_atomic_clauses_per_package": 15,
    "deviation_table_min_rows": 10,
    "narrative_sections_min_chars": 200,
    "evidence_per_item": 1,
    "config_items_min": 5,
    "config_description_min_sentences": 1,
}

# ── 富展开模式 ──
_RICH_EXPANSION_MODE = True
_TECH_POLLUTION_KEYWORDS = (
    "评分标准",
    "评分办法",
    "评分因素",
    "商务条款",
    "合同条款",
    "违约责任",
    "质疑",
    "投诉",
)
_PLACEHOLDER_FILL_ORDER = (
    "[投标方公司名称]",
    "[法定代表人]",
    "[授权代表]",
    "[联系电话]",
    "[联系地址]",
    "[公司注册地址]",
)
_COMPARATOR_PATTERNS: tuple[tuple[str, str], ...] = (
    ("≥", r"≥\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("≤", r"≤\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    (">=", r">=\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("<=", r"<=\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("不低于", r"不低于\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("不少于", r"不少于\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("不高于", r"不高于\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("不大于", r"不大于\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("至少", r"至少\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
)
_MIN_PROVEN_COMPLETION_RATE = 0.8
_CLAUSE_BOUNDARY_CHARS = "\n。；;!?！？"
_PENDING_RESPONSE_TEXT = "待核实（未匹配到已证实产品事实）"
_UNRESOLVED_DELIVERY_MARKERS = (
    _PENDING_RESPONSE_TEXT,
    "待补证",
    "投标方证据：未绑定",
    "未匹配到投标方证据",
)

_GENERIC_VALUE_PATTERNS = (
    "详见招标文件",
    "按招标文件要求",
    "见技术参数表",
    "参见附件",
    "见采购需求",
    "详见采购文件",
    "按采购文件",
    "详见技术要求",
)


def _is_generic_value(value: str) -> bool:
    """Check if a requirement value is a generic/collapsed placeholder rather than real content."""
    v = _safe_text(value)
    if not v or len(v) < 4:
        return True
    return any(pattern in v for pattern in _GENERIC_VALUE_PATTERNS)


def _llm_call(llm: ChatOpenAI, system_prompt: str, user_prompt: str) -> str:
    response = llm.invoke([SystemMessage(system_prompt), HumanMessage(user_prompt)])
    content = response.content
    if isinstance(content, list):
        content = "\n".join(
            item.get("text", "")
            for item in content
            if isinstance(item, dict)
        )
    return str(content).strip()


def _extract_json(text: str) -> dict[str, Any]:
    raw = text.strip()
    if raw.startswith("```json"):
        raw = raw[7:]
    if raw.startswith("```"):
        raw = raw[3:]
    if raw.endswith("```"):
        raw = raw[:-3]
    raw = raw.strip()
    return json.loads(raw)


def _ensure_str_list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def _stage_record(
    stage_code: str,
    stage_name: str,
    status: str,
    summary: str,
    data: dict[str, Any] | None = None,
    issues: list[str] | None = None,
) -> dict[str, Any]:
    return {
        "stage_code": stage_code,
        "stage_name": stage_name,
        "status": status,
        "summary": summary,
        "data": data or {},
        "issues": issues or [],
    }


def _to_int_or_none(value: Any) -> int | None:
    try:
        if value is None:
            return None
        return int(value)
    except (TypeError, ValueError):
        return None


def _prepare_citations(hits: list[dict[str, Any]], limit: int = 5) -> list[dict[str, Any]]:
    citations: list[dict[str, Any]] = []
    seen: set[tuple[str, int | None, str]] = set()

    for hit in hits:
        if not isinstance(hit, dict):
            continue

        metadata = hit.get("metadata")
        if not isinstance(metadata, dict):
            metadata = {}

        source = str(metadata.get("source") or "unknown").strip() or "unknown"
        chunk_index = _to_int_or_none(metadata.get("chunk_index"))

        score: float | None = None
        raw_score = hit.get("score")
        if raw_score is not None:
            try:
                score = round(float(raw_score), 6)
            except (TypeError, ValueError):
                score = None

        quote = str(hit.get("text") or "").replace("\n", " ").strip()
        if len(quote) > _MAX_CITATION_QUOTE_CHARS:
            quote = quote[:_MAX_CITATION_QUOTE_CHARS] + "..."
        if not quote:
            continue

        dedup_key = (source, chunk_index, quote)
        if dedup_key in seen:
            continue
        seen.add(dedup_key)

        citations.append(
            {
                "source": source,
                "chunk_index": chunk_index,
                "score": score,
                "quote": quote,
            }
        )

        if len(citations) >= max(1, limit):
            break

    return citations


def _fmt_money(amount: float) -> str:
    return f"{amount:,.2f}"


def _contains_any(text: str, keywords: tuple[str, ...]) -> bool:
    return any(keyword in text for keyword in keywords)


def _safe_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text or default


def _dedupe_texts(values: list[str]) -> list[str]:
    deduped: list[str] = []
    seen: set[str] = set()
    for value in values:
        normalized = _safe_text(value)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(normalized)
    return deduped


def _fact_match_keys(*values: str) -> list[str]:
    keys: list[str] = []
    seen: set[str] = set()
    for value in values:
        normalized = _safe_text(value)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        keys.append(normalized)
    return keys


def _parameter_name_tokens(value: str) -> list[str]:
    return [
        token
        for token in re.split(r"[，,、；;：:（）()\\[\\]\\s/]+", _safe_text(value))
        if len(token) >= 2 and not token.isdigit()
    ]


def _parameter_name_matches(left: str, right: str) -> bool:
    left_text = _safe_text(left)
    right_text = _safe_text(right)
    if not left_text or not right_text:
        return False
    if left_text == right_text or left_text in right_text or right_text in left_text:
        return True
    left_tokens = _parameter_name_tokens(left_text)
    right_tokens = _parameter_name_tokens(right_text)
    if not left_tokens or not right_tokens:
        return False
    return all(token in right_text for token in left_tokens[:3]) or all(token in left_text for token in right_tokens[:3])


def _extract_fact_value_from_quote(quote: str, fact_name: str = "") -> str:
    normalized_quote = _safe_text(quote)
    normalized_fact_name = _safe_text(fact_name)
    if not normalized_quote:
        return ""

    if normalized_fact_name:
        pattern = rf"{re.escape(normalized_fact_name)}\s*[：:]\s*(.+)"
        match = re.search(pattern, normalized_quote)
        if match:
            return _safe_text(match.group(1))

    parts = re.split(r"[：:]", normalized_quote, maxsplit=1)
    if len(parts) == 2:
        return _safe_text(parts[1])
    return normalized_quote


def _fact_matches_parameter(fact: dict[str, Any], parameter_name: str) -> bool:
    normalized_parameter = _safe_text(parameter_name)
    if not normalized_parameter:
        return False

    candidates = _fact_match_keys(
        _safe_text(fact.get("fact_name")),
        _safe_text(fact.get("evidence_type")),
        *[str(item) for item in fact.get("match_keys", []) if str(item).strip()],
    )
    for candidate in candidates:
        if (
            _parameter_name_matches(candidate, normalized_parameter)
            or candidate in normalized_parameter
            or normalized_parameter in candidate
        ):
            return True
    return False


def _lookup_package_fact_value(
    package_facts: dict[str, Any],
    parameter_name: str,
) -> tuple[str, str, str]:
    if not package_facts:
        return "", "", ""

    fact_rows: list[dict[str, Any]] = []
    for field in ("technical_facts", "identity_facts", "evidence_materials", "offered_facts"):
        for item in package_facts.get(field, []):
            if isinstance(item, dict):
                fact_rows.append(item)

    for fact in fact_rows:
        if not _fact_matches_parameter(fact, parameter_name):
            continue
        source = _safe_text(fact.get("evidence_source"), "投标方资料")
        quote = _safe_text(fact.get("evidence_quote"))
        value = _safe_text(fact.get("fact_value") or fact.get("evidence_value"))
        fact_name = _safe_text(fact.get("fact_name") or fact.get("evidence_type"), parameter_name)
        if not value:
            value = _extract_fact_value_from_quote(quote, fact_name)
        if value:
            return value, source, quote or f"{fact_name}：{value}"

    return "", "", ""


def _lookup_product_spec_value(product: ProductSpecification, parameter_name: str) -> str:
    normalized = _safe_text(parameter_name)
    if not normalized:
        return ""

    specs = product.specifications or {}
    if normalized in specs:
        return _safe_text(specs[normalized])

    short_name = normalized.split("：", 1)[0].strip()
    if short_name in specs:
        return _safe_text(specs[short_name])

    for key, value in specs.items():
        key_text = _safe_text(key)
        if not key_text:
            continue
        if key_text in normalized or normalized in key_text:
            return _safe_text(value)

    key_tokens = [token for token in re.split(r"[，,、；;：:（）()\\[\\]\\s/]+", short_name) if len(token) >= 2]
    for key, value in specs.items():
        key_text = _safe_text(key)
        if key_tokens and all(token in key_text or token in normalized for token in key_tokens[:3]):
            return _safe_text(value)

    return ""


def _resolve_bidder_evidence(
    requirement: str,
    company: CompanyProfile | None,
    products: dict[str, ProductSpecification],
    selected_packages: list[str],
) -> tuple[bool, str, str]:
    normalized = _safe_text(requirement)

    if company and _contains_any(normalized, ("营业执照", "许可证", "资质")):
        if company.licenses:
            license_names = "；".join(item.license_type for item in company.licenses[:3])
            return True, "企业证照", license_names
        if company.name.strip():
            return True, "企业主体信息", company.name

    if company and _contains_any(normalized, ("授权", "法定代表人")):
        if company.staff:
            return True, "人员/授权信息", company.staff[0].name
        if company.legal_representative.strip():
            return True, "法定代表人信息", company.legal_representative

    if company and _contains_any(normalized, ("社保", "税收")) and company.social_insurance_proof.strip():
        return True, "企业社保/税务材料", company.social_insurance_proof

    if company and _contains_any(normalized, ("截图", "信用")) and company.credit_check_time is not None:
        return True, "信用查询记录", company.credit_check_time.isoformat()

    for pkg_id in selected_packages:
        product = products.get(pkg_id)
        if product is None:
            continue

        if _contains_any(normalized, ("注册证", "备案证")) and product.registration_number.strip():
            return True, f"包{pkg_id} 注册证", product.registration_number
        if _contains_any(normalized, ("授权", "授权书")) and product.authorization_letter.strip():
            return True, f"包{pkg_id} 授权文件", product.authorization_letter

        for spec_key, spec_val in (product.specifications or {}).items():
            if spec_key and _parameter_name_matches(spec_key, normalized):
                return (
                    True,
                    f"包{pkg_id} 产品参数",
                    f"投标产品{product.product_name}的{spec_key}为{spec_val}，满足招标要求",
                )

        if product.model.strip() and product.model in normalized:
            return True, f"包{pkg_id} 产品型号", product.model

        candidate_facts = [
            {
                "fact_name": "产品名称",
                "fact_value": product.product_name,
                "evidence_source": f"包{pkg_id} 产品名称",
                "evidence_quote": f"产品名称：{product.product_name}",
                "match_keys": _fact_match_keys("产品名称", "货物名称"),
            },
            {
                "fact_name": "型号",
                "fact_value": product.model,
                "evidence_source": f"包{pkg_id} 产品型号",
                "evidence_quote": f"型号：{product.model}",
                "match_keys": _fact_match_keys("型号", "规格型号", "品牌型号"),
            },
            {
                "fact_name": "生产厂家",
                "fact_value": product.manufacturer,
                "evidence_source": f"包{pkg_id} 生产厂家",
                "evidence_quote": f"生产厂家：{product.manufacturer}",
                "match_keys": _fact_match_keys("生产厂家", "厂家", "制造商"),
            },
            {
                "fact_name": "原产地",
                "fact_value": product.origin,
                "evidence_source": f"包{pkg_id} 原产地",
                "evidence_quote": f"原产地：{product.origin}",
                "match_keys": _fact_match_keys("原产地", "产地", "来源地"),
            },
        ]
        if product.price > 0:
            candidate_facts.append(
                {
                    "fact_name": "单价",
                    "fact_value": _fmt_money(product.price),
                    "evidence_source": f"包{pkg_id} 报价信息",
                    "evidence_quote": f"单价：{_fmt_money(product.price)}",
                    "match_keys": _fact_match_keys("单价", "报价", "报价信息"),
                }
            )
        for certification in product.certifications:
            candidate_facts.append(
                {
                    "fact_name": "认证证书",
                    "fact_value": certification,
                    "evidence_source": f"包{pkg_id} 认证证书",
                    "evidence_quote": f"认证证书：{certification}",
                    "match_keys": _fact_match_keys("认证证书", "认证", "证书"),
                }
            )

        for candidate in candidate_facts:
            if _fact_matches_requirement_text(candidate, normalized):
                return True, _safe_text(candidate.get("evidence_source")), _safe_text(candidate.get("evidence_quote"))

        capability_words = ("具备", "支持", "提供", "配备", "配置", "满足", "可", "能够", "兼容")
        if product.product_name.strip() and any(keyword in normalized for keyword in capability_words):
            return (
                True,
                f"包{pkg_id} 产品能力推断",
                f"投标产品（{product.product_name}）具备该功能，满足招标要求",
            )

    return False, "未匹配到投标方证据", "需人工补充企业证照、产品参数或授权链"


def _fact_matches_requirement_text(fact: dict[str, Any], requirement_text: str) -> bool:
    normalized_requirement = _safe_text(requirement_text)
    if not normalized_requirement:
        return False

    if _fact_matches_parameter(fact, normalized_requirement):
        return True

    value = _safe_text(fact.get("fact_value") or fact.get("evidence_value"))
    quote = _safe_text(fact.get("evidence_quote"))
    if value and value in normalized_requirement:
        return True
    if quote and quote in normalized_requirement:
        return True
    return False


def _truncate_text(text: str, limit: int = _TEXT_SECTION_MAX_CHARS) -> str:
    normalized = re.sub(r"\s+", " ", _safe_text(text))
    if len(normalized) <= limit:
        return normalized
    return normalized[:limit] + "..."


def _load_document_units(file_path: str | Path | None) -> tuple[str, list[dict[str, Any]]]:
    if not file_path:
        return "", []

    path = Path(file_path)
    suffix = path.suffix.lower()
    units: list[dict[str, Any]] = []

    if suffix == ".pdf":
        with path.open("rb") as handle:
            reader = pypdf.PdfReader(handle)
            char_start = 0
            for idx, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                char_end = char_start + len(text)
                units.append(
                    {
                        "unit_type": "page",
                        "unit_index": idx,
                        "char_start": char_start,
                        "char_end": char_end,
                        "text": text,
                    }
                )
                char_start = char_end + 1
        return "pdf", units

    if suffix in {".docx", ".doc"} and _DOCX_AVAILABLE:
        document = _DocxDocument(str(path))
        char_start = 0
        unit_index = 1
        for paragraph in document.paragraphs:
            text = paragraph.text or ""
            if not text.strip():
                continue
            char_end = char_start + len(text)
            units.append(
                {
                    "unit_type": "paragraph",
                    "unit_index": unit_index,
                    "char_start": char_start,
                    "char_end": char_end,
                    "text": text,
                }
            )
            char_start = char_end + 1
            unit_index += 1
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if not row_text:
                    continue
                char_end = char_start + len(row_text)
                units.append(
                    {
                        "unit_type": "table_row",
                        "unit_index": unit_index,
                        "char_start": char_start,
                        "char_end": char_end,
                        "text": row_text,
                    }
                )
                char_start = char_end + 1
                unit_index += 1
        return "docx", units

    return suffix.lstrip("."), []


def _build_document_ingestion_view(
    raw_text: str,
    file_path: str | Path | None,
    tender_id: str,
) -> dict[str, Any]:
    try:
        source_format, units = _load_document_units(file_path)
    except Exception as exc:  # noqa: BLE001
        logger.warning("文档接入保真信息构建失败，回退到纯文本模式：%s", exc)
        source_format, units = "", []
    sample_segments: list[dict[str, Any]] = []
    if units:
        for unit in units[:_TEXT_SECTION_MAX_LINES]:
            sample_segments.append(
                {
                    "unit_type": unit["unit_type"],
                    "unit_index": unit["unit_index"],
                    "char_start": unit["char_start"],
                    "char_end": unit["char_end"],
                    "excerpt": _truncate_text(str(unit["text"]), 180),
                }
            )
    else:
        char_cursor = 0
        for idx, line in enumerate(raw_text.splitlines(), start=1):
            stripped = line.strip()
            if not stripped:
                char_cursor += len(line) + 1
                continue
            segment = {
                "unit_type": "line",
                "unit_index": idx,
                "char_start": char_cursor,
                "char_end": char_cursor + len(line),
                "excerpt": _truncate_text(stripped, 180),
            }
            sample_segments.append(segment)
            char_cursor += len(line) + 1
            if len(sample_segments) >= _TEXT_SECTION_MAX_LINES:
                break

    table_like_units = sum(
        1
        for item in units
        if "|" in _safe_text(item.get("text")) or "\t" in _safe_text(item.get("text"))
    )
    if not units:
        table_like_units = sum(
            1
            for line in raw_text.splitlines()
            if "|" in line or "\t" in line
        )

    summary = (
        f"文档接入完成，source=tender::{tender_id}，原文 {len(raw_text)} 字；"
        f"识别 {len(units) or len(sample_segments)} 个结构单元，表格样式行 {table_like_units} 条。"
    )
    return {
        "source_id": f"tender::{tender_id}",
        "source_format": source_format or "text",
        "raw_text_length": len(raw_text),
        "unit_count": len(units) or len(sample_segments),
        "page_or_block_count": len(units) if units else None,
        "table_like_unit_count": table_like_units,
        "sample_segments": sample_segments,
        "summary": summary,
    }


def _build_package_segmentation_view(
    tender: TenderDocument,
    raw_text: str,
    selected_packages: list[str],
) -> dict[str, Any]:
    target_package_ids = selected_packages or [pkg.package_id for pkg in tender.packages]
    package_views: list[dict[str, Any]] = []

    for pkg in tender.packages:
        if pkg.package_id not in target_package_ids:
            continue

        scope_text = TenderParser._extract_package_scope(raw_text, pkg.package_id, pkg.item_name)
        package_views.append(
            {
                "package_id": pkg.package_id,
                "item_name": pkg.item_name,
                "quantity": pkg.quantity,
                "budget": pkg.budget,
                "technical_requirement_count": len(pkg.technical_requirements or {}),
                "scope_length": len(scope_text),
                "scope_excerpt": _truncate_text(scope_text, 220),
                "is_isolated": bool(scope_text.strip()),
            }
        )

    missing_scope = [item["package_id"] for item in package_views if not item["is_isolated"]]
    summary = (
        f"包件切分完成，目标包 {len(package_views)} 个，"
        f"{'全部完成单包上下文切分' if not missing_scope else f'包{','.join(missing_scope)} 需人工补充定位'}。"
    )
    return {
        "selected_packages": target_package_ids,
        "packages": package_views,
        "missing_scope_packages": missing_scope,
        "summary": summary,
    }


def _parse_threshold(value: str) -> tuple[str, str, str]:
    normalized = _safe_text(value)
    for comparator, pattern in _COMPARATOR_PATTERNS:
        match = re.search(pattern, normalized)
        if match:
            threshold = _safe_text(match.group("threshold"))
            unit = _safe_text(match.group("unit"))
            return comparator, threshold, unit
    return "", "", ""


def _first_numeric_value(value: str) -> float | None:
    match = re.search(r"-?\d+(?:\.\d+)?", _safe_text(value))
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def _compare_numeric_requirement(comparator: str, threshold: float, response_value: float) -> bool | None:
    if comparator in {"≥", ">=", "不少于", "不低于", "至少"}:
        return response_value >= threshold
    if comparator in {"≤", "<=", "不高于", "不大于"}:
        return response_value <= threshold
    return None


def _evaluate_requirement_response(requirement_value: str, response_value: str) -> dict[str, Any]:
    normalized_requirement = _safe_text(requirement_value)
    normalized_response = _safe_text(response_value)
    if not normalized_response:
        return {
            "deviation_status": "待核实",
            "verified": False,
            "comparison_mode": "missing",
            "reason": "未匹配到可响应的产品事实",
        }

    comparator, threshold, _ = _parse_threshold(normalized_requirement)
    threshold_value = _first_numeric_value(threshold or normalized_requirement)
    response_numeric = _first_numeric_value(normalized_response)
    if comparator and threshold_value is not None and response_numeric is not None:
        compare_result = _compare_numeric_requirement(comparator, threshold_value, response_numeric)
        if compare_result:
            return {
                "deviation_status": "无偏离",
                "verified": True,
                "comparison_mode": "numeric",
                "reason": "已完成数值门槛校验",
            }
        if compare_result is False:
            return {
                "deviation_status": "有偏离",
                "verified": True,
                "comparison_mode": "numeric",
                "reason": "产品事实已证实，但数值未满足招标门槛",
            }

    compact_requirement = re.sub(r"\s+", "", normalized_requirement)
    compact_response = re.sub(r"\s+", "", normalized_response)
    if compact_requirement and compact_response:
        if compact_requirement == compact_response or compact_response in compact_requirement or compact_requirement in compact_response:
            return {
                "deviation_status": "无偏离",
                "verified": True,
                "comparison_mode": "text",
                "reason": "已完成文本级事实比对",
            }

    return {
        "deviation_status": "待核实",
        "verified": False,
        "comparison_mode": "ambiguous",
        "reason": "已匹配到产品事实，但未形成可自动判定的无偏离结论",
    }


def _looks_material_requirement(text: str) -> bool:
    normalized = _safe_text(text)
    return any(
        token in normalized
        for token in ("必须", "须", "应", "不得", "严禁", "≥", "≤", ">=", "<=", "不少于", "不低于", "至少")
    )


def _qualification_requirement_type(item: str) -> str:
    normalized = _safe_text(item)
    if _contains_any(normalized, ("执照", "许可证", "注册证", "备案凭证", "证书")):
        return "license"
    if _contains_any(normalized, ("声明", "承诺")):
        return "declaration"
    if _contains_any(normalized, ("授权",)):
        return "authorization"
    if _contains_any(normalized, ("截图", "查询")):
        return "screenshot"
    return "attachment"


def _workflow_context_text(tender: TenderDocument) -> str:
    parts = [
        tender.project_name,
        tender.project_number,
        tender.purchaser,
        tender.agency,
        tender.procurement_type,
        tender.special_requirements,
        " ".join(pkg.item_name for pkg in tender.packages),
        " ".join(" ".join(str(v) for v in pkg.technical_requirements.values()) for pkg in tender.packages),
        " ".join(f"{k}:{v}" for k, v in tender.evaluation_criteria.items()),
    ]
    return " ".join(str(part).strip() for part in parts if str(part).strip())


def _snippet_around(text: str, keyword: str, radius: int = 56) -> str:
    if not text.strip() or not keyword.strip():
        return ""

    lowered = text.lower()
    pos = lowered.find(keyword.lower())
    if pos < 0:
        return ""

    start = max(0, pos - radius)
    end = min(len(text), pos + len(keyword) + radius)
    snippet = text[start:end].replace("\n", " ").replace("\r", " ")
    snippet = re.sub(r"\s+", " ", snippet).strip()
    if len(snippet) > 160:
        snippet = snippet[:160] + "..."
    return snippet


def _extract_quoteable_segment(text: str, keyword: str) -> str:
    if not text.strip() or not keyword.strip():
        return ""

    lowered = text.lower()
    keyword_lower = keyword.lower()
    pos = lowered.find(keyword_lower)
    if pos < 0:
        return ""

    end_pos = pos + len(keyword)
    left_boundary = 0
    for marker in _CLAUSE_BOUNDARY_CHARS:
        boundary_pos = text.rfind(marker, 0, pos)
        if boundary_pos >= 0:
            left_boundary = max(left_boundary, boundary_pos + 1)

    right_candidates = [len(text)]
    for marker in _CLAUSE_BOUNDARY_CHARS:
        boundary_pos = text.find(marker, end_pos)
        if boundary_pos >= 0:
            right_candidates.append(boundary_pos)
    right_boundary = min(right_candidates)

    segment = text[left_boundary:right_boundary].strip()
    segment = re.sub(r"\s+", " ", segment)
    if 4 <= len(segment) <= 180:
        return segment
    return _snippet_around(text, keyword)


def _evidence_candidates(item: str) -> list[str]:
    candidates: list[str] = []

    def _append(candidate: str) -> None:
        normalized = candidate.strip()
        if len(normalized) < 2:
            return
        if normalized not in candidates:
            candidates.append(normalized)

    _append(item)
    compact = re.sub(r"[，,、；;：:（）()\[\]【】\s]+", "", item)
    _append(compact)

    parts = re.split(r"[：:]", item, maxsplit=1)
    for part in parts:
        _append(part)

    for token in re.split(r"[，,、；;：:（）()\[\]【】\s]+", item):
        _append(token)

    return candidates


def _locate_evidence_snippet(text: str, item: str) -> str:
    for candidate in _evidence_candidates(item):
        snippet = _extract_quoteable_segment(text, candidate)
        if snippet:
            return snippet
    return ""


def _branch_decision(
    decision_name: str,
    decision: str,
    basis: str,
    clause_type: str,
) -> dict[str, str]:
    return {
        "decision_name": decision_name,
        "decision": decision,
        "basis": basis,
        "clause_type": clause_type,
    }


_SERVICE_KEYWORDS = ("售后", "培训", "维修", "维保", "保修", "服务", "响应时间", "上门", "巡检", "技术支持")
_ACCEPTANCE_KEYWORDS = ("验收", "安装调试", "试运行", "到货验收", "终验", "初验")
_CONFIG_KEYWORDS = ("配置", "配件", "附件", "选配", "标配", "随机", "清单")


def _is_service_clause(text: str) -> bool:
    return any(kw in text for kw in _SERVICE_KEYWORDS)


def _is_acceptance_clause(text: str) -> bool:
    return any(kw in text for kw in _ACCEPTANCE_KEYWORDS)


def _is_config_clause(text: str) -> bool:
    return any(kw in text for kw in _CONFIG_KEYWORDS)


def _is_service_or_acceptance_clause(text: str) -> bool:
    return _is_service_clause(text) or _is_acceptance_clause(text) or _is_config_clause(text)
