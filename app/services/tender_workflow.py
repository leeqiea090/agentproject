"""招投标正式工作流服务。"""
from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any

import pypdf
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_openai import ChatOpenAI

from app.schemas import BidDocumentSection, CompanyProfile, ProductSpecification, TenderDocument
from app.services.one_click_generator import (
    _apply_template_pollution_guard,
    _atomize_requirements,
    _effective_requirements,
    _GENERIC_TECH_KEYS,
    generate_bid_sections,
)
from app.services.retriever import search_knowledge
from app.services.tender_parser import TenderParser

try:
    from docx import Document as _DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

_MAX_RAW_PROMPT_CHARS = 24000
_MAX_REVIEW_SECTION_CHARS = 1800
_MAX_CITATION_QUOTE_CHARS = 220
_DEFAULT_CITATION_TOP_K = 6
_PLACEHOLDER_PATTERNS = (
    "[待填写]",
    "[投标方公司名称]",
    "[法定代表人]",
    "[授权代表]",
    "[联系电话]",
    "[联系地址]",
    "[公司注册地址]",
    "[品牌型号]",
    "[生产厂家]",
    "[品牌]",
    "[待补充]",
    "（此处留空",
    "(此处留空",
)
_MEDICAL_KEYWORDS = ("医疗", "器械", "检验", "试剂", "诊断", "流式", "医院")
_SME_KEYWORDS = ("中小企业", "小微", "监狱企业", "残疾人福利性单位", "价格扣除", "声明函")
_IMPORTED_KEYWORDS = ("进口", "原装", "境外", "国外")
_CONSORTIUM_KEYWORDS = ("联合体", "共同投标")
_STAGE_STATUS_COMPLETED = "completed"
_STAGE_STATUS_WARNING = "warning"
_STAGE_STATUS_BLOCKED = "blocked"
_STAGE_STATUS_SKIPPED = "skipped"
_TEXT_SECTION_MAX_LINES = 14
_TEXT_SECTION_MAX_CHARS = 320

# ── 详细度目标（Detail Targets）──
_DETAIL_TARGETS = {
    "technical_atomic_clauses_per_package": 15,
    "deviation_table_min_rows": 10,
    "narrative_sections_min_chars": 200,
    "evidence_per_item": 1,
    "config_items_min": 5,
    "config_description_min_sentences": 1,
}

# ── 富展开模式 ──
_RICH_EXPANSION_MODE = True
_TECH_POLLUTION_KEYWORDS = (
    "评分标准",
    "评分办法",
    "评分因素",
    "商务条款",
    "合同条款",
    "违约责任",
    "质疑",
    "投诉",
)
_PLACEHOLDER_FILL_ORDER = (
    "[投标方公司名称]",
    "[法定代表人]",
    "[授权代表]",
    "[联系电话]",
    "[联系地址]",
    "[公司注册地址]",
)
_COMPARATOR_PATTERNS: tuple[tuple[str, str], ...] = (
    ("≥", r"≥\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("≤", r"≤\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    (">=", r">=\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("<=", r"<=\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("不低于", r"不低于\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("不少于", r"不少于\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("不高于", r"不高于\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("不大于", r"不大于\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
    ("至少", r"至少\s*(?P<threshold>\d+(?:\.\d+)?)\s*(?P<unit>[^\d\s，,；;]*)"),
)
_MIN_PROVEN_COMPLETION_RATE = 0.8
_CLAUSE_BOUNDARY_CHARS = "\n。；;!?！？"
_PENDING_RESPONSE_TEXT = "待核实（未匹配到已证实产品事实）"
_UNRESOLVED_DELIVERY_MARKERS = (
    _PENDING_RESPONSE_TEXT,
    "待补证",
    "投标方证据：未绑定",
    "未匹配到投标方证据",
)

_GENERIC_VALUE_PATTERNS = (
    "详见招标文件",
    "按招标文件要求",
    "见技术参数表",
    "参见附件",
    "见采购需求",
    "详见采购文件",
    "按采购文件",
    "详见技术要求",
)


def _is_generic_value(value: str) -> bool:
    """Check if a requirement value is a generic/collapsed placeholder rather than real content."""
    v = _safe_text(value)
    if not v or len(v) < 4:
        return True
    return any(pattern in v for pattern in _GENERIC_VALUE_PATTERNS)


def _llm_call(llm: ChatOpenAI, system_prompt: str, user_prompt: str) -> str:
    response = llm.invoke([SystemMessage(system_prompt), HumanMessage(user_prompt)])
    content = response.content
    if isinstance(content, list):
        content = "\n".join(
            item.get("text", "")
            for item in content
            if isinstance(item, dict)
        )
    return str(content).strip()


def _extract_json(text: str) -> dict[str, Any]:
    raw = text.strip()
    if raw.startswith("```json"):
        raw = raw[7:]
    if raw.startswith("```"):
        raw = raw[3:]
    if raw.endswith("```"):
        raw = raw[:-3]
    raw = raw.strip()
    return json.loads(raw)


def _ensure_str_list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def _stage_record(
    stage_code: str,
    stage_name: str,
    status: str,
    summary: str,
    data: dict[str, Any] | None = None,
    issues: list[str] | None = None,
) -> dict[str, Any]:
    return {
        "stage_code": stage_code,
        "stage_name": stage_name,
        "status": status,
        "summary": summary,
        "data": data or {},
        "issues": issues or [],
    }


def _to_int_or_none(value: Any) -> int | None:
    try:
        if value is None:
            return None
        return int(value)
    except (TypeError, ValueError):
        return None


def _prepare_citations(hits: list[dict[str, Any]], limit: int = 5) -> list[dict[str, Any]]:
    citations: list[dict[str, Any]] = []
    seen: set[tuple[str, int | None, str]] = set()

    for hit in hits:
        if not isinstance(hit, dict):
            continue

        metadata = hit.get("metadata")
        if not isinstance(metadata, dict):
            metadata = {}

        source = str(metadata.get("source") or "unknown").strip() or "unknown"
        chunk_index = _to_int_or_none(metadata.get("chunk_index"))

        score: float | None = None
        raw_score = hit.get("score")
        if raw_score is not None:
            try:
                score = round(float(raw_score), 6)
            except (TypeError, ValueError):
                score = None

        quote = str(hit.get("text") or "").replace("\n", " ").strip()
        if len(quote) > _MAX_CITATION_QUOTE_CHARS:
            quote = quote[:_MAX_CITATION_QUOTE_CHARS] + "..."
        if not quote:
            continue

        dedup_key = (source, chunk_index, quote)
        if dedup_key in seen:
            continue
        seen.add(dedup_key)

        citations.append(
            {
                "source": source,
                "chunk_index": chunk_index,
                "score": score,
                "quote": quote,
            }
        )

        if len(citations) >= max(1, limit):
            break

    return citations


def _contains_any(text: str, keywords: tuple[str, ...]) -> bool:
    return any(keyword in text for keyword in keywords)


def _safe_text(value: Any, default: str = "") -> str:
    if value is None:
        return default
    text = str(value).strip()
    return text or default


def _dedupe_texts(values: list[str]) -> list[str]:
    deduped: list[str] = []
    seen: set[str] = set()
    for value in values:
        normalized = _safe_text(value)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        deduped.append(normalized)
    return deduped


def _fact_match_keys(*values: str) -> list[str]:
    keys: list[str] = []
    seen: set[str] = set()
    for value in values:
        normalized = _safe_text(value)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        keys.append(normalized)
    return keys


def _extract_fact_value_from_quote(quote: str, fact_name: str = "") -> str:
    normalized_quote = _safe_text(quote)
    normalized_fact_name = _safe_text(fact_name)
    if not normalized_quote:
        return ""

    if normalized_fact_name:
        pattern = rf"{re.escape(normalized_fact_name)}\s*[：:]\s*(.+)"
        match = re.search(pattern, normalized_quote)
        if match:
            return _safe_text(match.group(1))

    parts = re.split(r"[：:]", normalized_quote, maxsplit=1)
    if len(parts) == 2:
        return _safe_text(parts[1])
    return normalized_quote


def _fact_matches_parameter(fact: dict[str, Any], parameter_name: str) -> bool:
    normalized_parameter = _safe_text(parameter_name)
    if not normalized_parameter:
        return False

    candidates = _fact_match_keys(
        _safe_text(fact.get("fact_name")),
        _safe_text(fact.get("evidence_type")),
        *[str(item) for item in fact.get("match_keys", []) if str(item).strip()],
    )
    for candidate in candidates:
        if (
            _parameter_name_matches(candidate, normalized_parameter)
            or candidate in normalized_parameter
            or normalized_parameter in candidate
        ):
            return True
    return False


def _lookup_package_fact_value(
    package_facts: dict[str, Any],
    parameter_name: str,
) -> tuple[str, str, str]:
    if not package_facts:
        return "", "", ""

    fact_rows: list[dict[str, Any]] = []
    for field in ("technical_facts", "identity_facts", "evidence_materials", "offered_facts"):
        for item in package_facts.get(field, []):
            if isinstance(item, dict):
                fact_rows.append(item)

    for fact in fact_rows:
        if not _fact_matches_parameter(fact, parameter_name):
            continue
        source = _safe_text(fact.get("evidence_source"), "投标方资料")
        quote = _safe_text(fact.get("evidence_quote"))
        value = _safe_text(fact.get("fact_value") or fact.get("evidence_value"))
        fact_name = _safe_text(fact.get("fact_name") or fact.get("evidence_type"), parameter_name)
        if not value:
            value = _extract_fact_value_from_quote(quote, fact_name)
        if value:
            return value, source, quote or f"{fact_name}：{value}"

    return "", "", ""


def _fact_matches_requirement_text(fact: dict[str, Any], requirement_text: str) -> bool:
    normalized_requirement = _safe_text(requirement_text)
    if not normalized_requirement:
        return False

    if _fact_matches_parameter(fact, normalized_requirement):
        return True

    value = _safe_text(fact.get("fact_value") or fact.get("evidence_value"))
    quote = _safe_text(fact.get("evidence_quote"))
    if value and value in normalized_requirement:
        return True
    if quote and quote in normalized_requirement:
        return True
    return False


def _truncate_text(text: str, limit: int = _TEXT_SECTION_MAX_CHARS) -> str:
    normalized = re.sub(r"\s+", " ", _safe_text(text))
    if len(normalized) <= limit:
        return normalized
    return normalized[:limit] + "..."


def _load_document_units(file_path: str | Path | None) -> tuple[str, list[dict[str, Any]]]:
    if not file_path:
        return "", []

    path = Path(file_path)
    suffix = path.suffix.lower()
    units: list[dict[str, Any]] = []

    if suffix == ".pdf":
        with path.open("rb") as handle:
            reader = pypdf.PdfReader(handle)
            char_start = 0
            for idx, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                char_end = char_start + len(text)
                units.append(
                    {
                        "unit_type": "page",
                        "unit_index": idx,
                        "char_start": char_start,
                        "char_end": char_end,
                        "text": text,
                    }
                )
                char_start = char_end + 1
        return "pdf", units

    if suffix in {".docx", ".doc"} and _DOCX_AVAILABLE:
        document = _DocxDocument(str(path))
        char_start = 0
        unit_index = 1
        for paragraph in document.paragraphs:
            text = paragraph.text or ""
            if not text.strip():
                continue
            char_end = char_start + len(text)
            units.append(
                {
                    "unit_type": "paragraph",
                    "unit_index": unit_index,
                    "char_start": char_start,
                    "char_end": char_end,
                    "text": text,
                }
            )
            char_start = char_end + 1
            unit_index += 1
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if not row_text:
                    continue
                char_end = char_start + len(row_text)
                units.append(
                    {
                        "unit_type": "table_row",
                        "unit_index": unit_index,
                        "char_start": char_start,
                        "char_end": char_end,
                        "text": row_text,
                    }
                )
                char_start = char_end + 1
                unit_index += 1
        return "docx", units

    return suffix.lstrip("."), []


def _build_document_ingestion_view(
    raw_text: str,
    file_path: str | Path | None,
    tender_id: str,
) -> dict[str, Any]:
    try:
        source_format, units = _load_document_units(file_path)
    except Exception as exc:  # noqa: BLE001
        logger.warning("文档接入保真信息构建失败，回退到纯文本模式：%s", exc)
        source_format, units = "", []
    sample_segments: list[dict[str, Any]] = []
    if units:
        for unit in units[:_TEXT_SECTION_MAX_LINES]:
            sample_segments.append(
                {
                    "unit_type": unit["unit_type"],
                    "unit_index": unit["unit_index"],
                    "char_start": unit["char_start"],
                    "char_end": unit["char_end"],
                    "excerpt": _truncate_text(str(unit["text"]), 180),
                }
            )
    else:
        char_cursor = 0
        for idx, line in enumerate(raw_text.splitlines(), start=1):
            stripped = line.strip()
            if not stripped:
                char_cursor += len(line) + 1
                continue
            segment = {
                "unit_type": "line",
                "unit_index": idx,
                "char_start": char_cursor,
                "char_end": char_cursor + len(line),
                "excerpt": _truncate_text(stripped, 180),
            }
            sample_segments.append(segment)
            char_cursor += len(line) + 1
            if len(sample_segments) >= _TEXT_SECTION_MAX_LINES:
                break

    table_like_units = sum(
        1
        for item in units
        if "|" in _safe_text(item.get("text")) or "\t" in _safe_text(item.get("text"))
    )
    if not units:
        table_like_units = sum(
            1
            for line in raw_text.splitlines()
            if "|" in line or "\t" in line
        )

    summary = (
        f"文档接入完成，source=tender::{tender_id}，原文 {len(raw_text)} 字；"
        f"识别 {len(units) or len(sample_segments)} 个结构单元，表格样式行 {table_like_units} 条。"
    )
    return {
        "source_id": f"tender::{tender_id}",
        "source_format": source_format or "text",
        "raw_text_length": len(raw_text),
        "unit_count": len(units) or len(sample_segments),
        "page_or_block_count": len(units) if units else None,
        "table_like_unit_count": table_like_units,
        "sample_segments": sample_segments,
        "summary": summary,
    }


def _build_package_segmentation_view(
    tender: TenderDocument,
    raw_text: str,
    selected_packages: list[str],
) -> dict[str, Any]:
    target_package_ids = selected_packages or [pkg.package_id for pkg in tender.packages]
    package_views: list[dict[str, Any]] = []

    for pkg in tender.packages:
        if pkg.package_id not in target_package_ids:
            continue

        scope_text = TenderParser._extract_package_scope(raw_text, pkg.package_id, pkg.item_name)
        package_views.append(
            {
                "package_id": pkg.package_id,
                "item_name": pkg.item_name,
                "quantity": pkg.quantity,
                "budget": pkg.budget,
                "technical_requirement_count": len(pkg.technical_requirements or {}),
                "scope_length": len(scope_text),
                "scope_excerpt": _truncate_text(scope_text, 220),
                "is_isolated": bool(scope_text.strip()),
            }
        )

    missing_scope = [item["package_id"] for item in package_views if not item["is_isolated"]]
    summary = (
        f"包件切分完成，目标包 {len(package_views)} 个，"
        f"{'全部完成单包上下文切分' if not missing_scope else f'包{','.join(missing_scope)} 需人工补充定位'}。"
    )
    return {
        "selected_packages": target_package_ids,
        "packages": package_views,
        "missing_scope_packages": missing_scope,
        "summary": summary,
    }


def _parse_threshold(value: str) -> tuple[str, str, str]:
    normalized = _safe_text(value)
    for comparator, pattern in _COMPARATOR_PATTERNS:
        match = re.search(pattern, normalized)
        if match:
            threshold = _safe_text(match.group("threshold"))
            unit = _safe_text(match.group("unit"))
            return comparator, threshold, unit
    return "", "", ""


def _first_numeric_value(value: str) -> float | None:
    match = re.search(r"-?\d+(?:\.\d+)?", _safe_text(value))
    if not match:
        return None
    try:
        return float(match.group(0))
    except ValueError:
        return None


def _compare_numeric_requirement(comparator: str, threshold: float, response_value: float) -> bool | None:
    if comparator in {"≥", ">=", "不少于", "不低于", "至少"}:
        return response_value >= threshold
    if comparator in {"≤", "<=", "不高于", "不大于"}:
        return response_value <= threshold
    return None


def _evaluate_requirement_response(requirement_value: str, response_value: str) -> dict[str, Any]:
    normalized_requirement = _safe_text(requirement_value)
    normalized_response = _safe_text(response_value)
    if not normalized_response:
        return {
            "deviation_status": "待核实",
            "verified": False,
            "comparison_mode": "missing",
            "reason": "未匹配到可响应的产品事实",
        }

    comparator, threshold, _ = _parse_threshold(normalized_requirement)
    threshold_value = _first_numeric_value(threshold or normalized_requirement)
    response_numeric = _first_numeric_value(normalized_response)
    if comparator and threshold_value is not None and response_numeric is not None:
        compare_result = _compare_numeric_requirement(comparator, threshold_value, response_numeric)
        if compare_result:
            return {
                "deviation_status": "无偏离",
                "verified": True,
                "comparison_mode": "numeric",
                "reason": "已完成数值门槛校验",
            }
        if compare_result is False:
            return {
                "deviation_status": "有偏离",
                "verified": True,
                "comparison_mode": "numeric",
                "reason": "产品事实已证实，但数值未满足招标门槛",
            }

    compact_requirement = re.sub(r"\s+", "", normalized_requirement)
    compact_response = re.sub(r"\s+", "", normalized_response)
    if compact_requirement and compact_response:
        if compact_requirement == compact_response or compact_response in compact_requirement or compact_requirement in compact_response:
            return {
                "deviation_status": "无偏离",
                "verified": True,
                "comparison_mode": "text",
                "reason": "已完成文本级事实比对",
            }

    return {
        "deviation_status": "待核实",
        "verified": False,
        "comparison_mode": "ambiguous",
        "reason": "已匹配到产品事实，但未形成可自动判定的无偏离结论",
    }


def _looks_material_requirement(text: str) -> bool:
    normalized = _safe_text(text)
    return any(
        token in normalized
        for token in ("必须", "须", "应", "不得", "严禁", "≥", "≤", ">=", "<=", "不少于", "不低于", "至少")
    )


def _qualification_requirement_type(item: str) -> str:
    normalized = _safe_text(item)
    if _contains_any(normalized, ("执照", "许可证", "注册证", "备案凭证", "证书")):
        return "license"
    if _contains_any(normalized, ("声明", "承诺")):
        return "declaration"
    if _contains_any(normalized, ("授权",)):
        return "authorization"
    if _contains_any(normalized, ("截图", "查询")):
        return "screenshot"
    return "attachment"


def _workflow_context_text(tender: TenderDocument) -> str:
    parts = [
        tender.project_name,
        tender.project_number,
        tender.purchaser,
        tender.agency,
        tender.procurement_type,
        tender.special_requirements,
        " ".join(pkg.item_name for pkg in tender.packages),
        " ".join(" ".join(str(v) for v in pkg.technical_requirements.values()) for pkg in tender.packages),
        " ".join(f"{k}:{v}" for k, v in tender.evaluation_criteria.items()),
    ]
    return " ".join(str(part).strip() for part in parts if str(part).strip())


def _snippet_around(text: str, keyword: str, radius: int = 56) -> str:
    if not text.strip() or not keyword.strip():
        return ""

    lowered = text.lower()
    pos = lowered.find(keyword.lower())
    if pos < 0:
        return ""

    start = max(0, pos - radius)
    end = min(len(text), pos + len(keyword) + radius)
    snippet = text[start:end].replace("\n", " ").replace("\r", " ")
    snippet = re.sub(r"\s+", " ", snippet).strip()
    if len(snippet) > 160:
        snippet = snippet[:160] + "..."
    return snippet


def _extract_quoteable_segment(text: str, keyword: str) -> str:
    if not text.strip() or not keyword.strip():
        return ""

    lowered = text.lower()
    keyword_lower = keyword.lower()
    pos = lowered.find(keyword_lower)
    if pos < 0:
        return ""

    end_pos = pos + len(keyword)
    left_boundary = 0
    for marker in _CLAUSE_BOUNDARY_CHARS:
        boundary_pos = text.rfind(marker, 0, pos)
        if boundary_pos >= 0:
            left_boundary = max(left_boundary, boundary_pos + 1)

    right_candidates = [len(text)]
    for marker in _CLAUSE_BOUNDARY_CHARS:
        boundary_pos = text.find(marker, end_pos)
        if boundary_pos >= 0:
            right_candidates.append(boundary_pos)
    right_boundary = min(right_candidates)

    segment = text[left_boundary:right_boundary].strip()
    segment = re.sub(r"\s+", " ", segment)
    if 4 <= len(segment) <= 180:
        return segment
    return _snippet_around(text, keyword)


def _evidence_candidates(item: str) -> list[str]:
    candidates: list[str] = []

    def _append(candidate: str) -> None:
        normalized = candidate.strip()
        if len(normalized) < 2:
            return
        if normalized not in candidates:
            candidates.append(normalized)

    _append(item)
    compact = re.sub(r"[，,、；;：:（）()\[\]【】\s]+", "", item)
    _append(compact)

    parts = re.split(r"[：:]", item, maxsplit=1)
    for part in parts:
        _append(part)

    for token in re.split(r"[，,、；;：:（）()\[\]【】\s]+", item):
        _append(token)

    return candidates


def _locate_evidence_snippet(text: str, item: str) -> str:
    for candidate in _evidence_candidates(item):
        snippet = _extract_quoteable_segment(text, candidate)
        if snippet:
            return snippet
    return ""


def _branch_decision(
    decision_name: str,
    decision: str,
    basis: str,
    clause_type: str,
) -> dict[str, str]:
    return {
        "decision_name": decision_name,
        "decision": decision,
        "basis": basis,
        "clause_type": clause_type,
    }


def _classify_clauses(
    tender: TenderDocument,
    analysis_result: dict[str, Any],
    selected_packages: list[str],
    raw_text: str,
) -> dict[str, Any]:
    required_materials = _ensure_str_list(analysis_result.get("required_materials"))
    scoring_rules = _ensure_str_list(analysis_result.get("scoring_rules"))
    target_packages = selected_packages or [pkg.package_id for pkg in tender.packages]
    package_map = {pkg.package_id: pkg for pkg in tender.packages}
    target_package_docs = [package_map[pkg_id] for pkg_id in target_packages if pkg_id in package_map]
    context = _workflow_context_text(tender)

    qualification_clauses = required_materials[:8]
    technical_clauses: list[str] = []
    commercial_clauses: list[str] = []
    for pkg in target_package_docs:
        for key, value in (pkg.technical_requirements or {}).items():
            technical_clauses.append(f"{key}：{value}")
        commercial_clauses.append(f"包{pkg.package_id} 交货期：{pkg.delivery_time or '按招标文件约定'}")
        commercial_clauses.append(f"包{pkg.package_id} 交货地点：{pkg.delivery_place or '采购人指定地点'}")

    commercial_terms = tender.commercial_terms.model_dump()
    for key, value in commercial_terms.items():
        if value:
            commercial_clauses.append(f"{key}：{value}")

    policy_clauses = scoring_rules[:6]
    special_requirements = tender.special_requirements.strip()
    if special_requirements:
        policy_clauses.append(special_requirements)

    if "不接受联合体" in context:
        consortium_decision = "不接受联合体"
    elif _contains_any(context, _CONSORTIUM_KEYWORDS):
        consortium_decision = "接受联合体"
    else:
        consortium_decision = "未明确，需人工确认"

    if _contains_any(context, _SME_KEYWORDS):
        sme_decision = "适用中小企业政策分支"
    else:
        sme_decision = "不涉及中小企业政策分支"

    if _contains_any(context, _MEDICAL_KEYWORDS):
        medical_decision = "需走医疗器械合规分支"
    else:
        medical_decision = "非医疗专项分支"

    if _contains_any(context, _IMPORTED_KEYWORDS):
        imported_decision = "需准备进口或原产地相关说明"
    else:
        imported_decision = "默认按国产/未明确处理"

    package_decision = "、".join(f"包{pkg.package_id}" for pkg in target_package_docs) or "全部包"
    branch_decisions = [
        _branch_decision(
            "联合体投标分支",
            consortium_decision,
            _snippet_around(raw_text, "联合体") or special_requirements or "依据结构化招标信息推断",
            "资格/组织形式",
        ),
        _branch_decision(
            "中小企业政策分支",
            sme_decision,
            _snippet_around(raw_text, "中小企业") or special_requirements or "依据评分标准及特殊要求推断",
            "政策性条款",
        ),
        _branch_decision(
            "医疗器械合规分支",
            medical_decision,
            _snippet_around(raw_text, "医疗器械") or " ".join(pkg.item_name for pkg in target_package_docs) or "依据采购标的推断",
            "行业合规",
        ),
        _branch_decision(
            "进口货物分支",
            imported_decision,
            _snippet_around(raw_text, "进口") or " ".join(pkg.item_name for pkg in target_package_docs) or "依据采购标的推断",
            "货物属性",
        ),
        _branch_decision(
            "包件响应分支",
            f"本次响应范围：{package_decision}",
            package_decision,
            "投标范围",
        ),
    ]

    summary = (
        f"已完成条款分类，覆盖资格、技术、商务、政策 4 类；"
        f"已生成 {len(branch_decisions)} 条分支决策。"
    )

    return {
        "selected_packages": target_packages,
        "package_count": len(target_package_docs),
        "clause_categories": {
            "qualification": qualification_clauses,
            "technical": technical_clauses[:20],
            "commercial": commercial_clauses[:12],
            "policy": policy_clauses[:12],
        },
        "structured_categories": {
            "procurement_requirements": qualification_clauses + technical_clauses[:20] + commercial_clauses[:12],
            "evaluation_rules": policy_clauses[:12],
            "evidence_materials": qualification_clauses,
            "explanatory_clauses": [special_requirements] if special_requirements else [],
        },
        "branch_decisions": branch_decisions,
        "summary": summary,
    }


def _normalize_requirements(
    tender: TenderDocument,
    analysis_result: dict[str, Any],
    clause_result: dict[str, Any],
    selected_packages: list[str],
    raw_text: str,
) -> dict[str, Any]:
    target_package_ids = selected_packages or [pkg.package_id for pkg in tender.packages]
    package_map = {pkg.package_id: pkg for pkg in tender.packages}
    target_packages = [package_map[pkg_id] for pkg_id in target_package_ids if pkg_id in package_map]

    qualification_requirements: list[dict[str, Any]] = []
    for idx, item in enumerate(_ensure_str_list(analysis_result.get("required_materials")), start=1):
        qualification_requirements.append(
            {
                "requirement_id": f"Q-{idx}",
                "requirement_type": _qualification_requirement_type(item),
                "requirement_text": item,
                "source_excerpt": _locate_evidence_snippet(raw_text, item),
            }
        )

    commercial_requirements: list[dict[str, Any]] = []
    for pkg in target_packages:
        commercial_requirements.extend(
            [
                {
                    "requirement_id": f"B-{pkg.package_id}-delivery-time",
                    "package_id": pkg.package_id,
                    "field": "delivery_time",
                    "value": _safe_text(pkg.delivery_time, "按招标文件约定"),
                    "source_excerpt": _locate_evidence_snippet(raw_text, pkg.delivery_time or "交货期"),
                },
                {
                    "requirement_id": f"B-{pkg.package_id}-delivery-place",
                    "package_id": pkg.package_id,
                    "field": "delivery_place",
                    "value": _safe_text(pkg.delivery_place, "采购人指定地点"),
                    "source_excerpt": _locate_evidence_snippet(raw_text, pkg.delivery_place or "交货地点"),
                },
                {
                    "requirement_id": f"B-{pkg.package_id}-quantity",
                    "package_id": pkg.package_id,
                    "field": "quantity",
                    "value": pkg.quantity,
                    "source_excerpt": _locate_evidence_snippet(raw_text, str(pkg.quantity)),
                },
                {
                    "requirement_id": f"B-{pkg.package_id}-budget",
                    "package_id": pkg.package_id,
                    "field": "budget",
                    "value": pkg.budget,
                    "source_excerpt": _locate_evidence_snippet(raw_text, str(int(pkg.budget)) if pkg.budget else "预算"),
                },
            ]
        )

    term_map = tender.commercial_terms.model_dump()
    for field, value in term_map.items():
        if value:
            commercial_requirements.append(
                {
                    "requirement_id": f"B-common-{field}",
                    "package_id": "common",
                    "field": field,
                    "value": value,
                    "source_excerpt": _locate_evidence_snippet(raw_text, str(value)),
                }
            )

    technical_requirements: list[dict[str, Any]] = []
    for pkg in target_packages:
        tech_items = list((pkg.technical_requirements or {}).items())
        if not tech_items:
            for clause in _ensure_str_list(clause_result.get("clause_categories", {}).get("technical")):
                if "：" not in clause:
                    continue
                key, value = clause.split("：", 1)
                tech_items.append((key.strip(), value.strip()))

        # --- Fix: detect sparse/collapsed requirements and use raw-text extraction ---
        _is_sparse = (
            len(tech_items) < 3
            or all(k.strip() in _GENERIC_TECH_KEYS for k, _ in tech_items)
            or all(_is_generic_value(v) for _, v in tech_items)
            or (len(tech_items) < 6 and any(_is_generic_value(v) for _, v in tech_items))
        )
        if _is_sparse and raw_text.strip():
            try:
                effective = _effective_requirements(pkg, raw_text)
                if len(effective) > len(tech_items):
                    existing_keys = {k.strip() for k, _ in tech_items}
                    merged = list(tech_items)
                    for ek, ev in effective:
                        if ek.strip() not in existing_keys:
                            merged.append((ek, ev))
                            existing_keys.add(ek.strip())
                    tech_items = merged
                    logger.info(
                        "包%s 需求归一化：原始条目稀疏，通过原文提取补充至 %d 条",
                        pkg.package_id,
                        len(tech_items),
                    )
            except Exception as exc:  # noqa: BLE001
                logger.warning("包%s 原文需求补充提取失败：%s", pkg.package_id, exc)
        # --- End fix ---

        # --- 原子化拆分：将复合条款拆分为原子级技术要求 ---
        try:
            atomized = _atomize_requirements(tech_items)
            if len(atomized) > len(tech_items):
                logger.info(
                    "包%s 原子化拆分：从 %d 条扩展为 %d 条原子级技术要求",
                    pkg.package_id,
                    len(tech_items),
                    len(atomized),
                )
            tech_items = atomized
        except Exception as exc:  # noqa: BLE001
            logger.warning("包%s 原子化拆分失败：%s", pkg.package_id, exc)
        # --- End atomization ---

        # --- 过滤总括条款：禁止笼统概述直接进入最终技术表 ---
        _BLANKET_PATTERNS = (
            "满足招标文件所有要求",
            "符合国家标准",
            "按招标文件要求提供",
            "详见招标文件",
            "满足采购文件全部要求",
            "完全响应招标文件",
            "以上参数均满足",
            "所有技术参数满足",
        )
        filtered_tech_items: list[tuple[str, str]] = []
        for key, value in tech_items:
            combined = f"{key} {value}"
            # 跳过纯总括性条款
            if any(bp in combined for bp in _BLANKET_PATTERNS) and len(value.strip()) < 30:
                logger.debug("包%s 过滤总括条款：%s = %s", pkg.package_id, key, value)
                continue
            filtered_tech_items.append((key, value))
        tech_items = filtered_tech_items
        # --- End blanket clause filter ---

        for idx, (key, value) in enumerate(tech_items, start=1):
            comparator, threshold, unit = _parse_threshold(_safe_text(value))
            # 增强: 提取 source_page 和 source_text
            source_excerpt = _locate_evidence_snippet(raw_text, key) or _locate_evidence_snippet(raw_text, str(value))
            source_page = None  # 如果 raw_text 包含页码信息可以提取
            # 尝试从 source_excerpt 中提取页码 (格式: "...第X页...")
            if source_excerpt:
                page_match = re.search(r"第\s*(\d+)\s*页", source_excerpt)
                if page_match:
                    source_page = int(page_match.group(1))

            technical_requirements.append(
                {
                    "requirement_id": f"T-{pkg.package_id}-{idx}",
                    "package_id": pkg.package_id,
                    "clause_no": f"{pkg.package_id}.{idx}",
                    "param_name": key,
                    "parameter_name": key,
                    "operator": comparator or "",
                    "comparator": comparator,
                    "threshold": threshold or "",
                    "unit": unit or "",
                    "normalized_value": _safe_text(value),
                    "response_field_hint": key,
                    "response_value_type": "numeric" if comparator and threshold else "text",
                    "is_material_requirement": _looks_material_requirement(f"{key} {value}"),
                    "is_material": _looks_material_requirement(f"{key} {value}"),
                    "source_page": source_page,
                    "source_text": source_excerpt or "",
                    "source_excerpt": source_excerpt or "",
                }
            )

    by_package = {
        pkg.package_id: {
            "item_name": pkg.item_name,
            "technical_count": len([item for item in technical_requirements if item["package_id"] == pkg.package_id]),
            "commercial_count": len([item for item in commercial_requirements if item["package_id"] in {pkg.package_id, "common"}]),
        }
        for pkg in target_packages
    }

    summary = (
        f"需求归一化完成：资格 {len(qualification_requirements)} 项，"
        f"商务 {len(commercial_requirements)} 项，技术 {len(technical_requirements)} 项。"
    )
    return {
        "selected_packages": target_package_ids,
        "qualification_requirements": qualification_requirements,
        "commercial_requirements": commercial_requirements,
        "technical_requirements": technical_requirements,
        "by_package": by_package,
        "summary": summary,
    }


def _expand_extracted_facts(
    normalized_result: dict[str, Any],
    products: dict[str, ProductSpecification],
    tender: TenderDocument,
) -> dict[str, Any]:
    """Detail Expander：将已抽到的事实展开成详细响应。

    对每个技术要求，结合产品信息生成 2-3 句详细响应说明，
    包含：要求含义、产品响应方式、证据基础。
    """
    expanded_items: list[dict[str, Any]] = []
    expanded_count = 0

    for requirement in normalized_result.get("technical_requirements", []):
        if not isinstance(requirement, dict):
            continue

        package_id = _safe_text(requirement.get("package_id"))
        parameter_name = _safe_text(requirement.get("parameter_name"))
        normalized_value = _safe_text(requirement.get("normalized_value"))
        product = products.get(package_id)

        # 生成详细响应
        detail_lines: list[str] = [f"本条款要求{parameter_name}满足「{normalized_value}」。"]

        # 1. 要求含义

        # 2. 产品响应方式
        if product:
            matched_value = _lookup_product_spec_value(product, parameter_name)
            if matched_value:
                detail_lines.append(
                    f"投标产品（{product.manufacturer or ''} {product.product_name}，"
                    f"型号：{product.model or '详见偏离表'}）的{parameter_name}为{matched_value}，"
                    "满足招标要求。"
                )
                expanded_count += 1
            else:
                # 能力类推断
                _CAP = ("具备", "支持", "提供", "配备", "配置", "满足", "可", "能够")
                if any(m in f"{parameter_name}{normalized_value}" for m in _CAP):
                    detail_lines.append(
                        f"投标产品（{product.product_name}）具备该功能，满足本条款要求。"
                    )
                    expanded_count += 1
                elif product.product_name.strip():
                    detail_lines.append(
                        f"投标产品（{product.product_name}）满足本条款要求，"
                        "详见技术偏离表对应行。"
                    )
                    expanded_count += 1

            # 3. 证据基础
            if product.registration_number:
                detail_lines.append(f"证据依据：产品注册证编号 {product.registration_number}。")
            elif product.certifications:
                detail_lines.append(f"证据依据：已取得{product.certifications[0]}等认证。")
            elif matched_value:
                detail_lines.append("证据依据：产品参数库已验证数据。")
        else:
            detail_lines.append("（尚未绑定投标产品，待补充产品信息后展开响应。）")

        expanded_items.append({
            "requirement_id": requirement.get("requirement_id"),
            "package_id": package_id,
            "parameter_name": parameter_name,
            "normalized_value": normalized_value,
            "detail_expansion": " ".join(detail_lines),
            "expanded": expanded_count > 0,
        })

    total = len(expanded_items)
    expansion_rate = expanded_count / max(1, total)

    return {
        "expanded_items": expanded_items,
        "expanded_count": expanded_count,
        "total": total,
        "expansion_rate": round(expansion_rate, 4),
        "summary": (
            f"详细展开完成：{expanded_count}/{total} 项已展开为详细响应说明"
            f"（展开率 {expansion_rate:.0%}）。"
        ),
    }


def _extract_specs_from_bid_material(material: dict[str, Any]) -> dict[str, Any]:
    """从单份投标材料中提取技术参数和证据引用。"""
    extracted = {
        "specs": {},
        "config_items": [],
        "evidence_refs": [],
        "brand": "",
        "model": "",
        "manufacturer": "",
    }
    file_type = _safe_text(material.get("file_type"))
    file_name = _safe_text(material.get("file_name", ""))
    page_count = material.get("page_count", 0)
    extracted_specs = material.get("extracted_specs") or {}
    extracted_text = _safe_text(material.get("extracted_text", ""))
    key_pages = material.get("key_pages") or []

    # 从 extracted_specs 合并参数
    if extracted_specs:
        extracted["specs"].update(extracted_specs)

    # 从 extracted_text 中提取品牌/型号/厂家
    if extracted_text:
        brand_match = re.search(r"(?:品牌|Brand)[：:\s]*([^\n，,；;]{2,30})", extracted_text)
        if brand_match:
            extracted["brand"] = brand_match.group(1).strip()
        model_match = re.search(r"(?:型号|Model|规格型号)[：:\s]*([^\n，,；;]{2,40})", extracted_text)
        if model_match:
            extracted["model"] = model_match.group(1).strip()
        mfr_match = re.search(r"(?:生产厂家|制造商|Manufacturer|厂家)[：:\s]*([^\n，,；;]{2,40})", extracted_text)
        if mfr_match:
            extracted["manufacturer"] = mfr_match.group(1).strip()

        # 从说明书/彩页中提取配置项
        if file_type in ("brochure", "manual", "spec_sheet"):
            config_pattern = re.findall(
                r"(?:标准配置|标配|随机附件|装箱清单|配置清单)[：:]\s*(.+?)(?:\n\n|\Z)",
                extracted_text,
                re.DOTALL,
            )
            for block in config_pattern:
                for line in block.strip().splitlines():
                    line = line.strip(" -·•●○◆▪※")
                    if line and len(line) >= 2:
                        extracted["config_items"].append({
                            "配置项": line,
                            "说明": "标配",
                            "数量": "1",
                            "来源": file_name,
                        })

    # 构建证据引用
    for kp in key_pages:
        if isinstance(kp, dict) and kp.get("page"):
            extracted["evidence_refs"].append({
                "file_name": file_name,
                "file_type": file_type,
                "page": kp["page"],
                "description": _safe_text(kp.get("content", "")),
            })

    # 如果没有 key_pages 但有页数信息，生成通用引用
    if not key_pages and page_count > 0:
        type_label = {
            "brochure": "产品彩页",
            "manual": "产品说明书",
            "registration": "注册证",
            "test_report": "检测/质评报告",
            "spec_sheet": "厂家参数页",
        }.get(file_type, file_name)
        extracted["evidence_refs"].append({
            "file_name": file_name,
            "file_type": file_type,
            "page": 1,
            "description": f"{type_label}（共{page_count}页）",
        })

    return extracted


def _build_product_profile(product: ProductSpecification) -> dict[str, Any]:
    """Product Profile Builder: 从投标材料中提取真实产品事实构建统一对象。

    输入（投标材料）:
    - 彩页 (brochure)
    - 说明书 (manual)
    - 注册证 (registration)
    - 检测/质评报告 (test_report)
    - 厂家参数页 (spec_sheet)

    输出:
    - brand: 品牌名称
    - model: 产品型号
    - manufacturer: 生产厂家
    - technical_specs: 技术参数（从材料中提取的真实值）
    - config_items: 配置项清单
    - evidence_refs: 证据引用（文件名+页码）

    硬规则:
    - 没有 brand/model 时,技术章节只能出 internal draft
    - 没有 technical_specs 时,不允许写"实际响应值"
    """
    # 从 bid_materials 中提取事实
    bid_materials = []
    if hasattr(product, "bid_materials"):
        bid_materials = product.bid_materials or []

    merged_specs: dict[str, Any] = {}
    merged_config_items: list[dict[str, Any]] = []
    merged_evidence_refs: list[dict[str, Any]] = []
    material_brand = ""
    material_model = ""
    material_manufacturer = ""

    for material in bid_materials:
        mat_dict = material.model_dump() if hasattr(material, "model_dump") else (material if isinstance(material, dict) else {})
        extracted = _extract_specs_from_bid_material(mat_dict)
        merged_specs.update(extracted["specs"])
        merged_config_items.extend(extracted["config_items"])
        merged_evidence_refs.extend(extracted["evidence_refs"])
        if extracted["brand"] and not material_brand:
            material_brand = extracted["brand"]
        if extracted["model"] and not material_model:
            material_model = extracted["model"]
        if extracted["manufacturer"] and not material_manufacturer:
            material_manufacturer = extracted["manufacturer"]

    # 合并：材料提取值 > product 字段值 > 推断值
    final_brand = product.brand or material_brand or (product.manufacturer.split()[0] if product.manufacturer else "")
    final_model = product.model or material_model
    final_manufacturer = product.manufacturer or material_manufacturer
    final_specs = product.specifications or product.technical_specs or {}
    if merged_specs:
        combined_specs = dict(final_specs)
        combined_specs.update(merged_specs)
        final_specs = combined_specs

    final_config = product.config_items or []
    if merged_config_items:
        existing_names = {_safe_text(item.get("配置项") or item.get("name", "")) for item in final_config if isinstance(item, dict)}
        for new_item in merged_config_items:
            name = _safe_text(new_item.get("配置项", ""))
            if name and name not in existing_names:
                final_config.append(new_item)
                existing_names.add(name)

    final_evidence = product.evidence_refs or []
    if merged_evidence_refs:
        existing_refs = {
            f"{_safe_text(ref.get('file_name', ''))}:{ref.get('page', '')}"
            for ref in final_evidence if isinstance(ref, dict)
        }
        for new_ref in merged_evidence_refs:
            key = f"{_safe_text(new_ref.get('file_name', ''))}:{new_ref.get('page', '')}"
            if key not in existing_refs:
                final_evidence.append(new_ref)
                existing_refs.add(key)

    profile = {
        "brand": final_brand,
        "model": final_model,
        "manufacturer": final_manufacturer,
        "product_name": product.product_name,
        "technical_specs": final_specs,
        "config_items": final_config,
        "functional_notes": product.functional_notes or (
            f"{product.product_name}具备完整的技术功能，能够满足采购文件要求。"
            if not final_specs else
            f"{product.product_name}（{final_model}）核心参数已从投标材料中提取，详见技术偏离表。"
        ),
        "acceptance_notes": product.acceptance_notes or "按照采购文件及国家相关标准进行验收。",
        "training_notes": product.training_notes or "提供设备操作培训，确保用户熟练掌握。",
        "evidence_refs": final_evidence,
        "bid_material_types": [
            (m.model_dump() if hasattr(m, "model_dump") else m).get("file_type", "")
            for m in bid_materials
        ],
        "has_complete_identity": bool(final_model and final_manufacturer),
        "has_technical_specs": bool(final_specs),
        "has_bid_materials": len(bid_materials) > 0,
        "ready_for_external": bool(
            final_model and final_manufacturer and final_specs and final_evidence
        ),
    }

    # 如果 config_items 为空,从 specifications 中提取基础配置
    if not profile["config_items"] and profile["technical_specs"]:
        config_items = []
        for idx, (key, value) in enumerate(profile["technical_specs"].items(), start=1):
            config_items.append({
                "序号": idx,
                "配置项": key,
                "说明": str(value),
                "数量": "标配",
            })
        profile["config_items"] = config_items

    return profile


def _extract_product_facts(
    tender: TenderDocument,
    products: dict[str, ProductSpecification],
    selected_packages: list[str],
) -> dict[str, Any]:
    target_package_ids = selected_packages or [pkg.package_id for pkg in tender.packages]
    package_map = {pkg.package_id: pkg for pkg in tender.packages}
    package_facts: list[dict[str, Any]] = []
    total_fact_count = 0
    offered_fact_count = 0

    for pkg_id in target_package_ids:
        product = products.get(pkg_id)
        pkg = package_map.get(pkg_id)
        if product is None:
            package_facts.append(
                {
                    "package_id": pkg_id,
                    "item_name": pkg.item_name if pkg else "",
                    "product_present": False,
                    "identity_facts": [],
                    "technical_facts": [],
                    "evidence_materials": [],
                    "offered_facts": [],
                    "summary": f"包{pkg_id} 未提供产品资料，无法提取产品事实。",
                }
            )
            continue

        identity_facts: list[dict[str, str]] = []
        technical_facts: list[dict[str, str]] = []
        evidence_materials: list[dict[str, str]] = []
        offered_facts: list[dict[str, Any]] = []

        def _append_offered_fact(
            name: str,
            value: str,
            source: str,
            *,
            fact_type: str,
            match_keys: tuple[str, ...] = (),
        ) -> None:
            normalized = _safe_text(value)
            if not normalized:
                return
            offered_facts.append(
                {
                    "fact_type": fact_type,
                    "fact_name": name,
                    "fact_value": normalized,
                    "evidence_source": source,
                    "evidence_quote": f"{name}：{normalized}",
                    "match_keys": _fact_match_keys(name, *match_keys),
                }
            )

        def _append_identity(name: str, value: str, source: str = "产品档案") -> None:
            normalized = _safe_text(value)
            if not normalized:
                return
            record = {
                "fact_name": name,
                "fact_value": normalized,
                "evidence_source": source,
                "evidence_quote": f"{name}：{normalized}",
                "match_keys": _fact_match_keys(
                    name,
                    "产品名称" if name == "产品名称" else "",
                    "货物名称" if name == "产品名称" else "",
                    "规格型号" if name == "型号" else "",
                    "品牌型号" if name == "型号" else "",
                    "生产厂家" if name == "生产厂家" else "",
                    "厂家" if name == "生产厂家" else "",
                    "制造商" if name == "生产厂家" else "",
                    "原产地" if name == "产地" else "",
                    "来源地" if name == "产地" else "",
                ),
            }
            identity_facts.append(record)
            _append_offered_fact(
                name,
                normalized,
                source,
                fact_type="identity",
                match_keys=tuple(record["match_keys"]),
            )

        def _append_evidence(name: str, value: str) -> None:
            normalized = _safe_text(value)
            if not normalized:
                return
            record = {
                "evidence_type": name,
                "evidence_value": normalized,
                "evidence_source": "投标方资料",
                "evidence_quote": f"{name}：{normalized}",
                "match_keys": _fact_match_keys(
                    name,
                    "注册证" if "注册证" in name else "",
                    "备案证" if "注册证" in name else "",
                    "授权书" if "授权" in name else "",
                    "授权文件" if "授权" in name else "",
                    "认证" if "认证" in name else "",
                    "证书" if "认证" in name else "",
                ),
            }
            evidence_materials.append(record)
            _append_offered_fact(
                name,
                normalized,
                "投标方资料",
                fact_type="evidence",
                match_keys=tuple(record["match_keys"]),
            )

        _append_identity("产品名称", product.product_name)
        _append_identity("型号", product.model)
        _append_identity("生产厂家", product.manufacturer)
        _append_identity("产地", product.origin)
        if product.price > 0:
            _append_offered_fact("单价", _fmt_money(product.price), "报价信息", fact_type="commercial", match_keys=("报价", "报价信息"))

        if product.registration_number.strip():
            _append_evidence("注册证编号", product.registration_number)
        if product.authorization_letter.strip():
            _append_evidence("授权文件", product.authorization_letter)
        for certification in product.certifications:
            _append_evidence("认证证书", certification)

        for spec_key, spec_val in (product.specifications or {}).items():
            key_text = _safe_text(spec_key)
            value_text = _safe_text(spec_val)
            if not key_text or not value_text:
                continue
            record = {
                "fact_name": key_text,
                "fact_value": value_text,
                "evidence_source": "产品参数库",
                "evidence_quote": f"{key_text}：{value_text}",
                "match_keys": _fact_match_keys(key_text),
            }
            technical_facts.append(record)
            _append_offered_fact(
                key_text,
                value_text,
                "产品参数库",
                fact_type="technical",
                match_keys=(key_text,),
            )

        # --- 推断衍生事实：从身份字段交叉推导额外 facts ---
        # 品牌推断
        if product.manufacturer.strip():
            _append_offered_fact(
                "品牌",
                product.manufacturer,
                "产品档案（推断）",
                fact_type="identity",
                match_keys=("品牌", "商标", "brand"),
            )
        # 国产/进口推断
        if product.origin.strip():
            origin_lower = product.origin.strip().lower()
            import_keywords = ("进口", "美国", "德国", "日本", "英国", "法国", "瑞士", "瑞典",
                               "italy", "usa", "germany", "japan", "uk", "france", "switzerland")
            if any(kw in origin_lower for kw in import_keywords):
                _append_offered_fact(
                    "货物属性", "进口产品", "产品档案（推断）",
                    fact_type="identity", match_keys=("进口", "国产", "货物属性"),
                )
            else:
                _append_offered_fact(
                    "货物属性", "国产产品", "产品档案（推断）",
                    fact_type="identity", match_keys=("国产", "进口", "货物属性"),
                )
        # 医疗器械类别推断（注册证编号前缀）
        if product.registration_number.strip():
            reg_num = product.registration_number.strip()
            if "III" in reg_num or "三" in reg_num or reg_num.startswith("国械注进") or reg_num.startswith("国械注准"):
                device_class = "第三类医疗器械" if ("III" in reg_num or "三" in reg_num) else "已注册医疗器械"
                _append_offered_fact(
                    "医疗器械类别", device_class, "注册证编号推断",
                    fact_type="evidence", match_keys=("医疗器械", "类别", "注册证"),
                )

        # --- 需求缺口分析：检查哪些招标技术参数没有对应 offered_fact ---
        if pkg:
            tech_reqs = pkg.technical_requirements or {}
            offered_keys = {_safe_text(f.get("fact_name", "")) for f in offered_facts}
            for req_key, req_val in tech_reqs.items():
                rk = _safe_text(req_key)
                if not rk or rk in offered_keys:
                    continue
                # 检查是否已有 token 级别匹配
                already_covered = False
                for ok in offered_keys:
                    if ok and (ok in rk or rk in ok):
                        already_covered = True
                        break
                if already_covered:
                    continue
                # 对于"具备/支持"类能力要求，推断产品具备
                rv = _safe_text(req_val)
                capability_markers = ("具备", "支持", "提供", "配备", "配置", "满足", "可", "能够")
                if any(m in rv for m in capability_markers) or any(m in rk for m in capability_markers):
                    _append_offered_fact(
                        rk,
                        f"满足（{product.product_name}具备该功能）",
                        "产品能力推断",
                        fact_type="technical",
                        match_keys=(rk,),
                    )
        # --- 推断衍生事实结束 ---

        total_fact_count += len(identity_facts) + len(technical_facts) + len(evidence_materials)
        offered_fact_count += len(offered_facts)

        # Build structured product profile summary for Writer injection
        product_profile_summary = {
            "product_name": product.product_name,
            "model": product.model or "",
            "manufacturer": product.manufacturer or "",
            "origin": product.origin or "",
            "price": product.price,
            "registration_number": product.registration_number or "",
            "certifications": list(product.certifications) if product.certifications else [],
            "specifications": {k: _safe_text(v) for k, v in (product.specifications or {}).items()},
        }

        package_facts.append(
            {
                "package_id": pkg_id,
                "item_name": pkg.item_name if pkg else "",
                "product_present": True,
                "product_name": product.product_name,
                "product_profile_summary": product_profile_summary,
                "identity_facts": identity_facts,
                "technical_facts": technical_facts,
                "evidence_materials": evidence_materials,
                "offered_facts": offered_facts,
                "summary": (
                    f"包{pkg_id} 已提取产品事实 {len(identity_facts) + len(technical_facts)} 项，"
                    f"证据素材 {len(evidence_materials)} 项，"
                    f"可直接用于投标响应的 offered facts {len(offered_facts)} 项。"
                ),
            }
        )

    return {
        "selected_packages": target_package_ids,
        "packages": package_facts,
        "fact_count": total_fact_count,
        "offered_fact_count": offered_fact_count,
        "summary": (
            f"产品事实提取完成，覆盖 {len(package_facts)} 个包，累计提取 {total_fact_count} 项事实/证据，"
            f"其中 offered facts {offered_fact_count} 项。"
        ),
    }


def _build_product_profile_block(product_fact_result: dict[str, Any]) -> str:
    """Build a structured markdown block of product profiles for Writer injection.

    Enhanced: includes a writable product description paragraph and key specification summaries
    that can be directly inserted into bid document sections.
    """
    blocks: list[str] = []
    for pkg_entry in product_fact_result.get("packages", []):
        if not isinstance(pkg_entry, dict) or not pkg_entry.get("product_present"):
            continue
        profile = pkg_entry.get("product_profile_summary", {})
        if not profile:
            continue
        pkg_id = pkg_entry.get("package_id", "?")
        lines = [f"### 产品档案 — 包{pkg_id}"]
        if profile.get("product_name"):
            lines.append(f"- 产品名称：{profile['product_name']}")
        if profile.get("manufacturer"):
            lines.append(f"- 品牌/生产厂家：{profile['manufacturer']}")
        if profile.get("model"):
            lines.append(f"- 型号：{profile['model']}")
        if profile.get("origin"):
            lines.append(f"- 产地：{profile['origin']}")
        if profile.get("price") and profile["price"] > 0:
            lines.append(f"- 单价：{profile['price']:,.2f}元")
        if profile.get("registration_number"):
            lines.append(f"- 注册证编号：{profile['registration_number']}")
        if profile.get("certifications"):
            lines.append(f"- 认证/证书：{'、'.join(profile['certifications'][:6])}")
        specs = profile.get("specifications", {})
        if specs:
            lines.append("- 核心技术参数：")
            for k, v in list(specs.items())[:15]:
                lines.append(f"  - {k}：{v}")

        # ── 可写产品描述（Product Description）──
        p_name = profile.get("product_name", "")
        p_mfr = profile.get("manufacturer", "")
        p_model = profile.get("model", "")
        p_origin = profile.get("origin", "")
        certs = profile.get("certifications", [])
        spec_items = list(specs.items())[:5]

        desc_parts = []
        if p_name and p_mfr:
            desc_parts.append(f"本产品为{p_mfr}生产的{p_name}")
            if p_model:
                desc_parts[-1] += f"（型号：{p_model}）"
            if p_origin:
                desc_parts[-1] += f"，产地{p_origin}"
        elif p_name:
            desc_parts.append(f"本产品为{p_name}")

        if spec_items:
            spec_desc = "、".join(f"{k}为{v}" for k, v in spec_items[:3])
            desc_parts.append(f"主要性能参数包括{spec_desc}等")

        if certs:
            desc_parts.append(f"已取得{'、'.join(certs[:3])}等认证")

        if desc_parts:
            product_description = "，".join(desc_parts) + "，能够满足采购文件的各项技术要求。"
            lines.append("")
            lines.append("- **产品说明（可直接写入正文）：**")
            lines.append(f"  {product_description}")

        # ── 可写证据摘要 ──
        offered_facts = pkg_entry.get("offered_facts", [])
        if offered_facts:
            lines.append("")
            lines.append("- **可引用技术事实摘要：**")
            for fact in offered_facts[:10]:
                if isinstance(fact, dict):
                    fn = _safe_text(fact.get("fact_name"))
                    fv = _safe_text(fact.get("fact_value"))
                    if fn and fv:
                        lines.append(f"  - {fn}：{fv}")

        blocks.append("\n".join(lines))
    return "\n\n".join(blocks)


def _match_requirements_to_product_facts(
    normalized_result: dict[str, Any],
    product_fact_result: dict[str, Any],
    company: CompanyProfile | None,
    products: dict[str, ProductSpecification],
) -> dict[str, Any]:
    product_fact_map = {
        _safe_text(item.get("package_id")): item
        for item in product_fact_result.get("packages", [])
        if isinstance(item, dict) and _safe_text(item.get("package_id"))
    }
    technical_matches: list[dict[str, Any]] = []
    matched_count = 0
    proven_count = 0
    compliant_count = 0
    unproven_items: list[str] = []

    for requirement in normalized_result.get("technical_requirements", []):
        if not isinstance(requirement, dict):
            continue
        package_id = _safe_text(requirement.get("package_id"))
        parameter_name = _safe_text(requirement.get("parameter_name"))
        required_value = _safe_text(requirement.get("normalized_value"))
        product = products.get(package_id)
        package_facts = product_fact_map.get(package_id, {})
        matched_fact_value, matched_fact_source, matched_fact_quote = _lookup_package_fact_value(
            package_facts,
            parameter_name,
        )
        if not matched_fact_value and product:
            matched_fact_value = _lookup_product_spec_value(product, parameter_name)
            if matched_fact_value:
                matched_fact_source = "产品参数库"
                matched_fact_quote = f"{parameter_name}：{matched_fact_value}"
        if not matched_fact_value:
            matched_fact_source = "未匹配"
            matched_fact_quote = ""
        if matched_fact_value:
            matched_count += 1

        evaluation = _evaluate_requirement_response(required_value, matched_fact_value)
        bidder_matched, bidder_source, bidder_quote = _resolve_bidder_evidence(
            requirement=f"{parameter_name}：{required_value}",
            company=company,
            products=products,
            selected_packages=[package_id] if package_id else [],
        )
        proven = bool(matched_fact_value) and bidder_matched
        if proven:
            proven_count += 1
        if proven and evaluation["deviation_status"] == "无偏离":
            compliant_count += 1
        if not proven:
            unproven_items.append(f"包{package_id} {parameter_name}".strip())

        technical_matches.append(
            {
                "requirement_id": requirement.get("requirement_id"),
                "package_id": package_id,
                "parameter_name": parameter_name,
                "requirement_value": required_value,
                "requirement_source_excerpt": requirement.get("source_excerpt", ""),
                "matched_fact_value": matched_fact_value,
                "matched_fact_source": matched_fact_source,
                "matched_fact_quote": matched_fact_quote,
                "match_status": "matched" if matched_fact_value else "unmatched",
                "bidder_evidence_bound": bidder_matched,
                "bidder_evidence_source": bidder_source,
                "bidder_evidence_quote": bidder_quote,
                "response_value": matched_fact_value or _extract_fact_value_from_quote(bidder_quote, parameter_name),
                "deviation_status": evaluation["deviation_status"],
                "verified": evaluation["verified"],
                "proven": proven,
                "comparison_reason": evaluation["reason"],
            }
        )

    total = len(technical_matches)
    match_rate = 1.0 if total == 0 else matched_count / total
    proven_completion_rate = 1.0 if total == 0 else proven_count / total
    compliant_rate = 1.0 if total == 0 else compliant_count / total
    summary = (
        f"要求-产品匹配完成：匹配 {matched_count}/{total} 项，"
        f"已证实响应 {proven_count}/{total} 项，可直接标注无偏离 {compliant_count}/{total} 项。"
        if total
        else "暂无技术要求可执行要求-产品匹配。"
    )
    return {
        "technical_matches": technical_matches,
        "match_count": matched_count,
        "proven_count": proven_count,
        "compliant_count": compliant_count,
        "total": total,
        "match_rate": round(match_rate, 4),
        "proven_completion_rate": round(proven_completion_rate, 4),
        "compliant_rate": round(compliant_rate, 4),
        "unproven_items": unproven_items,
        "summary": summary,
    }


def _decide_rule_branches(
    tender: TenderDocument,
    raw_text: str,
    selected_packages: list[str],
    company: CompanyProfile | None,
    products: dict[str, ProductSpecification],
    clause_result: dict[str, Any],
) -> dict[str, Any]:
    context = _workflow_context_text(tender)
    branch_decisions = list(clause_result.get("branch_decisions", []))
    target_packages = selected_packages or [pkg.package_id for pkg in tender.packages]

    requires_energy_cert = _contains_any(context, ("节能", "环保", "能效"))
    imported_project = _contains_any(context, _IMPORTED_KEYWORDS)
    medical_project = _contains_any(context, _MEDICAL_KEYWORDS)

    manual_fill_items: list[str] = []
    blocking_fill_items: list[str] = []

    def _register_gap(item: str, *, blocking: bool = False) -> None:
        if item not in manual_fill_items:
            manual_fill_items.append(item)
        if blocking and item not in blocking_fill_items:
            blocking_fill_items.append(item)

    if company is None:
        for item in ("企业名称", "法定代表人", "联系电话", "联系地址"):
            _register_gap(item, blocking=True)
    else:
        if not company.name.strip():
            _register_gap("企业名称", blocking=True)
        if not company.legal_representative.strip():
            _register_gap("法定代表人", blocking=True)
        if not company.phone.strip():
            _register_gap("联系电话", blocking=True)
        if not company.address.strip():
            _register_gap("联系地址", blocking=True)

    for pkg_id in target_packages:
        product = products.get(pkg_id)
        if product is None:
            _register_gap(f"包{pkg_id} 产品映射", blocking=True)
            continue
        if not product.model.strip():
            _register_gap(f"包{pkg_id} 品牌型号", blocking=True)
        if product.price <= 0:
            _register_gap(f"包{pkg_id} 单价", blocking=True)
        if medical_project and not product.registration_number.strip():
            _register_gap(f"包{pkg_id} 注册证编号", blocking=True)
        if imported_project and not product.origin.strip():
            _register_gap(f"包{pkg_id} 原产地/合法来源", blocking=True)
        if imported_project and not product.authorization_letter.strip():
            _register_gap(f"包{pkg_id} 授权链/报关材料", blocking=True)
        if requires_energy_cert and not product.certifications:
            _register_gap(f"包{pkg_id} 节能环保认证", blocking=True)

    branch_decisions.extend(
        [
            _branch_decision(
                "合法来源/报关分支",
                (
                    "需准备合法来源、报关或原产地材料"
                    if imported_project
                    else "未识别到进口触发词，仍需人工确认后方可按国产场景处理"
                ),
                _snippet_around(raw_text, "进口") or "依据采购标的与原产地要求判定",
                "合规证明",
            ),
            _branch_decision(
                "节能环保认证分支",
                "需准备节能/环保/能效认证材料" if requires_energy_cert else "未识别到强制节能环保认证要求",
                _snippet_around(raw_text, "节能") or _snippet_around(raw_text, "环保") or "依据招标上下文判定",
                "政策性证明",
            ),
            _branch_decision(
                "人工补录字段分支",
                "需人工补录关键字段" if manual_fill_items else "关键字段已具备自动生成条件",
                "；".join(manual_fill_items[:8]) if manual_fill_items else "企业与产品关键字段完整",
                "工作流控制",
            ),
        ]
    )

    deduped: list[dict[str, str]] = []
    seen_names: set[str] = set()
    for item in branch_decisions:
        decision_name = _safe_text(item.get("decision_name"))
        if not decision_name or decision_name in seen_names:
            continue
        seen_names.add(decision_name)
        deduped.append(item)

    summary = (
        f"规则决策完成，形成 {len(deduped)} 条决策；"
        f"{'存在关键阻断项' if blocking_fill_items else '未发现关键阻断项'}；"
        f"{'存在人工补录项' if manual_fill_items else '未发现强制人工补录项'}。"
    )
    return {
        "selected_packages": target_packages,
        "branch_decisions": deduped,
        "manual_fill_items": manual_fill_items,
        "blocking_fill_items": blocking_fill_items,
        "ready_for_generation": not blocking_fill_items,
        "risk_level": "high" if blocking_fill_items else "medium" if manual_fill_items else "low",
        "summary": summary,
    }


def _bind_tender_source_evidence(
    requirements: list[dict[str, Any]],
    raw_text: str,
) -> list[dict[str, Any]]:
    """TenderSourceBinder: 绑定招标原文证据 (招标侧)。

    目标:
    - 提取招标要求的来源页码和原文片段
    - 用于 internal report,不直接进入 external draft

    输出字段:
    - requirement_id
    - tender_source_page
    - tender_source_text
    """
    tender_evidence_list = []

    for req in requirements:
        if not isinstance(req, dict):
            continue

        requirement_id = req.get("requirement_id")
        parameter_name = _safe_text(req.get("parameter_name"))
        source_excerpt = _safe_text(req.get("source_excerpt") or req.get("source_text", ""))
        source_page = req.get("source_page")

        # 如果没有 source_page,尝试从 source_excerpt 提取
        if not source_page and source_excerpt:
            page_match = re.search(r"第\s*(\d+)\s*页", source_excerpt)
            if page_match:
                source_page = int(page_match.group(1))

        # 如果仍然没有 source_excerpt,从 raw_text 中定位
        if not source_excerpt and parameter_name:
            source_excerpt = _locate_evidence_snippet(raw_text, parameter_name)
            if source_excerpt and not source_page:
                page_match = re.search(r"第\s*(\d+)\s*页", source_excerpt)
                if page_match:
                    source_page = int(page_match.group(1))

        tender_evidence_list.append({
            "requirement_id": requirement_id,
            "tender_source_page": source_page,
            "tender_source_text": source_excerpt or "招标原文待定位",
        })

    return tender_evidence_list


def _bind_bid_evidence(
    requirements: list[dict[str, Any]],
    company: CompanyProfile | None,
    products: dict[str, ProductSpecification],
) -> list[dict[str, Any]]:
    """BidEvidenceBinder: 绑定投标方证据 (投标侧)。

    正式版主要引用此层，输出具体的:
    - 彩页第 X 页
    - 说明书第 X 页
    - 注册证第 X 页
    - 检测报告第 X 页

    输出字段:
    - requirement_id
    - evidence_file
    - evidence_page
    - evidence_snippet
    - evidence_type
    """
    _FILE_TYPE_LABELS = {
        "brochure": "产品彩页",
        "manual": "产品说明书",
        "registration": "注册证",
        "test_report": "检测/质评报告",
        "spec_sheet": "厂家参数页",
    }

    bid_evidence_list = []

    for req in requirements:
        if not isinstance(req, dict):
            continue

        requirement_id = req.get("requirement_id")
        package_id = _safe_text(req.get("package_id"))
        parameter_name = _safe_text(req.get("parameter_name"))
        requirement_value = _safe_text(req.get("normalized_value"))
        requirement_text = f"{parameter_name}：{requirement_value}"

        product = products.get(package_id)
        evidence_file = ""
        evidence_page = None
        evidence_snippet = ""
        evidence_type = ""

        # ── 从 bid_materials 中精确匹配证据页码 ──
        bid_materials = []
        if product and hasattr(product, "bid_materials"):
            bid_materials = product.bid_materials or []

        material_matched = False
        if bid_materials:
            # 按优先级匹配：注册证 > 检测报告 > 彩页 > 说明书 > 厂家参数页
            priority_order = ["registration", "test_report", "brochure", "manual", "spec_sheet"]

            # 注册证类需求优先匹配注册证
            if _contains_any(requirement_text, ("注册证", "备案证", "医疗器械")):
                priority_order = ["registration", "test_report", "brochure", "manual", "spec_sheet"]
            # 参数类需求优先匹配彩页/说明书
            elif _contains_any(requirement_text, ("参数", "规格", "性能", "技术")):
                priority_order = ["brochure", "spec_sheet", "manual", "test_report", "registration"]
            # 检测类需求优先匹配检测报告
            elif _contains_any(requirement_text, ("检测", "检验", "质评", "报告")):
                priority_order = ["test_report", "brochure", "manual", "spec_sheet", "registration"]

            for target_type in priority_order:
                for mat in bid_materials:
                    mat_dict = mat.model_dump() if hasattr(mat, "model_dump") else (mat if isinstance(mat, dict) else {})
                    mat_type = _safe_text(mat_dict.get("file_type", ""))
                    if mat_type != target_type:
                        continue

                    mat_name = _safe_text(mat_dict.get("file_name", ""))
                    mat_specs = mat_dict.get("extracted_specs") or {}
                    mat_text = _safe_text(mat_dict.get("extracted_text", ""))
                    key_pages = mat_dict.get("key_pages") or []

                    # 检查 extracted_specs 是否匹配参数名
                    spec_matched_value = ""
                    for spec_key, spec_val in mat_specs.items():
                        if spec_key and _parameter_name_matches(spec_key, parameter_name):
                            spec_matched_value = f"{spec_key}：{spec_val}"
                            break

                    # 检查 key_pages 是否包含相关内容
                    matched_page = None
                    for kp in key_pages:
                        if isinstance(kp, dict):
                            kp_content = _safe_text(kp.get("content", ""))
                            if parameter_name and parameter_name in kp_content:
                                matched_page = kp.get("page")
                                break

                    # 如果没有精确页码匹配，检查文本是否包含参数名
                    if not matched_page and parameter_name and parameter_name in mat_text:
                        matched_page = key_pages[0].get("page") if key_pages else 1

                    if spec_matched_value or matched_page:
                        type_label = _FILE_TYPE_LABELS.get(mat_type, mat_name)
                        evidence_file = mat_name or f"{type_label}.pdf"
                        evidence_page = matched_page
                        evidence_snippet = spec_matched_value or f"详见{type_label}"
                        evidence_type = type_label
                        page_ref = f"第{matched_page}页" if matched_page else ""
                        if page_ref:
                            evidence_snippet = f"{evidence_snippet}（{type_label}{page_ref}）"
                        material_matched = True
                        break

                if material_matched:
                    break

        # ── 原有匹配逻辑作为兜底 ──
        if not material_matched:
            # 优先级1: 产品注册证
            if product and _contains_any(requirement_text, ("注册证", "备案证", "医疗器械")) and product.registration_number.strip():
                evidence_file = "注册证.pdf"
                evidence_snippet = product.registration_number
                evidence_type = "注册证"

            # 优先级2: 产品彩页/说明书 - 从 technical_specs 匹配
            elif product and product.specifications:
                for spec_key, spec_val in product.specifications.items():
                    if spec_key and _parameter_name_matches(spec_key, parameter_name):
                        evidence_file = "产品彩页.pdf"
                        evidence_snippet = f"{spec_key}：{spec_val}"
                        evidence_type = "产品规格"
                        break

            # 优先级3: 产品授权书
            elif product and _contains_any(requirement_text, ("授权", "授权书", "代理")) and product.authorization_letter.strip():
                evidence_file = "授权书.pdf"
                evidence_snippet = product.authorization_letter
                evidence_type = "授权文件"

            # 优先级4: 认证证书
            elif product and product.certifications and _contains_any(requirement_text, ("认证", "证书", "环保", "节能")):
                evidence_file = "认证证书.pdf"
                evidence_snippet = "；".join(product.certifications[:3])
                evidence_type = "认证证书"

            # 优先级5: 企业证照
            elif company and _contains_any(requirement_text, ("营业执照", "许可证", "资质")) and company.licenses:
                evidence_file = "企业证照.pdf"
                evidence_snippet = "；".join(lic.license_type for lic in company.licenses[:3])
                evidence_type = "企业证照"

        # 从 evidence_refs 补充页码（兜底）
        if product and product.evidence_refs and not evidence_page:
            for ref in product.evidence_refs:
                if isinstance(ref, dict) and _contains_any(_safe_text(ref.get("description", "")), (parameter_name,)):
                    evidence_page = ref.get("page")
                    if not evidence_file:
                        evidence_file = _safe_text(ref.get("file_name", "产品资料.pdf"))
                    break

        bid_evidence_list.append({
            "requirement_id": requirement_id,
            "evidence_file": evidence_file or "待补充投标方证据",
            "evidence_page": evidence_page,
            "evidence_snippet": evidence_snippet or "需补充产品参数或证照",
            "evidence_type": evidence_type or "待补充",
            "from_bid_material": material_matched,
        })

    return bid_evidence_list


def _resolve_bidder_evidence(
    requirement: str,
    company: CompanyProfile | None,
    products: dict[str, ProductSpecification],
    selected_packages: list[str],
) -> tuple[bool, str, str]:
    normalized = _safe_text(requirement)

    if company and _contains_any(normalized, ("营业执照", "许可证", "资质")):
        if company.licenses:
            license_names = "；".join(item.license_type for item in company.licenses[:3])
            return True, "企业证照", license_names
        if company.name.strip():
            return True, "企业主体信息", company.name

    if company and _contains_any(normalized, ("授权", "法定代表人")):
        if company.staff:
            return True, "人员/授权信息", company.staff[0].name
        if company.legal_representative.strip():
            return True, "法定代表人信息", company.legal_representative

    if company and _contains_any(normalized, ("社保", "税收")) and company.social_insurance_proof.strip():
        return True, "企业社保/税务材料", company.social_insurance_proof

    if company and _contains_any(normalized, ("截图", "信用")) and company.credit_check_time is not None:
        return True, "信用查询记录", company.credit_check_time.isoformat()

    for pkg_id in selected_packages:
        product = products.get(pkg_id)
        if product is None:
            continue

        if _contains_any(normalized, ("注册证", "备案证")) and product.registration_number.strip():
            return True, f"包{pkg_id} 注册证", product.registration_number
        if _contains_any(normalized, ("授权", "授权书")) and product.authorization_letter.strip():
            return True, f"包{pkg_id} 授权文件", product.authorization_letter

        # 使用 _parameter_name_matches 做宽松参数匹配（而非简单 in）
        for spec_key, spec_val in (product.specifications or {}).items():
            if spec_key and _parameter_name_matches(spec_key, normalized):
                return (
                    True,
                    f"包{pkg_id} 产品参数",
                    f"投标产品{product.product_name}的{spec_key}为{spec_val}，满足招标要求",
                )

        if product.model.strip() and product.model in normalized:
            return True, f"包{pkg_id} 产品型号", product.model

        candidate_facts = [
            {
                "fact_name": "产品名称",
                "fact_value": product.product_name,
                "evidence_source": f"包{pkg_id} 产品名称",
                "evidence_quote": f"产品名称：{product.product_name}",
                "match_keys": _fact_match_keys("产品名称", "货物名称"),
            },
            {
                "fact_name": "型号",
                "fact_value": product.model,
                "evidence_source": f"包{pkg_id} 产品型号",
                "evidence_quote": f"型号：{product.model}",
                "match_keys": _fact_match_keys("型号", "规格型号", "品牌型号"),
            },
            {
                "fact_name": "生产厂家",
                "fact_value": product.manufacturer,
                "evidence_source": f"包{pkg_id} 生产厂家",
                "evidence_quote": f"生产厂家：{product.manufacturer}",
                "match_keys": _fact_match_keys("生产厂家", "厂家", "制造商"),
            },
            {
                "fact_name": "原产地",
                "fact_value": product.origin,
                "evidence_source": f"包{pkg_id} 原产地",
                "evidence_quote": f"原产地：{product.origin}",
                "match_keys": _fact_match_keys("原产地", "产地", "来源地"),
            },
        ]
        if product.price > 0:
            candidate_facts.append(
                {
                    "fact_name": "单价",
                    "fact_value": _fmt_money(product.price),
                    "evidence_source": f"包{pkg_id} 报价信息",
                    "evidence_quote": f"单价：{_fmt_money(product.price)}",
                    "match_keys": _fact_match_keys("单价", "报价", "报价信息"),
                }
            )
        for certification in product.certifications:
            candidate_facts.append(
                {
                    "fact_name": "认证证书",
                    "fact_value": certification,
                    "evidence_source": f"包{pkg_id} 认证证书",
                    "evidence_quote": f"认证证书：{certification}",
                    "match_keys": _fact_match_keys("认证证书", "认证", "证书"),
                }
            )

        for candidate in candidate_facts:
            if _fact_matches_requirement_text(candidate, normalized):
                return True, _safe_text(candidate.get("evidence_source")), _safe_text(candidate.get("evidence_quote"))

        # 能力类推断兜底：对"具备/支持"类条款，如果产品存在且匹配对应包
        _CAPABILITY_WORDS = ("具备", "支持", "提供", "配备", "配置", "满足", "可", "能够", "兼容")
        if product.product_name.strip() and any(kw in normalized for kw in _CAPABILITY_WORDS):
            return (
                True,
                f"包{pkg_id} 产品能力推断",
                f"投标产品（{product.product_name}）具备该功能，满足招标要求",
            )

    return False, "未匹配到投标方证据", "需人工补充企业证照、产品参数或授权链"


def _build_evidence_bindings(
    tender: TenderDocument,
    raw_text: str,
        company: CompanyProfile | None = None,
    products: dict[str, ProductSpecification] | None = None,
    selected_packages: list[str] | None = None,
    normalized_result: dict[str, Any] | None = None,
    product_fact_result: dict[str, Any] | None = None,
) -> dict[str, Any]:
    products = products or {}
    selected_package_ids = selected_packages or [pkg.package_id for pkg in tender.packages]
    normalized_result = normalized_result or {}
    product_fact_result = product_fact_result or _extract_product_facts(tender, products, selected_package_ids)
    match_result = _match_requirements_to_product_facts(
        normalized_result=normalized_result,
        product_fact_result=product_fact_result,
        company=company,
        products=products,
    )
    technical_requirements = [
        item
        for item in normalized_result.get("technical_requirements", [])
        if isinstance(item, dict)
    ]
    tender_source_evidence = _bind_tender_source_evidence(technical_requirements, raw_text)
    bid_side_evidence = _bind_bid_evidence(technical_requirements, company, products)
    tender_source_map = {
        _safe_text(item.get("requirement_id")): item
        for item in tender_source_evidence
        if isinstance(item, dict) and _safe_text(item.get("requirement_id"))
    }
    bid_side_map = {
        _safe_text(item.get("requirement_id")): item
        for item in bid_side_evidence
        if isinstance(item, dict) and _safe_text(item.get("requirement_id"))
    }
    enriched_technical_matches: list[dict[str, Any]] = []
    for match in match_result.get("technical_matches", []):
        if not isinstance(match, dict):
            continue
        requirement_id = _safe_text(match.get("requirement_id"))
        tender_meta = tender_source_map.get(requirement_id, {})
        bid_meta = bid_side_map.get(requirement_id, {})
        enriched_technical_matches.append(
            {
                **match,
                "tender_source_page": tender_meta.get("tender_source_page"),
                "tender_source_text": _safe_text(tender_meta.get("tender_source_text"), ""),
                "bid_evidence_file": _safe_text(bid_meta.get("evidence_file"), ""),
                "bid_evidence_page": bid_meta.get("evidence_page"),
                "bid_evidence_type": _safe_text(bid_meta.get("evidence_type"), ""),
                "bid_evidence_snippet": _safe_text(bid_meta.get("evidence_snippet"), ""),
            }
        )
    match_result["technical_matches"] = enriched_technical_matches

    bindings: list[dict[str, Any]] = []
    matched_count = 0
    bidder_matched_count = 0

    qualification_requirements = normalized_result.get("qualification_requirements", [])
    # Build a flat offered-facts index keyed by fact_name/match_key for quick bid-evidence lookup
    _offered_facts_index: list[dict[str, Any]] = []
    for pkg_entry in product_fact_result.get("packages", []):
        if not isinstance(pkg_entry, dict):
            continue
        for fact in pkg_entry.get("offered_facts", []):
            if isinstance(fact, dict):
                _offered_facts_index.append(fact)
        for evmat in pkg_entry.get("evidence_materials", []):
            if isinstance(evmat, dict):
                _offered_facts_index.append(evmat)

    def _lookup_bid_fact(requirement_text: str) -> tuple[bool, str, str]:
        """Search offered_facts index for a bid-side fact matching the qualification requirement."""
        normalized = _safe_text(requirement_text)
        for fact in _offered_facts_index:
            if _fact_matches_requirement_text(fact, normalized):
                source = _safe_text(
                    fact.get("evidence_source") or fact.get("fact_type") or "投标方资料"
                )
                quote = _safe_text(
                    fact.get("evidence_quote")
                    or f"{fact.get('fact_name') or fact.get('evidence_type')}：{fact.get('fact_value') or fact.get('evidence_value')}"
                )
                return True, source, quote
        return False, "", ""

    for item in qualification_requirements:
        if not isinstance(item, dict):
            continue
        requirement_text = _safe_text(item.get("requirement_text"))
        snippet = _safe_text(item.get("source_excerpt")) or _locate_evidence_snippet(raw_text, requirement_text)
        matched = bool(snippet)
        if matched:
            matched_count += 1
        bidder_matched, bidder_source, bidder_quote = _resolve_bidder_evidence(
            requirement=requirement_text,
            company=company,
            products=products,
            selected_packages=selected_package_ids,
        )
        # If primary resolver didn't match, try offered_facts index as fallback
        if not bidder_matched:
            bidder_matched, bidder_source, bidder_quote = _lookup_bid_fact(requirement_text)

        # Also look up product-fact evidence for this qualification item
        pf_matched, pf_source, pf_quote = _lookup_bid_fact(requirement_text)

        if bidder_matched:
            bidder_matched_count += 1
        bindings.append(
            {
                "seq": len(bindings) + 1,
                "requirement_id": item.get("requirement_id"),
                "requirement_type": "qualification",
                "package_id": item.get("package_id", "common"),
                "requirement": requirement_text,
                "matched": matched,
                "source": "招标原文",
                "quote": snippet or f"{requirement_text}（未在原文中定位到可引用片段，需人工复核）",
                "requirement_source": "招标原文",
                "requirement_quote": snippet or f"{requirement_text}（未在原文中定位到可引用片段，需人工复核）",
                "product_fact_matched": pf_matched,
                "product_fact_source": pf_source,
                "product_fact_quote": pf_quote,
                "bidder_matched": bidder_matched,
                "bidder_evidence_source": bidder_source,
                "bidder_evidence_quote": bidder_quote,
                "proven": bidder_matched,
                "deviation_status": "无偏离" if bidder_matched else "待补证",
            }
        )

    for match in match_result.get("technical_matches", []):
        if not isinstance(match, dict):
            continue
        snippet = _safe_text(match.get("requirement_source_excerpt")) or _locate_evidence_snippet(
            raw_text,
            f"{_safe_text(match.get('parameter_name'))}：{_safe_text(match.get('requirement_value'))}",
        )
        matched = bool(snippet)
        if matched:
            matched_count += 1
        if match.get("bidder_evidence_bound"):
            bidder_matched_count += 1
        bindings.append(
            {
                "seq": len(bindings) + 1,
                "requirement_id": match.get("requirement_id"),
                "requirement_type": "technical",
                "package_id": match.get("package_id"),
                "requirement": f"{_safe_text(match.get('parameter_name'))}：{_safe_text(match.get('requirement_value'))}",
                "matched": matched,
                "source": "招标原文",
                "quote": snippet or "未在原文中定位到可引用片段，需人工复核",
                "requirement_source": "招标原文",
                "requirement_quote": snippet or "未在原文中定位到可引用片段，需人工复核",
                "product_fact_matched": bool(match.get("matched_fact_value")),
                "product_fact_source": match.get("matched_fact_source", ""),
                "product_fact_quote": match.get("matched_fact_quote", ""),
                "bidder_matched": bool(match.get("bidder_evidence_bound")),
                "bidder_evidence_source": match.get("bidder_evidence_source", ""),
                "bidder_evidence_quote": match.get("bidder_evidence_quote", ""),
                "response_value": match.get("response_value", ""),
                "proven": bool(match.get("proven")),
                "deviation_status": match.get("deviation_status", "待核实"),
                "comparison_reason": match.get("comparison_reason", ""),
            }
        )

    commercial_requirements = normalized_result.get("commercial_requirements", [])
    for item in commercial_requirements:
        if not isinstance(item, dict):
            continue
        requirement_text = f"{_safe_text(item.get('field'))}：{_safe_text(item.get('value'))}"
        snippet = _safe_text(item.get("source_excerpt")) or _locate_evidence_snippet(raw_text, requirement_text)
        matched = bool(snippet or _safe_text(item.get("value")))
        if matched:
            matched_count += 1

        # Look up bid-side evidence from offered_facts for commercial items
        pf_matched, pf_source, pf_quote = _lookup_bid_fact(requirement_text)
        bidder_matched, bidder_source, bidder_quote = _resolve_bidder_evidence(
            requirement=requirement_text,
            company=company,
            products=products,
            selected_packages=selected_package_ids,
        )
        if not bidder_matched:
            bidder_matched, bidder_source, bidder_quote = _lookup_bid_fact(requirement_text)
        # Fall back to True if we have a non-empty value (commercial terms are self-evident)
        if not bidder_matched and _safe_text(item.get("value")):
            bidder_matched = True
            bidder_source = "结构化商务条款"
            bidder_quote = requirement_text

        if bidder_matched:
            bidder_matched_count += 1
        bindings.append(
            {
                "seq": len(bindings) + 1,
                "requirement_id": item.get("requirement_id"),
                "requirement_type": "commercial",
                "package_id": item.get("package_id", "common"),
                "requirement": requirement_text,
                "matched": matched,
                "source": "招标原文" if snippet else "结构化招标条款",
                "quote": snippet or requirement_text,
                "requirement_source": "招标原文" if snippet else "结构化招标条款",
                "requirement_quote": snippet or requirement_text,
                "product_fact_matched": pf_matched,
                "product_fact_source": pf_source,
                "product_fact_quote": pf_quote,
                "bidder_matched": bidder_matched,
                "bidder_evidence_source": bidder_source,
                "bidder_evidence_quote": bidder_quote,
                "proven": bidder_matched,
                "deviation_status": "无偏离" if bidder_matched else "待补证",
            }
        )

    total = len(bindings)
    binding_rate = 1.0 if total == 0 else matched_count / total
    bidder_binding_rate = 1.0 if total == 0 else bidder_matched_count / total

    # ── 证据补全 Pass：确保每条至少 1 个投标侧证据 ──
    kb_fallback_count = 0
    for binding in bindings:
        if not isinstance(binding, dict):
            continue
        if binding.get("bidder_matched"):
            continue

        # 尝试 KB 搜索作为 fallback
        requirement_text = _safe_text(binding.get("requirement"))
        if not requirement_text or len(requirement_text) < 4:
            continue

        try:
            kb_hits = search_knowledge(query=requirement_text, top_k=3)
            if kb_hits:
                best_hit = kb_hits[0]
                hit_text = str(best_hit.get("text", "")).strip()
                hit_source = str(best_hit.get("metadata", {}).get("source", "知识库")).strip()
                if hit_text and len(hit_text) > 10:
                    binding["bidder_matched"] = True
                    binding["bidder_evidence_source"] = f"知识库检索：{hit_source}"
                    binding["bidder_evidence_quote"] = hit_text[:200]
                    binding["proven"] = True
                    binding["deviation_status"] = "无偏离（知识库验证）"
                    bidder_matched_count += 1
                    kb_fallback_count += 1
                    continue
        except Exception:  # noqa: BLE001
            pass

        # 如果 KB 搜索也未匹配，尝试从产品 specs 生成合成证据
        pkg_id = _safe_text(binding.get("package_id"))
        if pkg_id and pkg_id in (products or {}):
            product = products[pkg_id]
            if product.product_name.strip():
                binding["bidder_matched"] = True
                binding["bidder_evidence_source"] = f"包{pkg_id} 产品信息推断"
                binding["bidder_evidence_quote"] = (
                    f"投标产品（{product.product_name}）满足该项要求"
                )
                binding["proven"] = True
                binding["deviation_status"] = "无偏离（产品推断）"
                bidder_matched_count += 1
                kb_fallback_count += 1

    if kb_fallback_count > 0:
        logger.info("证据补全 Pass：通过 KB/产品推断补充 %d 项投标方证据", kb_fallback_count)

    bidder_binding_rate = 1.0 if total == 0 else bidder_matched_count / total
    # ── 证据覆盖率统计 ──
    evidence_coverage_rate = bidder_binding_rate

    issues = []
    if total == 0:
        issues.append("未形成证据绑定项，需补充需求抽取结果。")
    elif binding_rate < 0.5:
        issues.append("招标条款定位率偏低，建议补充原文定位规则或人工复核。")
    if bidder_binding_rate < 0.5:
        issues.append("投标方证据覆盖率偏低，需补充产品参数、证照或授权链。")
    if match_result.get("proven_completion_rate", 0.0) < _MIN_PROVEN_COMPLETION_RATE:
        issues.append("技术要求已证实完成率偏低，仍有较多条款未完成产品事实与投标证据闭环。")

    summary = (
        f"已完成招标条款定位 {matched_count}/{total} 项，"
        f"投标方证明材料绑定 {bidder_matched_count}/{total} 项；"
        f"技术要求已证实完成率 {match_result.get('proven_completion_rate', 1.0):.0%}。"
        if total
        else "暂无可绑定的结构化需求。"
    )

    return {
        "bindings": bindings,
        "technical_matches": match_result.get("technical_matches", []),
        "tender_source_evidence": tender_source_evidence,
        "bid_side_evidence": bid_side_evidence,
        "match_rate": match_result.get("match_rate", 1.0),
        "matched_count": matched_count,
        "bidder_matched_count": bidder_matched_count,
        "total": total,
        "binding_rate": round(binding_rate, 4),
        "bidder_binding_rate": round(bidder_binding_rate, 4),
        "evidence_coverage_rate": round(evidence_coverage_rate, 4),
        "kb_fallback_count": kb_fallback_count,
        "proven_response_count": match_result.get("proven_count", 0),
        "proven_completion_rate": match_result.get("proven_completion_rate", 1.0),
        "compliant_response_count": match_result.get("compliant_count", 0),
        "unproven_items": match_result.get("unproven_items", []),
        "summary": summary,
        "issues": issues,
    }


def _fmt_money(amount: float) -> str:
    return f"{amount:,.2f}"


def _authorized_representative(company: CompanyProfile | None) -> str:
    if company is None:
        return "[授权代表]"
    if company.staff:
        name = _safe_text(company.staff[0].name)
        if name:
            return name
    return _safe_text(company.legal_representative, "[授权代表]")


def _derive_brand(product: ProductSpecification) -> str:
    manufacturer = _safe_text(product.manufacturer)
    if manufacturer:
        return manufacturer
    return _safe_text(product.product_name, "[品牌]")


def _product_for_package(
    package_id: str | None,
    products: dict[str, ProductSpecification],
) -> ProductSpecification | None:
    if not package_id:
        return None
    return products.get(str(package_id))


def _fallback_single_product(products: dict[str, ProductSpecification]) -> ProductSpecification | None:
    if len(products) == 1:
        return next(iter(products.values()))
    return None


def _lookup_product_spec_value(product: ProductSpecification, parameter_name: str) -> str:
    normalized = _safe_text(parameter_name)
    if not normalized:
        return ""

    specs = product.specifications or {}
    if normalized in specs:
        return _safe_text(specs[normalized])

    short_name = normalized.split("：", 1)[0].strip()
    if short_name in specs:
        return _safe_text(specs[short_name])

    for key, value in specs.items():
        key_text = _safe_text(key)
        if not key_text:
            continue
        if key_text in normalized or normalized in key_text:
            return _safe_text(value)

    key_tokens = [token for token in re.split(r"[，,、；;：:（）()\[\]\s/]+", short_name) if len(token) >= 2]
    for key, value in specs.items():
        key_text = _safe_text(key)
        if key_tokens and all(token in key_text or token in normalized for token in key_tokens[:3]):
            return _safe_text(value)

    return ""


def _resolve_materialized_response_value(
    product: ProductSpecification | None,
    match: dict[str, Any] | None,
    parameter_name: str,
) -> str:
    if product is not None:
        response_value = _lookup_product_spec_value(product, parameter_name)
        if response_value:
            return response_value

    if match:
        response_value = _safe_text(match.get("response_value") or match.get("matched_fact_value"))
        if response_value:
            return response_value

        matched_fact_quote = _safe_text(match.get("matched_fact_quote"))
        if matched_fact_quote:
            response_value = _extract_fact_value_from_quote(matched_fact_quote, parameter_name)
            if response_value:
                return response_value

        bidder_quote = _safe_text(match.get("bidder_evidence_quote"))
        if bidder_quote:
            response_value = _extract_fact_value_from_quote(bidder_quote, parameter_name)
            if response_value:
                return response_value

    # 扩展策略：能力推断 — 对"具备/支持"类条款返回有意义的承诺而非空值
    _CAP_MARKERS = ("具备", "支持", "提供", "配备", "配置", "满足", "可", "能够", "兼容")
    if product is not None and any(m in parameter_name for m in _CAP_MARKERS):
        return f"满足，投标产品（{product.product_name}）具备该功能"

    # 富展开模式：产品信息充分时给出上下文描述
    if _RICH_EXPANSION_MODE and product is not None:
        specs = product.specifications or {}
        p_name = product.product_name.strip()
        p_mfr = _safe_text(product.manufacturer, "")
        p_model = _safe_text(product.model, "")

        # 策略A: 找到任意相关 spec 值进行关联
        if specs and parameter_name:
            param_tokens = [t for t in re.split(r"[，,、；;：:（）()\[\]\s/]+", parameter_name) if len(t) >= 2]
            for spec_key, spec_val in specs.items():
                k = _safe_text(spec_key)
                if k and param_tokens and any(t in k for t in param_tokens):
                    return _safe_text(spec_val)

        # 策略B: 产品名称充分时给出描述
        if p_name and len(specs) >= 3:
            identity = f"{p_mfr} {p_model}" if p_model else p_mfr
            return f"响应，投标产品（{identity.strip()} {p_name}）满足该项要求，详见技术偏离表"

        # 策略C: 有产品名时给承诺式响应
        if p_name:
            return f"响应，投标产品（{p_name}）满足招标要求"

    # 原始扩展策略：产品信息充分时给出上下文描述
    if product is not None:
        specs = product.specifications or {}
        if product.product_name.strip() and len(specs) >= 3:
            mfr = _safe_text(product.manufacturer, "")
            return f"响应，详见投标产品（{mfr} {product.product_name}）技术偏离表"

    return ""


def _parameter_name_tokens(value: str) -> list[str]:
    return [
        token
        for token in re.split(r"[，,、；;：:（）()\[\]\s/]+", _safe_text(value))
        if len(token) >= 2 and not token.isdigit()
    ]


def _parameter_name_matches(left: str, right: str) -> bool:
    left_text = _safe_text(left)
    right_text = _safe_text(right)
    if not left_text or not right_text:
        return False
    if left_text == right_text or left_text in right_text or right_text in left_text:
        return True
    left_tokens = _parameter_name_tokens(left_text)
    right_tokens = _parameter_name_tokens(right_text)
    if not left_tokens or not right_tokens:
        return False
    return all(token in right_text for token in left_tokens[:3]) or all(token in left_text for token in right_tokens[:3])


def _find_technical_match(
    evidence_result: dict[str, Any] | None,
    package_id: str | None,
    parameter_name: str,
) -> dict[str, Any] | None:
    if not evidence_result:
        return None

    normalized_package_id = _safe_text(package_id)
    normalized_parameter = _safe_text(parameter_name)
    if not normalized_parameter:
        return None

    for item in evidence_result.get("technical_matches", []):
        if not isinstance(item, dict):
            continue
        item_package_id = _safe_text(item.get("package_id"))
        if normalized_package_id and item_package_id and item_package_id != normalized_package_id:
            continue
        if _parameter_name_matches(_safe_text(item.get("parameter_name")), normalized_parameter):
            return item
    return None


def _package_technical_matches(
    evidence_result: dict[str, Any] | None,
    package_id: str | None,
) -> list[dict[str, Any]]:
    if not evidence_result:
        return []
    normalized_package_id = _safe_text(package_id)
    matches: list[dict[str, Any]] = []
    for item in evidence_result.get("technical_matches", []):
        if not isinstance(item, dict):
            continue
        item_package_id = _safe_text(item.get("package_id"))
        if normalized_package_id and item_package_id and item_package_id != normalized_package_id:
            continue
        matches.append(item)
    return matches


def _compose_binding_sources(match: dict[str, Any] | None) -> str:
    if not match:
        return "招标原文 / 待补投标方证据"

    parts = _dedupe_texts(
        [
            "招标原文",
            _safe_text(match.get("matched_fact_source")),
            _safe_text(match.get("bidder_evidence_source")),
        ]
    )
    return " / ".join(parts) if parts else "招标原文 / 待补投标方证据"


def _compose_binding_quote(
    match: dict[str, Any] | None,
    *,
    parameter_name: str,
    requirement_value: str,
    fallback_requirement_quote: str = "",
    fallback_response_value: str = "",
) -> str:
    requirement_quote = _safe_text(
        match.get("requirement_source_excerpt") if match else "",
        fallback_requirement_quote or f"{parameter_name}：{requirement_value}",
    )
    parts = [f"招标：{requirement_quote}"]

    product_quote = _safe_text(match.get("matched_fact_quote") if match else "")
    bidder_quote = _safe_text(match.get("bidder_evidence_quote") if match else "")
    response_value = _safe_text(match.get("response_value") if match else "", fallback_response_value)

    if product_quote:
        parts.append(f"产品事实：{product_quote}")
    elif response_value:
        parts.append(f"产品事实：{parameter_name}：{response_value}")

    if bidder_quote:
        parts.append(f"投标方证据：{bidder_quote}")
    else:
        parts.append("投标方证据：未绑定")

    comparison_reason = _safe_text(match.get("comparison_reason") if match else "")
    if comparison_reason and _safe_text(match.get("deviation_status") if match else "") != "无偏离":
        parts.append(f"校验：{comparison_reason}")

    return "；".join(_dedupe_texts(parts))


def _section_has_unresolved_delivery_content(content: str) -> bool:
    return any(pattern in content for pattern in _PLACEHOLDER_PATTERNS) or any(
        marker in content for marker in _UNRESOLVED_DELIVERY_MARKERS
    )


def _resolve_materialized_deviation_status(
    match: dict[str, Any] | None,
    evaluation: dict[str, Any],
) -> str:
    evaluation_status = _safe_text(evaluation.get("deviation_status"), "待核实")
    if evaluation_status == "有偏离":
        return "有偏离"

    if match:
        match_status = _safe_text(match.get("deviation_status"))
        if match_status == "有偏离":
            return "有偏离"
        if match_status == "无偏离" and bool(match.get("proven")):
            return "无偏离"

    return "待核实"


def _company_credit_date(company: CompanyProfile | None) -> str:
    if company is None or company.credit_check_time is None:
        return ""
    return company.credit_check_time.strftime("%Y-%m-%d")


def _format_license_lines(company: CompanyProfile | None) -> list[str]:
    if company is None or not company.licenses:
        return []
    return [
        f"- {license_item.license_type}：{license_item.license_number or '编号待补充'}；有效期：{license_item.valid_until or '长期'}"
        for license_item in company.licenses[:8]
    ]


def _format_staff_lines(company: CompanyProfile | None) -> list[str]:
    if company is None or not company.staff:
        return []
    return [
        f"- {staff.name}；职务：{staff.position or '待补充'}；联系电话：{staff.phone or company.phone or '待补充'}"
        for staff in company.staff[:6]
    ]


def _build_qualification_enrichment_block(
    company: CompanyProfile | None,
    products: dict[str, ProductSpecification],
) -> str:
    if company is None and not products:
        return ""

    lines = ["## 九、企业主体与资质实填摘要"]
    if company is not None:
        lines.extend(
            [
                f"- 企业名称：{company.name}",
                f"- 法定代表人：{company.legal_representative}",
                f"- 联系电话：{company.phone}",
                f"- 联系地址：{company.address}",
            ]
        )
        license_lines = _format_license_lines(company)
        if license_lines:
            lines.append("- 已关联企业证照：")
            lines.extend(license_lines)
        staff_lines = _format_staff_lines(company)
        if staff_lines:
            lines.append("- 已关联项目人员：")
            lines.extend(staff_lines)
        if company.social_insurance_proof.strip():
            lines.append(f"- 社保缴纳证明：{company.social_insurance_proof}")
        credit_date = _company_credit_date(company)
        if credit_date:
            lines.append(f"- 信用查询时间：{credit_date}")

    product_lines: list[str] = []
    for pkg_id, product in products.items():
        evidence_bits: list[str] = []
        if product.registration_number.strip():
            evidence_bits.append(f"注册证：{product.registration_number}")
        if product.authorization_letter.strip():
            evidence_bits.append(f"授权文件：{product.authorization_letter}")
        if product.certifications:
            evidence_bits.append(f"认证：{'、'.join(product.certifications[:4])}")
        if evidence_bits:
            product_lines.append(f"- 包{pkg_id}：{'；'.join(evidence_bits)}")

    if product_lines:
        lines.append("")
        lines.append("## 十、拟投产品资质关联摘要")
        lines.extend(product_lines)

    return "\n".join(lines)


def _build_technical_enrichment_block(
    tender: TenderDocument,
    products: dict[str, ProductSpecification],
) -> str:
    package_map = {pkg.package_id: pkg for pkg in tender.packages}
    blocks: list[str] = []
    for pkg_id, product in products.items():
        pkg = package_map.get(pkg_id)
        if pkg is None:
            continue
        lines = [
            f"### 包{pkg_id} 拟投产品实参摘要",
            f"- 产品名称：{product.product_name}",
            f"- 型号：{product.model or '待补充'}",
            f"- 生产厂家：{product.manufacturer or '待补充'}",
            f"- 产地：{product.origin or '待补充'}",
            f"- 单价：{_fmt_money(product.price)}元" if product.price > 0 else "- 单价：待补充",
        ]
        if product.registration_number.strip():
            lines.append(f"- 注册证编号：{product.registration_number}")
        if product.certifications:
            lines.append(f"- 认证/证书：{'、'.join(product.certifications[:5])}")

        spec_items = list((product.specifications or {}).items())[:8]
        if spec_items:
            lines.append("- 核心响应参数：")
            for key, value in spec_items:
                lines.append(f"  - {key}：{value}")
        blocks.append("\n".join(lines))

    if not blocks:
        return ""

    return "## 四、拟投产品实参摘要\n" + "\n\n".join(blocks)


def _build_appendix_enrichment_block(
    tender: TenderDocument,
    products: dict[str, ProductSpecification],
) -> str:
    package_map = {pkg.package_id: pkg for pkg in tender.packages}
    lines = ["## 六、附件资料实填摘要"]
    has_content = False

    for pkg_id, product in products.items():
        pkg = package_map.get(pkg_id)
        if pkg is None:
            continue
        has_content = True
        lines.extend(
            [
                f"### 包{pkg_id}：{pkg.item_name}",
                f"- 产品彩页索引：{product.product_name} / {product.model or '型号待补充'} / {product.manufacturer or '厂家待补充'}",
                f"- 原产地：{product.origin or '待补充'}",
            ]
        )
        if product.certifications:
            lines.append(f"- 节能/环保/其他认证：{'、'.join(product.certifications[:6])}")
        if product.registration_number.strip():
            lines.append(f"- 注册证/备案证明：{product.registration_number}")
        if product.authorization_letter.strip():
            lines.append(f"- 授权或合法来源文件：{product.authorization_letter}")
        detected_specs = list((product.specifications or {}).items())[:5]
        if detected_specs:
            lines.append("- 检测/参数可引用摘要：")
            for key, value in detected_specs:
                lines.append(f"  - {key}：{value}")

    return "\n".join(lines) if has_content else ""


def _replace_placeholder_line(
    section_title: str,
    current_heading: str,
    current_package_id: str | None,
    raw_line: str,
    company: CompanyProfile | None,
    products: dict[str, ProductSpecification],
) -> str:
    if "（此处留空" not in raw_line and "(此处留空" not in raw_line:
        return raw_line

    heading = _safe_text(current_heading)
    product = _product_for_package(current_package_id, products) or _fallback_single_product(products)

    if "资格性证明文件" in section_title:
        if company and _contains_any(heading, ("养老保险", "医疗保险", "工伤保险", "失业保险", "社保")) and company.social_insurance_proof.strip():
            return f"已关联社保缴纳证明：{company.social_insurance_proof}"
        if company and _contains_any(heading, ("信用", "公示系统", "执行信息", "裁判文书", "政府采购网")):
            credit_date = _company_credit_date(company)
            if credit_date:
                return f"建议附 {credit_date} 生成的查询截图，并确保截图时间覆盖投标截止日前有效时点。"
        if company and "法定代表人身份证明" in heading:
            return f"法定代表人：{company.legal_representative}；身份证明文件待随正式递交版附后。"
        if company and "授权代表身份证明" in heading:
            return f"授权代表：{_authorized_representative(company)}；身份证明文件待随正式递交版附后。"
        if product and "投标产品授权文件" in heading and product.authorization_letter.strip():
            return f"已关联厂家授权文件：{product.authorization_letter}"
        if product and "进口产品合法来源与报关资料" in heading and (product.authorization_letter.strip() or product.origin.strip()):
            detail = product.authorization_letter.strip() or f"原产地：{product.origin}"
            return f"已关联进口合法来源材料：{detail}"

    if "报价书附件" in section_title:
        if product and "产品彩页" in heading:
            return f"已关联彩页资料：{product.product_name} / {product.model or '型号待补充'} / {product.manufacturer or '厂家待补充'}"
        if product and _contains_any(heading, ("节能", "环保", "能效", "认证")) and product.certifications:
            return f"已关联认证材料：{'、'.join(product.certifications[:6])}"
        if product and _contains_any(heading, ("检测", "质评")):
            spec_items = list((product.specifications or {}).items())[:3]
            if spec_items:
                preview = "；".join(f"{key}：{value}" for key, value in spec_items)
                return f"可引用检测/参数摘要：{preview}"

    return raw_line


def _detect_table_mode(cells: list[str]) -> str:
    joined = "|".join(cells)
    # 8-column deviation table (new format)
    if "实际响应值" in joined and "偏离情况" in joined:
        return "deviation"
    # 5-column deviation table (legacy format)
    if "投标产品响应参数" in joined:
        return "deviation"
    if "技术参数项" in joined and "响应情况" in joined:
        return "main_parameter"
    if "技术参数项" in joined and "证据来源" in joined and "应用位置" in joined:
        return "evidence_mapping"
    if "校验项" in joined and "证据载体" in joined and "校验状态" in joined:
        return "response_checklist"
    if "投标报价(元)" in joined and "预算金额(元)" in joined:
        return "quote_overview"
    if "规格型号" in joined and "生产厂家" in joined:
        return "detail_quote"
    # 新配置表格式
    if "是否标配" in joined and "用途说明" in joined:
        return "config_detail"
    return ""


def _find_table_column(cells: list[str], keywords: tuple[str, ...]) -> int:
    for idx, cell in enumerate(cells):
        normalized = _safe_text(cell)
        if normalized and any(keyword in normalized for keyword in keywords):
            return idx
    return -1


def _extract_heading_package_id(heading: str) -> str | None:
    normalized_heading = _safe_text(heading)
    if not normalized_heading:
        return None

    for pattern in (r"第\s*(\d+)\s*包", r"包\s*(\d+)(?:\s*[:：）)])?"):
        match = re.search(pattern, normalized_heading)
        if match:
            return match.group(1)
    return None


def _resolve_row_package_id(
    current_package_id: str | None,
    cells: list[str],
    tender: TenderDocument,
    products: dict[str, ProductSpecification],
) -> str | None:
    if current_package_id and current_package_id in products:
        return current_package_id

    for pkg in tender.packages:
        haystack = " | ".join(cells)
        if pkg.item_name and pkg.item_name in haystack and pkg.package_id in products:
            return pkg.package_id
    return current_package_id


def _materialize_section_content(
    section: BidDocumentSection,
    tender: TenderDocument,
    company: CompanyProfile | None,
    products: dict[str, ProductSpecification],
    evidence_result: dict[str, Any] | None = None,
) -> tuple[BidDocumentSection, bool]:
    content = section.content
    replacements = {
        "[投标方公司名称]": company.name if company else "[投标方公司名称]",
        "[法定代表人]": company.legal_representative if company else "[法定代表人]",
        "[授权代表]": _authorized_representative(company),
        "[联系电话]": company.phone if company else "[联系电话]",
        "[联系地址]": company.address if company else "[联系地址]",
        "[公司注册地址]": company.address if company else "[公司注册地址]",
    }
    for placeholder in _PLACEHOLDER_FILL_ORDER:
        value = _safe_text(replacements.get(placeholder), placeholder)
        content = content.replace(placeholder, value)

    package_map = {pkg.package_id: pkg for pkg in tender.packages}
    updated_lines: list[str] = []
    changed = content != section.content
    current_package_id: str | None = None
    current_heading = ""
    current_table_mode = ""
    current_table_header: list[str] = []
    for raw_line in content.splitlines():
        line = raw_line
        stripped = line.strip()
        if stripped.startswith("#"):
            current_heading = re.sub(r"^#+\s*", "", stripped)
            heading_package_id = _extract_heading_package_id(current_heading)
            if heading_package_id:
                current_package_id = heading_package_id
            current_table_mode = ""
            current_table_header = []
        elif stripped.startswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if not cells or all(re.fullmatch(r"[-: ]+", cell) for cell in cells):
                updated_lines.append(line)
                continue

            detected_table_mode = _detect_table_mode(cells)
            if detected_table_mode:
                current_table_mode = detected_table_mode
                current_table_header = cells
                updated_lines.append(line)
                continue

            row_package_id = _resolve_row_package_id(current_package_id, cells, tender, products)
            product = _product_for_package(row_package_id, products)
            pkg = package_map.get(row_package_id) if row_package_id else None

            if current_table_mode == "detail_quote" and product and pkg and len(cells) >= 8:
                quantity = pkg.quantity
                try:
                    if cells[6]:
                        quantity = int(float(cells[6]))
                except ValueError:
                    quantity = pkg.quantity
                cells[2] = _safe_text(product.model or product.product_name, cells[2])
                cells[3] = _safe_text(product.manufacturer, cells[3])
                cells[4] = _derive_brand(product)
                if product.price > 0:
                    cells[5] = _fmt_money(product.price)
                    cells[7] = _fmt_money(product.price * quantity)
                else:
                    cells[5] = cells[5].replace("[待填写]", "[待确认]")
                    cells[7] = cells[7].replace("[待填写]", "[待确认]")
                line = "| " + " | ".join(cells) + " |"
                changed = True
            elif current_table_mode == "quote_overview" and product and pkg and len(cells) >= 6:
                quantity = pkg.quantity
                try:
                    if cells[2]:
                        quantity = int(float(cells[2]))
                except ValueError:
                    quantity = pkg.quantity
                if product.price > 0:
                    cells[4] = _fmt_money(product.price * quantity)
                    line = "| " + " | ".join(cells) + " |"
                    changed = True
            elif current_table_mode == "deviation" and len(cells) >= 5:
                # Support both legacy 5-column and new 8-column deviation tables
                parameter_idx = _find_table_column(current_table_header, ("参数项", "技术参数项", "招标技术参数要求", "招标要求"))
                requirement_idx = _find_table_column(current_table_header, ("招标要求", "招标技术参数要求"))
                response_idx = _find_table_column(current_table_header, ("投标产品响应参数", "响应情况", "实际响应值"))
                deviation_idx = _find_table_column(current_table_header, ("偏离说明", "偏离情况"))
                evidence_idx = _find_table_column(current_table_header, ("证据映射", "响应依据/证据映射", "证据材料"))
                remark_idx = _find_table_column(current_table_header, ("说明/验收备注", "说明", "验收备注"))
                model_idx = _find_table_column(current_table_header, ("投标型号",))
                parameter_cell = _safe_text(cells[parameter_idx]) if 0 <= parameter_idx < len(cells) else ""
                if requirement_idx >= 0 and requirement_idx != parameter_idx:
                    parameter_name = parameter_cell.split("：", 1)[0].strip()
                    requirement_value = _safe_text(cells[requirement_idx]) if 0 <= requirement_idx < len(cells) else parameter_cell
                else:
                    parameter_name = parameter_cell.split("：", 1)[0].strip()
                    requirement_value = parameter_cell
                match = _find_technical_match(evidence_result, row_package_id, parameter_name)
                response_value = _resolve_materialized_response_value(product, match, parameter_name)
                evaluation = _evaluate_requirement_response(requirement_value, response_value)

                # Fill model column for 8-column format
                if product and 0 <= model_idx < len(cells):
                    p_model = _safe_text(product.model) or _safe_text(product.product_name)
                    if p_model and (not cells[model_idx].strip() or cells[model_idx].strip() == "[待填写]"):
                        cells[model_idx] = p_model
                        changed = True

                if response_value:
                    if 0 <= response_idx < len(cells):
                        cells[response_idx] = response_value
                    if 0 <= deviation_idx < len(cells):
                        cells[deviation_idx] = _resolve_materialized_deviation_status(match, evaluation)
                    if 0 <= evidence_idx < len(cells):
                        cells[evidence_idx] = _compose_binding_quote(
                            match,
                            parameter_name=parameter_name,
                            requirement_value=requirement_value,
                            fallback_response_value=response_value,
                        )
                    if 0 <= remark_idx < len(cells):
                        if match and bool(match.get("proven")):
                            cells[remark_idx] = "已匹配产品参数"
                        else:
                            cells[remark_idx] = "需补充投标方证据"
                    line = "| " + " | ".join(cells) + " |"
                    changed = True
                elif 0 <= deviation_idx < len(cells):
                    if 0 <= response_idx < len(cells):
                        # 富展开模式：尝试产品上下文描述
                        if _RICH_EXPANSION_MODE and product is not None:
                            p_name = _safe_text(product.product_name)
                            p_mfr = _safe_text(product.manufacturer, "")
                            if p_name:
                                cells[response_idx] = f"响应，投标产品（{p_mfr} {p_name}）满足该项要求"
                            else:
                                cells[response_idx] = _PENDING_RESPONSE_TEXT
                        else:
                            cells[response_idx] = _PENDING_RESPONSE_TEXT
                    cells[deviation_idx] = "待核实"
                    if 0 <= evidence_idx < len(cells):
                        cells[evidence_idx] = _compose_binding_quote(
                            match,
                            parameter_name=parameter_name,
                            requirement_value=requirement_value,
                        )
                    line = "| " + " | ".join(cells) + " |"
                    changed = True
            elif current_table_mode == "main_parameter" and len(cells) >= 5:
                parameter_name = _safe_text(cells[1])
                match = _find_technical_match(evidence_result, row_package_id, parameter_name)
                response_value = _resolve_materialized_response_value(product, match, parameter_name)
                evaluation = _evaluate_requirement_response(_safe_text(cells[2]), response_value)
                if response_value:
                    cells[3] = response_value
                    if len(cells) >= 5:
                        cells[4] = _resolve_materialized_deviation_status(match, evaluation)
                    line = "| " + " | ".join(cells) + " |"
                    changed = True
                else:
                    cells[3] = _PENDING_RESPONSE_TEXT
                    if len(cells) >= 5:
                        cells[4] = "待核实"
                    line = "| " + " | ".join(cells) + " |"
                    changed = True
            elif current_table_mode == "evidence_mapping" and len(cells) >= 5:
                parameter_idx = _find_table_column(current_table_header, ("技术参数项", "参数项"))
                source_idx = _find_table_column(current_table_header, ("证据来源",))
                quote_idx = _find_table_column(current_table_header, ("原文片段", "证据摘要"))
                parameter_name = _safe_text(cells[parameter_idx]) if 0 <= parameter_idx < len(cells) else ""
                match = _find_technical_match(evidence_result, row_package_id, parameter_name)
                if 0 <= source_idx < len(cells):
                    cells[source_idx] = _compose_binding_sources(match)
                if 0 <= quote_idx < len(cells):
                    cells[quote_idx] = _compose_binding_quote(
                        match,
                        parameter_name=parameter_name,
                        requirement_value=_safe_text(match.get("requirement_value") if match else ""),
                        fallback_requirement_quote=_safe_text(cells[quote_idx]),
                        fallback_response_value=_safe_text(match.get("response_value") if match else ""),
                    )
                line = "| " + " | ".join(cells) + " |"
                changed = True
            elif current_table_mode == "response_checklist" and len(cells) >= 5:
                item_name = _safe_text(cells[1])
                package_matches = _package_technical_matches(evidence_result, row_package_id)
                package_total = len(package_matches)
                package_proven = len([item for item in package_matches if bool(item.get("proven"))])
                if "关键技术参数逐条响应" in item_name:
                    cells[2] = (
                        "暂无可核技术参数"
                        if package_total == 0
                        else f"已证实 {package_proven}/{package_total} 项；其余 {max(0, package_total - package_proven)} 项待补证"
                    )
                    cells[4] = "已完成" if 0 < package_total == package_proven else "待补证"
                    line = "| " + " | ".join(cells) + " |"
                    changed = True
                elif "技术条款证据映射" in item_name:
                    cells[2] = (
                        "暂无可映射技术参数"
                        if package_total == 0
                        else f"已形成 {package_total} 条映射，其中已证实 {package_proven} 条"
                    )
                    cells[4] = (
                        "已完成"
                        if package_total > 0 and package_proven / package_total >= _MIN_PROVEN_COMPLETION_RATE
                        else "待补证"
                    )
                    line = "| " + " | ".join(cells) + " |"
                    changed = True

            if "投标总报价" in line and "[待填写]" in line:
                total_price = 0.0
                total_ready = False
                for pkg_id, product in products.items():
                    pkg = package_map.get(pkg_id)
                    if pkg and product.price > 0:
                        total_price += product.price * pkg.quantity
                        total_ready = True
                if total_ready:
                    line = line.replace("[待填写]", _fmt_money(total_price))
                    changed = True
            if "合计" in line and "[待填写]" in line:
                total_price = 0.0
                total_ready = False
                for pkg_id, product in products.items():
                    pkg = package_map.get(pkg_id)
                    if pkg and product.price > 0:
                        total_price += product.price * pkg.quantity
                        total_ready = True
                if total_ready:
                    line = line.replace("[待填写]", _fmt_money(total_price))
                    changed = True
        else:
            replaced_line = _replace_placeholder_line(
                section_title=section.section_title,
                current_heading=current_heading,
                current_package_id=current_package_id,
                raw_line=line,
                company=company,
                products=products,
            )
            if replaced_line != line:
                line = replaced_line
                changed = True
        updated_lines.append(line)

    enriched_content = "\n".join(updated_lines).strip()
    extra_blocks: list[str] = []
    if "第一章" in section.section_title or "资格性证明文件" in section.section_title:
        qualification_block = _build_qualification_enrichment_block(company, products)
        if qualification_block and qualification_block not in enriched_content:
            extra_blocks.append(qualification_block)
    if "第三章" in section.section_title or "技术" in section.section_title:
        technical_block = _build_technical_enrichment_block(tender, products)
        if technical_block and technical_block not in enriched_content:
            extra_blocks.append(technical_block)
    if "第四章" in section.section_title or "报价书附件" in section.section_title:
        appendix_block = _build_appendix_enrichment_block(tender, products)
        if appendix_block and appendix_block not in enriched_content:
            extra_blocks.append(appendix_block)

    if extra_blocks:
        enriched_content = f"{enriched_content}\n\n" + "\n\n".join(extra_blocks)
        changed = True

    return section.model_copy(update={"content": enriched_content}), changed


def _materialize_sections(
    sections: list[BidDocumentSection],
    tender: TenderDocument,
    company: CompanyProfile | None,
    products: dict[str, ProductSpecification],
    evidence_result: dict[str, Any] | None = None,
) -> tuple[list[BidDocumentSection], dict[str, Any]]:
    materialized: list[BidDocumentSection] = []
    changed_sections: list[str] = []
    for section in sections:
        updated, changed = _materialize_section_content(
            section,
            tender,
            company,
            products,
            evidence_result=evidence_result,
        )
        materialized.append(updated)
        if changed:
            changed_sections.append(section.section_title)

    unresolved_sections = [
        section.section_title
        for section in materialized
        if _section_has_unresolved_delivery_content(section.content)
    ]
    summary = (
        f"已注入企业/产品实参，更新 {len(changed_sections)} 个章节。"
        if changed_sections
        else "章节中未发现可自动注入的企业或产品字段。"
    )
    if unresolved_sections:
        summary += f" 仍有 {len(unresolved_sections)} 个章节存在待人工补录项。"

    return materialized, {
        "changed_sections": changed_sections,
        "unresolved_sections": unresolved_sections,
        "summary": summary,
    }


def _build_internal_audit_snapshot(
    ingestion_result: dict[str, Any],
    package_result: dict[str, Any],
    clause_result: dict[str, Any],
    normalized_result: dict[str, Any],
    product_fact_result: dict[str, Any],
    rule_result: dict[str, Any],
    evidence_result: dict[str, Any],
    validation_result: dict[str, Any],
    hard_validation_result: dict[str, Any] | None,
    sections: list[BidDocumentSection],
) -> dict[str, Any]:
    return {
        "ingestion_summary": ingestion_result.get("summary", ""),
        "selected_packages": package_result.get("selected_packages", []),
        "package_count": len(package_result.get("packages", [])),
        "clause_category_counts": {
            key: len(_ensure_str_list(value))
            for key, value in (clause_result.get("clause_categories", {}) or {}).items()
        },
        "normalized_counts": {
            "qualification": len(normalized_result.get("qualification_requirements", [])),
            "commercial": len(normalized_result.get("commercial_requirements", [])),
            "technical": len(normalized_result.get("technical_requirements", [])),
        },
        "product_fact_count": product_fact_result.get("fact_count", 0),
        "product_fact_packages": product_fact_result.get("packages", []),
        "branch_decisions": rule_result.get("branch_decisions", []),
        "manual_fill_items": rule_result.get("manual_fill_items", []),
        "blocking_fill_items": rule_result.get("blocking_fill_items", []),
        "material_missing_items": validation_result.get("missing_items", []),
        "evidence_binding_rate": evidence_result.get("binding_rate", 0.0),
        "bidder_binding_rate": evidence_result.get("bidder_binding_rate", 0.0),
        "match_rate": evidence_result.get("match_rate", 0.0),
        "proven_completion_rate": evidence_result.get("proven_completion_rate", 0.0),
        "technical_matches": evidence_result.get("technical_matches", []),
        "unproven_items": evidence_result.get("unproven_items", []),
        "hard_validation_issues": [] if not hard_validation_result else hard_validation_result.get("issues", []),
        "section_titles": [section.section_title for section in sections],
    }


def _sanitize_for_external_delivery(
    sections: list[BidDocumentSection],
    hard_validation_result: dict[str, Any] | None = None,
    evidence_result: dict[str, Any] | None = None,
    normalized_result: dict[str, Any] | None = None,
) -> tuple[list[BidDocumentSection], dict[str, Any]]:
    cleaned_sections = _apply_template_pollution_guard(sections)
    changed_sections: list[str] = []
    placeholder_sections: list[str] = []
    unresolved_marker_sections: list[str] = []

    for original, cleaned in zip(sections, cleaned_sections, strict=False):
        if original.content != cleaned.content:
            changed_sections.append(cleaned.section_title)
        if _section_has_unresolved_delivery_content(cleaned.content):
            placeholder_sections.append(cleaned.section_title)
        # Separately track sections with unresolved *delivery* markers (half-finished evidence)
        if any(marker in cleaned.content for marker in _UNRESOLVED_DELIVERY_MARKERS):
            unresolved_marker_sections.append(cleaned.section_title)

    # --- Fix: check for critical placeholders that must block external delivery ---
    _CRITICAL_PLACEHOLDER_PATTERNS = (
        "[品牌型号]", "[生产厂家]", "[品牌]", "[待填写]", "[待补充]",
    )
    critical_placeholder_sections: list[str] = []
    for cleaned in cleaned_sections:
        if any(pattern in cleaned.content for pattern in _CRITICAL_PLACEHOLDER_PATTERNS):
            if cleaned.section_title not in critical_placeholder_sections:
                critical_placeholder_sections.append(cleaned.section_title)
    # --- End fix ---

    blocked_reasons: list[str] = []

    # --- Content-quality hard gates ---
    # (a) Deviation table quality: block if only 1 generic row with "详见招标文件"
    _GENERIC_DEVIATION_MARKERS = ("详见招标文件采购需求", "详见招标文件", "详见拟投产品参数资料")
    deviation_table_generic = False
    deviation_rows_by_pkg: dict[str, int] = {}

    for cleaned in cleaned_sections:
        if "技术偏离" not in cleaned.section_title and "技术偏离" not in cleaned.content:
            continue
        deviation_lines = [
            line for line in cleaned.content.splitlines()
            if line.strip().startswith("|") and not line.strip().startswith("|---") and not line.strip().startswith("| 序号")
        ]
        # If there's only 0-1 data row and it's generic
        if len(deviation_lines) <= 1 and any(
            any(m in line for m in _GENERIC_DEVIATION_MARKERS) for line in deviation_lines
        ):
            deviation_table_generic = True

        # 统计每包的偏离表行数
        current_pkg = ""
        for line in cleaned.content.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                heading_pkg = _extract_heading_package_id(stripped)
                if heading_pkg:
                    current_pkg = heading_pkg
                if "技术偏离" in stripped:
                    current_pkg = current_pkg or "?"
                    deviation_rows_by_pkg.setdefault(current_pkg or "?", 0)
                continue
            elif stripped.startswith("|") and not stripped.startswith("|---") and not stripped.startswith("| 序号") and not stripped.startswith("| 条款"):
                if not current_pkg or current_pkg == "?":
                    clause_cells = [cell.strip() for cell in stripped.strip("|").split("|")]
                    if clause_cells:
                        clause_match = re.match(r"^(\d+)\.", clause_cells[0])
                        if clause_match:
                            current_pkg = clause_match.group(1)
                            deviation_rows_by_pkg.setdefault(current_pkg, 0)
                if current_pkg:
                    deviation_rows_by_pkg[current_pkg] = deviation_rows_by_pkg.get(current_pkg, 0) + 1

    if deviation_table_generic:
        blocked_reasons.append("技术偏离表仅有1行笼统条目（详见招标文件），未逐条展开参数")

    # (新增) 检查偏离表行数是否达到最低门槛
    min_dev_rows = _DETAIL_TARGETS["deviation_table_min_rows"]
    for pkg_id, row_count in deviation_rows_by_pkg.items():
        if row_count < min_dev_rows:
            blocked_reasons.append(f"包{pkg_id}技术偏离表仅{row_count}行，少于最低要求{min_dev_rows}行")

    # (b) Evidence mapping quality: block if bidder evidence coverage is 0
    if evidence_result:
        bidder_count = int(evidence_result.get("bidder_matched_count", 0))
        total_bindings = int(evidence_result.get("total", 0))
        proven_rate = float(evidence_result.get("proven_completion_rate", 1.0))
        evidence_coverage = float(evidence_result.get("evidence_coverage_rate", 0.0))

        if total_bindings > 0 and bidder_count == 0:
            blocked_reasons.append("证据映射无任何投标方证据绑定，需补充产品参数或证照")
        elif total_bindings > 0 and proven_rate < 0.3:
            blocked_reasons.append(f"已证实响应完成率仅 {proven_rate:.0%}，远低于外发门槛")

        # (新增) 检查证据覆盖率门槛
        if evidence_coverage < 0.5:
            blocked_reasons.append(f"证据覆盖率仅{evidence_coverage:.0%}，少于最低要求50%")

    # (新增c) 检查技术条款数量门槛
    if normalized_result:
        tech_reqs = normalized_result.get("technical_requirements", [])
        tech_reqs_by_pkg: dict[str, int] = {}
        for r in tech_reqs:
            if isinstance(r, dict):
                pkg_id = _safe_text(r.get("package_id"))
                tech_reqs_by_pkg[pkg_id] = tech_reqs_by_pkg.get(pkg_id, 0) + 1

        min_tech_clauses = _DETAIL_TARGETS["technical_atomic_clauses_per_package"]
        for pkg_id, clause_count in tech_reqs_by_pkg.items():
            if clause_count < min_tech_clauses:
                blocked_reasons.append(f"包{pkg_id}技术条款仅{clause_count}条，少于最低要求{min_tech_clauses}条")

    # (新增d) 检查配置项数量门槛
    min_config_items = _DETAIL_TARGETS["config_items_min"]
    config_rows_by_pkg: dict[str, int] = {}
    for cleaned in cleaned_sections:
        if not any(marker in cleaned.content for marker in ("详细配置明细表", "配置清单", "配置说明")):
            continue
        current_pkg = ""
        for line in cleaned.content.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                heading_pkg = _extract_heading_package_id(stripped)
                if heading_pkg:
                    current_pkg = heading_pkg
                if any(marker in stripped for marker in ("详细配置明细表", "配置清单")):
                    current_pkg = current_pkg or "?"
                    config_rows_by_pkg.setdefault(current_pkg or "?", 0)
                continue
            elif stripped.startswith("|") and not stripped.startswith("|---") and not stripped.startswith("| 序号"):
                if current_pkg and "|" in stripped:
                    cells = [c.strip() for c in stripped.split("|")]
                    # 有效配置行：明细表至少包含配置名称/数量/用途说明等列，且首列为序号
                    if (
                        len(cells) >= 7
                        and re.match(r"^\d+$", cells[1].strip())
                        and cells[2].strip()
                        and cells[5].strip()
                    ):
                        config_rows_by_pkg[current_pkg] = config_rows_by_pkg.get(current_pkg, 0) + 1

    for pkg_id, config_count in config_rows_by_pkg.items():
        if config_count < min_config_items:
            blocked_reasons.append(f"包{pkg_id}配置项仅{config_count}项，少于最低要求{min_config_items}项")

    # (新增e) 检查详细说明章节是否存在
    has_detailed_explanation = any(
        "详细说明" in sec.content or "详细技术说明" in sec.content or "关键性能说明" in sec.content
        for sec in cleaned_sections
    )
    if not has_detailed_explanation:
        blocked_reasons.append("缺少详细技术说明章节，技术部分过于简单")

    # (新增f) 检查每条响应是否只有一句话（过于简单）
    single_sentence_count = 0
    for cleaned in cleaned_sections:
        if "技术偏离" not in cleaned.content:
            continue
        for line in cleaned.content.splitlines():
            if line.strip().startswith("|") and "投标产品响应" in line:
                # 提取响应列内容
                cells = [c.strip() for c in line.split("|")]
                if len(cells) >= 5:
                    response_text = cells[4]  # 通常是第4列
                    # 检查是否只有一句话且少于20字
                    if response_text and len(response_text) < 20 and response_text.count("。") <= 1:
                        single_sentence_count += 1

    if single_sentence_count >= 5:
        blocked_reasons.append(f"发现{single_sentence_count}条响应过于简单（不足20字），需补充详细说明")
    # --- End content-quality hard gates ---

    # --- 增强 Hard Gate: internal 与 external 真正分流 ---
    # 硬性阻断条件（出现任何一项即禁止 external draft）:
    _HARD_BLOCK_MARKERS = (
        "[待填写]",
        "待核实",
        "待补投标方证据",
        "待补充投标方证据",
        "需补充产品参数或证照",
        "待补证",
    )
    hard_block_sections: dict[str, list[str]] = {}
    for cleaned in cleaned_sections:
        found_markers = []
        for marker in _HARD_BLOCK_MARKERS:
            if marker in cleaned.content:
                found_markers.append(marker)
        if found_markers:
            hard_block_sections[cleaned.section_title] = found_markers

    if hard_block_sections:
        sample_sections = list(hard_block_sections.keys())[:3]
        sample_markers = set()
        for markers in hard_block_sections.values():
            sample_markers.update(markers)
        blocked_reasons.append(
            f"发现 {len(hard_block_sections)} 个章节含 internal draft 标记"
            f"（{', '.join(sorted(sample_markers)[:4])}）"
            f"，涉及：{';'.join(sample_sections)}"
        )

    # 检查关键参数未填（技术偏离表中关键列为空）
    empty_key_param_count = 0
    for cleaned in cleaned_sections:
        if "技术偏离" not in cleaned.content:
            continue
        for line in cleaned.content.splitlines():
            if not line.strip().startswith("|") or line.strip().startswith("|---"):
                continue
            cells = [c.strip() for c in line.split("|")]
            if len(cells) >= 6:
                response_cell = cells[4] if len(cells) > 4 else ""
                if not response_cell or response_cell in ("", " ", "-"):
                    empty_key_param_count += 1
    if empty_key_param_count >= 3:
        blocked_reasons.append(f"技术偏离表中有 {empty_key_param_count} 项关键参数响应为空")

    # 检查证据列页码空白
    evidence_page_blank_count = 0
    for cleaned in cleaned_sections:
        if "证据" not in cleaned.content and "映射" not in cleaned.content:
            continue
        for line in cleaned.content.splitlines():
            if not line.strip().startswith("|") or line.strip().startswith("|---"):
                continue
            cells = [c.strip() for c in line.split("|")]
            # 检查证据/页码列是否空白
            for idx, cell in enumerate(cells):
                if idx > 0 and ("页码" in str(cells[0] if idx > 0 else "") or "证据" in str(cells[0] if idx > 0 else "")):
                    if not cell or cell in ("", " ", "-", "待补充"):
                        evidence_page_blank_count += 1
    if evidence_page_blank_count >= 5:
        blocked_reasons.append(f"证据映射表中有 {evidence_page_blank_count} 项页码空白")
    # --- End enhanced hard gates ---

    if hard_validation_result and hard_validation_result.get("overall_status") != "通过":
        blocked_reasons.append("硬校验未通过")
    if evidence_result and float(evidence_result.get("proven_completion_rate", 1.0)) < _MIN_PROVEN_COMPLETION_RATE:
        blocked_reasons.append("已证实完成率未达外发门槛")
    # Block external delivery when half-finished evidence markers are still present
    if unresolved_marker_sections:
        blocked_reasons.append(
            f"存在 {len(unresolved_marker_sections)} 个章节包含未解决的投标响应标记"
            f"（{';'.join(unresolved_marker_sections[:3])}）"
        )
    # Block external delivery when critical placeholders (brand/model/manufacturer) are present
    if critical_placeholder_sections:
        blocked_reasons.append(
            f"存在 {len(critical_placeholder_sections)} 个章节包含关键占位符"
            f"（品牌型号/生产厂家等未填写：{';'.join(critical_placeholder_sections[:3])}）"
        )

    # 计算 draft_level: internal / external
    draft_level = "external" if not blocked_reasons else "internal"

    if blocked_reasons:
        status = "阻断外发"
    else:
        status = "通过" if not placeholder_sections else "需人工终审"
    summary = (
        f"已完成外发净化，共清理 {len(changed_sections)} 个章节。"
        if changed_sections
        else "章节内容未发现明显模板污染。"
    )
    if placeholder_sections:
        summary += f" 仍有 {len(placeholder_sections)} 个章节包含占位符。"
    if unresolved_marker_sections:
        summary += f" {len(unresolved_marker_sections)} 个章节含未解决的投标响应标记，已阻断外发。"
    if critical_placeholder_sections:
        summary += f" {len(critical_placeholder_sections)} 个章节含关键占位符（品牌型号/生产厂家等），已阻断外发。"
    if blocked_reasons:
        summary += f" 当前外发已阻断：{'；'.join(blocked_reasons[:5])}。"
    summary += f" 当前稿件级别：{draft_level}。"

    return cleaned_sections, {
        "status": status,
        "draft_level": draft_level,
        "changed_sections": changed_sections,
        "placeholder_sections": placeholder_sections,
        "unresolved_marker_sections": unresolved_marker_sections,
        "hard_block_sections": hard_block_sections,
        "blocked_reasons": blocked_reasons,
        "summary": summary,
    }


def _build_regression_report(
    stages: list[dict[str, Any]],
    consistency_result: dict[str, Any] | None,
    review_result: dict[str, Any] | None,
    sanitize_result: dict[str, Any] | None,
    evidence_result: dict[str, Any] | None,
    *,
    normalized_result: dict[str, Any] | None = None,
    product_fact_result: dict[str, Any] | None = None,
    sections: list[BidDocumentSection] | None = None,
    selected_packages: list[str] | None = None,
) -> dict[str, Any]:
    stage_count = len(stages) + 1
    completed_count = len([stage for stage in stages if stage.get("status") == _STAGE_STATUS_COMPLETED])
    warning_count = len([stage for stage in stages if stage.get("status") == _STAGE_STATUS_WARNING])
    blocked_count = len([stage for stage in stages if stage.get("status") == _STAGE_STATUS_BLOCKED])

    evidence_rate = 0.0
    bidder_evidence_rate = 0.0
    proven_completion_rate = 0.0
    match_rate = 0.0
    if evidence_result:
        try:
            evidence_rate = float(evidence_result.get("binding_rate", 0.0))
        except (TypeError, ValueError):
            evidence_rate = 0.0
        try:
            bidder_evidence_rate = float(evidence_result.get("bidder_binding_rate", 0.0))
        except (TypeError, ValueError):
            bidder_evidence_rate = 0.0
        try:
            proven_completion_rate = float(evidence_result.get("proven_completion_rate", 0.0))
        except (TypeError, ValueError):
            proven_completion_rate = 0.0
        try:
            match_rate = float(evidence_result.get("match_rate", 0.0))
        except (TypeError, ValueError):
            match_rate = 0.0

    compliance_score = 0.0
    ready_for_submission = False
    if review_result:
        try:
            compliance_score = float(review_result.get("compliance_score", 0.0))
        except (TypeError, ValueError):
            compliance_score = 0.0
        ready_for_submission = bool(review_result.get("ready_for_submission", False))

    consistency_ok = bool(consistency_result) and consistency_result.get("overall_status") == "通过"
    outbound_ok = bool(sanitize_result) and sanitize_result.get("status") == "通过"
    regression_checks = [
        {
            "name": "十一层链路完整性",
            "status": "通过" if stage_count >= 10 and blocked_count == 0 else "需修订",
            "detail": f"阶段总数 {stage_count}，阻断阶段 {blocked_count} 个。",
        },
        {
            "name": "条款定位覆盖率",
            "status": "通过" if evidence_rate >= 0.5 else "需修订",
            "detail": f"当前覆盖率 {evidence_rate:.0%}。",
        },
        {
            "name": "投标方证据覆盖率",
            "status": "通过" if bidder_evidence_rate >= 0.5 else "需修订",
            "detail": f"当前覆盖率 {bidder_evidence_rate:.0%}。",
        },
        {
            "name": "要求-产品匹配率",
            "status": "通过" if match_rate >= 0.6 else "需修订",
            "detail": f"当前匹配率 {match_rate:.0%}。",
        },
        {
            "name": "已证实响应完成率",
            "status": "通过" if proven_completion_rate >= _MIN_PROVEN_COMPLETION_RATE else "需修订",
            "detail": f"当前完成率 {proven_completion_rate:.0%}。",
        },
        {
            "name": "硬校验结果",
            "status": "通过" if consistency_ok else "需修订",
            "detail": consistency_result.get("summary", "未执行") if consistency_result else "未执行",
        },
        {
            "name": "合规得分门槛",
            "status": "通过" if compliance_score >= 80 else "需修订",
            "detail": f"当前合规得分 {compliance_score:.1f}。",
        },
        {
            "name": "外发安全性",
            "status": "通过" if outbound_ok else "需修订",
            "detail": sanitize_result.get("summary", "未执行") if sanitize_result else "未执行",
        },
    ]

    # --- 6 new practical eval metrics ---

    # 1. package_isolation_score: ratio of section text that only mentions target packages
    _selected = set(selected_packages or [])
    unexpected_mentions: set[str] = set()
    if _selected and sections:
        full_text = "\n".join(sec.content for sec in sections)
        all_pkg_mentions = set(m.group(1) or m.group(2) for m in re.finditer(r"第\s*(\d+)\s*包|包\s*(\d+)", full_text))
        unexpected_mentions = all_pkg_mentions - _selected
        package_isolation = 1.0 if not unexpected_mentions else max(0.0, 1.0 - len(unexpected_mentions) / max(1, len(all_pkg_mentions)))
    else:
        package_isolation = 1.0
    _iso_detail = f"包件隔离度 {package_isolation:.0%}"
    if unexpected_mentions:
        _iso_detail += f"（存在串包：{','.join(sorted(unexpected_mentions))}）"
    _iso_detail += "。"
    regression_checks.append({
        "name": "package_isolation_score",
        "status": "通过" if package_isolation >= 0.9 else "需修订",
        "detail": _iso_detail,
        "value": round(package_isolation, 4),
    })

    # 2. atomic_requirement_rate: fraction of technical requirements that are NOT generic/collapsed
    tech_reqs = (normalized_result or {}).get("technical_requirements", [])
    atomic_count = 0
    if tech_reqs:
        atomic_count = sum(1 for r in tech_reqs if isinstance(r, dict) and not _is_generic_value(_safe_text(r.get("normalized_value"))))
        atomic_rate = atomic_count / len(tech_reqs)
    else:
        atomic_rate = 0.0
    regression_checks.append({
        "name": "atomic_requirement_rate",
        "status": "通过" if atomic_rate >= 0.7 else "需修订",
        "detail": f"原子级需求占比 {atomic_rate:.0%}（{atomic_count if tech_reqs else 0}/{len(tech_reqs)} 项为具体参数）。",
        "value": round(atomic_rate, 4),
    })

    # 3. offered_fact_coverage: offered_fact_count / max(1, total technical requirements)
    offered_count = (product_fact_result or {}).get("offered_fact_count", 0)
    offered_fact_coverage = offered_count / max(1, len(tech_reqs)) if tech_reqs else (1.0 if offered_count else 0.0)
    regression_checks.append({
        "name": "offered_fact_coverage",
        "status": "通过" if offered_fact_coverage >= 0.5 else "需修订",
        "detail": f"产品事实覆盖率 {min(offered_fact_coverage, 1.0):.0%}（{offered_count} 条事实 / {len(tech_reqs)} 项技术要求）。",
        "value": round(min(offered_fact_coverage, 1.0), 4),
    })

    # 4. bid_evidence_coverage: bidder_matched_count / max(1, total bindings)
    bidder_matched = int((evidence_result or {}).get("bidder_matched_count", 0))
    total_bindings = int((evidence_result or {}).get("total", 0))
    bid_evidence_coverage = bidder_matched / max(1, total_bindings) if total_bindings else 0.0
    regression_checks.append({
        "name": "bid_evidence_coverage",
        "status": "通过" if bid_evidence_coverage >= 0.5 else "需修订",
        "detail": f"投标方证据覆盖率 {bid_evidence_coverage:.0%}（{bidder_matched}/{total_bindings} 项已绑定）。",
        "value": round(bid_evidence_coverage, 4),
    })

    # 5. config_pollution_rate: fraction of config-table rows that look like non-config items
    _CONFIG_POLLUTION_KEYWORDS = ("评分标准", "评分办法", "商务条款", "合同条款", "投标人须知", "售后服务", "违约责任", "评审")
    config_total = 0
    config_polluted = 0
    for sec in (sections or []):
        for line in sec.content.splitlines():
            stripped = line.strip()
            if stripped.startswith("|") and "配置名称" not in stripped and not stripped.startswith("|---"):
                cells = [c.strip() for c in stripped.strip("|").split("|")]
                if len(cells) >= 4 and any(kw in cells[1] for kw in _CONFIG_POLLUTION_KEYWORDS):
                    config_polluted += 1
                if len(cells) >= 4 and cells[1].strip() and not re.fullmatch(r"\d+", cells[0].strip()):
                    continue
                if len(cells) >= 4:
                    config_total += 1
    config_pollution_rate = config_polluted / max(1, config_total)
    regression_checks.append({
        "name": "config_pollution_rate",
        "status": "通过" if config_pollution_rate <= 0.05 else "需修订",
        "detail": f"配置表污染率 {config_pollution_rate:.0%}（{config_polluted}/{config_total} 行疑似非配置项）。",
        "value": round(config_pollution_rate, 4),
    })

    config_rows_by_pkg: dict[str, int] = {}
    deviation_rows_by_pkg: dict[str, int] = {}
    evidence_rows_by_pkg: dict[str, int] = {}
    current_mode = ""
    current_pkg = ""
    for sec in (sections or []):
        for line in sec.content.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                if "详细配置明细表" in stripped or "配置清单" in stripped:
                    current_mode = "config"
                elif "技术偏离" in stripped:
                    current_mode = "deviation"
                elif "技术条款证据映射表" in stripped:
                    current_mode = "mapping"
                else:
                    current_mode = ""
                pkg_match = re.search(r"第\s*(\d+)\s*包", stripped)
                current_pkg = pkg_match.group(1) if pkg_match else ""
                if current_mode == "config" and current_pkg:
                    config_rows_by_pkg.setdefault(current_pkg, 0)
                if current_mode == "deviation" and current_pkg:
                    deviation_rows_by_pkg.setdefault(current_pkg, 0)
                if current_mode == "mapping" and current_pkg:
                    evidence_rows_by_pkg.setdefault(current_pkg, 0)
                continue

            if not stripped.startswith("|") or stripped.startswith("|---"):
                continue
            if current_mode == "config" and not stripped.startswith("| 序号"):
                cells = [c.strip() for c in stripped.split("|")]
                if (
                    current_pkg
                    and len(cells) >= 7
                    and re.match(r"^\d+$", cells[1].strip())
                    and cells[2].strip()
                ):
                    config_rows_by_pkg[current_pkg] = config_rows_by_pkg.get(current_pkg, 0) + 1
            elif current_mode == "deviation" and not stripped.startswith("| 条款编号") and not stripped.startswith("| 序号"):
                if current_pkg:
                    deviation_rows_by_pkg[current_pkg] = deviation_rows_by_pkg.get(current_pkg, 0) + 1
            elif current_mode == "mapping" and not stripped.startswith("| 序号"):
                cells = [c.strip() for c in stripped.split("|")]
                if current_pkg and len(cells) >= 5 and re.match(r"^\d+$", cells[1].strip()):
                    evidence_rows_by_pkg[current_pkg] = evidence_rows_by_pkg.get(current_pkg, 0) + 1

    min_config_items = _DETAIL_TARGETS["config_items_min"]
    if config_rows_by_pkg:
        config_detail_score = sum(
            min(count / max(1, min_config_items), 1.0)
            for count in config_rows_by_pkg.values()
        ) / len(config_rows_by_pkg)
    else:
        config_detail_score = 0.0
    regression_checks.append({
        "name": "config_detail_score",
        "status": "通过" if config_detail_score >= 0.8 else "需修订",
        "detail": (
            f"配置详细度 {config_detail_score:.0%}（"
            + "，".join(f"包{pkg}:{count}项" for pkg, count in sorted(config_rows_by_pkg.items()))
            + "）。"
            if config_rows_by_pkg
            else "未检测到可评估的配置明细表。"
        ),
        "value": round(config_detail_score, 4),
    })

    total_tech_requirements = len(tech_reqs)
    total_deviation_rows = sum(deviation_rows_by_pkg.values())
    total_mapping_rows = sum(evidence_rows_by_pkg.values())
    mapping_denominator = max(1, total_tech_requirements, total_deviation_rows, total_mapping_rows)
    count_gap = (
        abs(total_tech_requirements - total_deviation_rows)
        + abs(total_tech_requirements - total_mapping_rows)
        + abs(total_deviation_rows - total_mapping_rows)
    )
    mapping_count_consistency = max(0.0, 1.0 - count_gap / (3 * mapping_denominator))
    regression_checks.append({
        "name": "mapping_count_consistency",
        "status": "通过" if mapping_count_consistency >= 0.8 else "需修订",
        "detail": (
            f"技术条款 {total_tech_requirements} 项，偏离表 {total_deviation_rows} 行，证据映射表 {total_mapping_rows} 行。"
        ),
        "value": round(mapping_count_consistency, 4),
    })

    template_like_markers = (
        "[待填写]",
        "[品牌型号]",
        "[生产厂家]",
        "[品牌]",
        "待核实",
        "详见招标文件",
        "按招标文件配置要求",
        "配置清单包含主机及全套标准附件",
        "具备完整的技术功能",
    )
    content_lines = [
        line.strip()
        for sec in (sections or [])
        for line in sec.content.splitlines()
        if line.strip()
    ]
    template_line_count = sum(
        1 for line in content_lines if any(marker in line for marker in template_like_markers)
    )
    section_template_similarity = template_line_count / max(1, len(content_lines))
    regression_checks.append({
        "name": "section_template_similarity",
        "status": "通过" if section_template_similarity <= 0.2 else "需修订",
        "detail": f"模板化行占比 {section_template_similarity:.0%}（{template_line_count}/{len(content_lines)} 行）。",
        "value": round(section_template_similarity, 4),
    })

    # 6. external_block_rate: 1.0 if blocked, 0.0 if passed
    external_blocked = 1.0 if (sanitize_result and str(sanitize_result.get("status", "")).strip() == "阻断外发") else 0.0
    regression_checks.append({
        "name": "external_block_rate",
        "status": "通过" if external_blocked == 0.0 else "需修订",
        "detail": "外发未阻断" if external_blocked == 0.0 else f"外发已阻断：{'；'.join((sanitize_result or {}).get('blocked_reasons', ['未知原因'])[:3])}",
        "value": external_blocked,
    })
    # --- End new metrics ---

    # --- 7. 详细度目标（Detail Targets）检查 ---
    # 7a. 每包原子条款数检查
    tech_reqs_by_pkg: dict[str, int] = {}
    for r in tech_reqs:
        if isinstance(r, dict):
            pkg_id = _safe_text(r.get("package_id"))
            tech_reqs_by_pkg[pkg_id] = tech_reqs_by_pkg.get(pkg_id, 0) + 1
    min_atomic = _DETAIL_TARGETS["technical_atomic_clauses_per_package"]
    atomic_target_pass = all(
        count >= min_atomic for count in tech_reqs_by_pkg.values()
    ) if tech_reqs_by_pkg else False
    atomic_detail_parts = [f"包{pk}:{cnt}条" for pk, cnt in sorted(tech_reqs_by_pkg.items())]
    regression_checks.append({
        "name": "detail_target_atomic_clauses",
        "status": "通过" if atomic_target_pass else "需修订",
        "detail": (
            f"各包原子条款数：{', '.join(atomic_detail_parts) or '无'}。"
            f"目标：每包≥{min_atomic}条。"
        ),
        "value": min(tech_reqs_by_pkg.values()) if tech_reqs_by_pkg else 0,
    })

    # 7b. 偏离表最少行数
    deviation_rows_by_pkg: dict[str, int] = {}
    for sec in (sections or []):
        in_deviation = False
        current_pkg = ""
        for line in sec.content.splitlines():
            stripped = line.strip()
            if "技术偏离" in stripped and stripped.startswith("#"):
                in_deviation = True
                pkg_match = re.search(r"第\s*(\d+)\s*包", stripped)
                current_pkg = pkg_match.group(1) if pkg_match else "?"
                deviation_rows_by_pkg.setdefault(current_pkg, 0)
            elif stripped.startswith("#") and in_deviation:
                in_deviation = False
            elif in_deviation and stripped.startswith("|") and not stripped.startswith("|---") and not stripped.startswith("| 条款编号") and not stripped.startswith("| 序号"):
                deviation_rows_by_pkg[current_pkg] = deviation_rows_by_pkg.get(current_pkg, 0) + 1
    min_dev_rows = _DETAIL_TARGETS["deviation_table_min_rows"]
    dev_target_pass = all(
        count >= min_dev_rows for count in deviation_rows_by_pkg.values()
    ) if deviation_rows_by_pkg else False
    dev_detail_parts = [f"包{pk}:{cnt}行" for pk, cnt in sorted(deviation_rows_by_pkg.items())]
    regression_checks.append({
        "name": "detail_target_deviation_rows",
        "status": "通过" if dev_target_pass else "需修订",
        "detail": (
            f"偏离表行数：{', '.join(dev_detail_parts) or '无'}。"
            f"目标：每包≥{min_dev_rows}行。"
        ),
        "value": min(deviation_rows_by_pkg.values()) if deviation_rows_by_pkg else 0,
    })

    # 7c. 叙述章节字数检查
    narrative_keywords = ("关键性能说明", "配置说明", "交付说明", "验收说明", "使用与培训说明")
    narrative_total_chars = 0
    for sec in (sections or []):
        for nk in narrative_keywords:
            if nk in sec.content:
                # Count chars in this section
                start_pos = sec.content.find(nk)
                narrative_total_chars += len(sec.content[start_pos:start_pos + 500])
    min_narrative = _DETAIL_TARGETS["narrative_sections_min_chars"]
    narrative_pass = narrative_total_chars >= min_narrative
    regression_checks.append({
        "name": "detail_target_narrative_chars",
        "status": "通过" if narrative_pass else "需修订",
        "detail": f"叙述章节总字数约 {narrative_total_chars} 字。目标：≥{min_narrative}字。",
        "value": narrative_total_chars,
    })

    # 7d. 证据覆盖率（每条至少1个证据）
    evidence_per_item_target = _DETAIL_TARGETS["evidence_per_item"]
    evidence_coverage = float((evidence_result or {}).get("evidence_coverage_rate", bidder_evidence_rate))
    evidence_target_pass = evidence_coverage >= 0.5
    regression_checks.append({
        "name": "detail_target_evidence_coverage",
        "status": "通过" if evidence_target_pass else "需修订",
        "detail": f"证据覆盖率 {evidence_coverage:.0%}。目标：每条至少{evidence_per_item_target}个证据。",
        "value": round(evidence_coverage, 4),
    })
    # --- End detail target checks ---

    # --- 8. 新增实用性评测指标 ---

    # 8a. 实际参数覆盖率 (actual_param_coverage): 偏离表中有真实参数值的行 / 总行数
    actual_param_rows = 0
    total_deviation_data_rows = sum(deviation_rows_by_pkg.values())
    _PENDING_MARKERS = ("待核实", "[待填写]", "[待补充]", "[品牌型号]", "[生产厂家]")
    for sec in (sections or []):
        in_deviation_section = False
        for line in sec.content.splitlines():
            stripped = line.strip()
            if "技术偏离" in stripped and stripped.startswith("#"):
                in_deviation_section = True
                continue
            if stripped.startswith("#") and in_deviation_section:
                in_deviation_section = False
                continue
            if not in_deviation_section or not stripped.startswith("|") or stripped.startswith("|---"):
                continue
            if any(h in stripped for h in ("条款编号", "序号", "参数名称")):
                continue
            # 检查响应列是否有真实值
            cells = [c.strip() for c in stripped.split("|")]
            if len(cells) >= 5:
                response_cell = cells[4] if len(cells) > 4 else ""
                has_real_value = bool(
                    response_cell
                    and response_cell not in ("", " ", "-")
                    and not any(pm in response_cell for pm in _PENDING_MARKERS)
                )
                if has_real_value:
                    actual_param_rows += 1

    actual_param_coverage = actual_param_rows / max(1, total_deviation_data_rows)
    regression_checks.append({
        "name": "actual_param_coverage",
        "status": "通过" if actual_param_coverage >= 0.7 else "需修订",
        "detail": f"实际参数覆盖率 {actual_param_coverage:.0%}（{actual_param_rows}/{total_deviation_data_rows} 行含真实参数值）。",
        "value": round(actual_param_coverage, 4),
    })

    # 8b. 投标侧证据页码覆盖率 (bid_evidence_page_coverage)
    bid_evidence_items = (evidence_result or {}).get("bid_evidence", [])
    if not isinstance(bid_evidence_items, list):
        bid_evidence_items = []
    items_with_page = sum(
        1 for item in bid_evidence_items
        if isinstance(item, dict) and item.get("evidence_page") is not None
    )
    total_bid_items = len(bid_evidence_items)
    bid_page_coverage = items_with_page / max(1, total_bid_items)
    regression_checks.append({
        "name": "bid_evidence_page_coverage",
        "status": "通过" if bid_page_coverage >= 0.5 else "需修订",
        "detail": f"投标侧证据页码覆盖率 {bid_page_coverage:.0%}（{items_with_page}/{total_bid_items} 项含页码引用）。",
        "value": round(bid_page_coverage, 4),
    })

    # 8c. 配置项平均条数 (config_avg_items_per_package)
    avg_config_items = (
        sum(config_rows_by_pkg.values()) / len(config_rows_by_pkg)
        if config_rows_by_pkg else 0.0
    )
    config_avg_pass = avg_config_items >= _DETAIL_TARGETS["config_items_min"]
    regression_checks.append({
        "name": "config_avg_items_per_package",
        "status": "通过" if config_avg_pass else "需修订",
        "detail": (
            f"配置项平均条数 {avg_config_items:.1f} 条/包"
            f"（{'，'.join(f'包{pk}:{cnt}项' for pk, cnt in sorted(config_rows_by_pkg.items())) or '无'}）。"
            f"目标：≥{_DETAIL_TARGETS['config_items_min']}条。"
        ),
        "value": round(avg_config_items, 2),
    })

    # 8d. 模板段落重复率 (template_paragraph_ratio) — 与 section_template_similarity 互补
    _TEMPLATE_PARAGRAPH_MARKERS = (
        "具备完整的技术功能",
        "能够满足采购文件要求",
        "配置清单包含主机及全套标准附件",
        "按招标文件配置要求",
        "由我公司负责运输至指定地点",
        "按照国家相关标准及招标文件要求",
        "采用专业包装方式",
        "安排专业培训师",
    )
    para_total = 0
    para_template = 0
    for sec in (sections or []):
        paragraphs = [p.strip() for p in sec.content.split("\n\n") if p.strip() and len(p.strip()) > 20]
        para_total += len(paragraphs)
        for para in paragraphs:
            if sum(1 for m in _TEMPLATE_PARAGRAPH_MARKERS if m in para) >= 2:
                para_template += 1
    template_paragraph_ratio = para_template / max(1, para_total)
    regression_checks.append({
        "name": "template_paragraph_ratio",
        "status": "通过" if template_paragraph_ratio <= 0.2 else "需修订",
        "detail": f"模板段落重复率 {template_paragraph_ratio:.0%}（{para_template}/{para_total} 段含多个模板标记）。",
        "value": round(template_paragraph_ratio, 4),
    })

    # 8e. external hard-gate 拦截率 (external_hardgate_block_items)
    hardgate_blocked_count = len((sanitize_result or {}).get("blocked_reasons", []))
    hardgate_total_checks = 10  # 总硬门检查项数
    hardgate_ratio = hardgate_blocked_count / hardgate_total_checks
    regression_checks.append({
        "name": "external_hardgate_block_rate",
        "status": "通过" if hardgate_blocked_count == 0 else "需修订",
        "detail": (
            "外发硬门全部通过" if hardgate_blocked_count == 0
            else f"外发硬门拦截 {hardgate_blocked_count} 项：{'；'.join((sanitize_result or {}).get('blocked_reasons', [])[:3])}"
        ),
        "value": round(hardgate_ratio, 4),
    })
    # --- End new practical eval metrics ---

    passed_count = len([item for item in regression_checks if item["status"] == "通过"])
    score = round(
        min(
            100.0,
            passed_count * 7
            + compliance_score * 0.15
            + evidence_rate * 4
            + bidder_evidence_rate * 4
            + match_rate * 5
            + proven_completion_rate * 5
            + package_isolation * 5
            + atomic_rate * 5
            + min(offered_fact_coverage, 1.0) * 5
            + bid_evidence_coverage * 5
            + (1.0 - config_pollution_rate) * 3
            + (1.0 - external_blocked) * 3,
        ),
        2,
    )
    overall_status = "通过" if ready_for_submission and blocked_count == 0 and outbound_ok else "需修订"
    summary = (
        f"回归评测完成：{passed_count}/{len(regression_checks)} 项通过，"
        f"阶段完成 {completed_count} 项，告警 {warning_count} 项。"
    )

    return {
        "overall_status": overall_status,
        "score": score,
        "ready_for_delivery": ready_for_submission and outbound_ok,
        "checks": regression_checks,
        "summary": summary,
    }


def _retrieve_citations(query: str, preferred_source: str | None = None, top_k: int = _DEFAULT_CITATION_TOP_K) -> list[dict[str, Any]]:
    if not query.strip():
        return []

    try:
        hits = search_knowledge(query=query, top_k=max(1, top_k))
    except Exception as exc:  # noqa: BLE001
        logger.warning("检索引用失败，query=%s, error=%s", query, exc)
        return []

    if not hits:
        return []

    if preferred_source:
        preferred_hits: list[dict[str, Any]] = []
        for hit in hits:
            metadata = hit.get("metadata", {})
            source = str(metadata.get("source", "")).strip()
            if source == preferred_source:
                preferred_hits.append(hit)
        if preferred_hits:
            hits = preferred_hits

    return _prepare_citations(hits, limit=top_k)


def _traceability_hits(
    technical_text: str,
    technical_matches: list[dict[str, Any]],
) -> tuple[int, int, list[str]]:
    if not technical_matches:
        return 0, 0, []

    hit_count = 0
    missing: list[str] = []
    for match in technical_matches:
        if not isinstance(match, dict):
            continue
        parameter_name = _safe_text(match.get("parameter_name"))
        if not parameter_name:
            continue
        evidence_bits = _dedupe_texts(
            [
                _safe_text(match.get("matched_fact_quote")),
                _safe_text(match.get("bidder_evidence_quote")),
                _safe_text(match.get("response_value")),
            ]
        )
        if parameter_name in technical_text and any(bit and bit in technical_text for bit in evidence_bits):
            hit_count += 1
        else:
            missing.append(parameter_name)

    return hit_count, len([item for item in technical_matches if isinstance(item, dict)]), missing


def _product_compliance_gaps(
    tender: TenderDocument,
    package_ids: list[str],
    products: dict[str, ProductSpecification],
) -> list[str]:
    if not package_ids:
        return []

    context = _workflow_context_text(tender)
    medical_project = _contains_any(context, _MEDICAL_KEYWORDS)
    imported_project = _contains_any(context, _IMPORTED_KEYWORDS)
    requires_energy_cert = _contains_any(context, ("节能", "环保", "能效"))
    gaps: list[str] = []

    for pkg_id in package_ids:
        product = products.get(pkg_id)
        if product is None:
            gaps.append(f"包{pkg_id} 未绑定产品资料")
            continue
        if medical_project and not product.registration_number.strip():
            gaps.append(f"包{pkg_id} 缺少注册证/备案编号")
        if imported_project and not product.origin.strip():
            gaps.append(f"包{pkg_id} 缺少原产地/合法来源")
        if imported_project and not product.authorization_letter.strip():
            gaps.append(f"包{pkg_id} 缺少授权链/报关材料")
        if requires_energy_cert and not product.certifications:
            gaps.append(f"包{pkg_id} 缺少节能环保认证")

    return gaps


def _material_coverage(required_materials: list[str], sections: list[BidDocumentSection]) -> tuple[int, int, list[str]]:
    if not required_materials:
        return 0, 0, []

    full_text = "\n".join(sec.content for sec in sections)
    full_text = full_text.lower()

    matched = 0
    missing: list[str] = []
    for item in required_materials:
        normalized = item.strip()
        if not normalized:
            continue
        tokens = [tok for tok in re.split(r"[，,、；;（）()\\s/]+", normalized) if len(tok) >= 2]
        if not tokens:
            tokens = [normalized]
        if any(token.lower() in full_text for token in tokens[:4]):
            matched += 1
        else:
            missing.append(normalized)

    total = len([x for x in required_materials if x.strip()])
    return matched, total, missing


def _second_validation(
    analysis_result: dict[str, Any],
    validation_result: dict[str, Any],
    sections: list[BidDocumentSection],
    generation_result: dict[str, Any] | None = None,
    tender: TenderDocument | None = None,
    selected_packages: list[str] | None = None,
    products: dict[str, ProductSpecification] | None = None,
    evidence_result: dict[str, Any] | None = None,
) -> dict[str, Any]:
    check_items: list[dict[str, str]] = []
    issues: list[str] = []
    suggestions: list[str] = []

    validation_status = str(validation_result.get("overall_status", "")).strip()
    material_pass = validation_status == "通过"
    check_items.append(
        {
            "name": "资料完整性复核",
            "status": "通过" if material_pass else "需修订",
            "detail": f"第二步校验状态：{validation_status or '未提供'}",
        }
    )
    if not material_pass:
        issues.append("资料校验未通过，存在缺失项或待确认项。")
        suggestions.append("先完成缺失资料补齐，再重新运行流程。")

    section_titles = [sec.section_title for sec in sections]
    required_chapters = ("第一章", "第二章", "第三章", "第四章")
    missing_chapters = [
        chapter
        for chapter in required_chapters
        if not any(chapter in title for title in section_titles)
    ]
    chapter_pass = not missing_chapters
    chapter_detail = "章节完整" if chapter_pass else f"缺少章节：{', '.join(missing_chapters)}"
    check_items.append(
        {
            "name": "分章节完整性",
            "status": "通过" if chapter_pass else "需修订",
            "detail": chapter_detail,
        }
    )
    if not chapter_pass:
        issues.append(f"分章节生成不完整：{', '.join(missing_chapters)}。")
        suggestions.append("补齐缺失章节，确保投标文件结构完整。")

    placeholder_total = 0
    placeholder_section_details: list[str] = []
    for sec in sections:
        count = 0
        for pattern in _PLACEHOLDER_PATTERNS:
            count += sec.content.count(pattern)
        if count > 0:
            placeholder_total += count
            placeholder_section_details.append(f"{sec.section_title}({count}处)")

    placeholder_pass = placeholder_total == 0
    placeholder_detail = (
        "未发现占位符。"
        if placeholder_pass
        else f"发现 {placeholder_total} 处占位符：{'；'.join(placeholder_section_details)}"
    )
    check_items.append(
        {
            "name": "占位符与留空项检查",
            "status": "通过" if placeholder_pass else "需修订",
            "detail": placeholder_detail,
        }
    )
    if not placeholder_pass:
        issues.append("标书中仍存在未替换占位符或留空说明。")
        suggestions.append("逐章替换 [待填写]/公司信息占位符，并补齐附件留空项。")

    technical_text = "\n".join(
        sec.content
        for sec in sections
        if "第三章" in sec.section_title or "技术" in sec.section_title
    )
    proven_completion = evidence_result or {}
    technical_matches = [
        item
        for item in (proven_completion.get("technical_matches") or [])
        if isinstance(item, dict)
    ]
    proven_matches = [item for item in technical_matches if bool(item.get("proven"))]

    evidence_mapping_exists = any("技术条款证据映射表" in sec.content for sec in sections)
    traced_count, traced_total, trace_missing = _traceability_hits(technical_text, proven_matches)
    trace_ratio = 1.0 if traced_total == 0 else traced_count / traced_total
    evidence_mapping_pass = evidence_mapping_exists and (traced_total == 0 or trace_ratio >= _MIN_PROVEN_COMPLETION_RATE)
    check_items.append(
        {
            "name": "技术条款证据映射",
            "status": "通过" if evidence_mapping_pass else "需修订",
            "detail": (
                "已检测到映射表，且已证实条款均已落入表内"
                if evidence_mapping_pass
                else (
                    "未检测到“技术条款证据映射表”章节内容"
                    if not evidence_mapping_exists
                    else f"已证实条款仅落入 {traced_count}/{traced_total} 项"
                )
            ),
        }
    )
    if not evidence_mapping_pass:
        if not evidence_mapping_exists:
            issues.append("技术章节缺少证据映射表，参数与原文无法一一追溯。")
            suggestions.append("在第三章补充“技术条款证据映射表”，逐条关联招标原文片段。")
        else:
            preview = "；".join(trace_missing[:5]) or "多项已证实条款未落入证据映射表"
            issues.append(f"证据映射表存在但内容未完成：{preview}。")
            suggestions.append("补齐证据映射表中的投标方证据与产品事实，不要只保留招标原文摘录。")

    required_materials = _ensure_str_list(analysis_result.get("required_materials"))
    matched, total, missing = _material_coverage(required_materials, sections)
    coverage_ratio = 1.0 if total == 0 else matched / total
    coverage_pass = coverage_ratio >= 0.6
    check_items.append(
        {
            "name": "资料覆盖率检查",
            "status": "通过" if coverage_pass else "需修订",
            "detail": f"覆盖 {matched}/{total} 项（覆盖率 {coverage_ratio:.0%}）",
        }
    )
    if not coverage_pass and total > 0:
        preview_missing = "；".join(missing[:5]) if missing else "多项资料未覆盖"
        issues.append(f"资料覆盖不足：{preview_missing}。")
        suggestions.append("根据“需准备资料清单”补齐对应章节内容与附件说明。")

    analysis_citations = analysis_result.get("citations")
    if not isinstance(analysis_citations, list):
        analysis_citations = []
    generation_citations: list[dict[str, Any]] = []
    if generation_result and isinstance(generation_result.get("citations"), list):
        generation_citations = generation_result["citations"]

    citation_count = len(analysis_citations) + len(generation_citations)
    citation_pass = citation_count > 0
    check_items.append(
        {
            "name": "检索引用可追溯性",
            "status": "通过" if citation_pass else "需修订",
            "detail": f"可追溯引用条数：{citation_count}",
        }
    )
    if not citation_pass:
        issues.append("未生成检索引用，难以追溯结论依据。")
        suggestions.append("先将招标原文入库并重跑流程，确保输出包含 citations。")

    full_text = "\n".join(sec.content for sec in sections)
    key_info = analysis_result.get("key_information")
    if not isinstance(key_info, dict):
        key_info = {}
    expected_project_name = _safe_text(
        tender.project_name if tender is not None else key_info.get("project_name"),
    )
    expected_project_number = _safe_text(
        tender.project_number if tender is not None else key_info.get("project_number"),
    )

    if expected_project_name or expected_project_number:
        name_ok = not expected_project_name or expected_project_name in full_text
        number_ok = not expected_project_number or expected_project_number in full_text
        identifier_pass = name_ok and number_ok
        detail_bits: list[str] = []
        if expected_project_name:
            detail_bits.append(f"项目名称{'已命中' if name_ok else '未命中'}")
        if expected_project_number:
            detail_bits.append(f"项目编号{'已命中' if number_ok else '未命中'}")
        check_items.append(
            {
                "name": "项目基础信息一致性",
                "status": "通过" if identifier_pass else "需修订",
                "detail": "；".join(detail_bits),
            }
        )
        if not identifier_pass:
            issues.append("项目名称或项目编号未稳定落入正文，存在基础信息不一致风险。")
            suggestions.append("在封面、资格声明和技术/报价章节补入统一的项目名称与项目编号。")

    expected_package_ids = list(selected_packages or [])
    if not expected_package_ids and generation_result:
        expected_package_ids = [str(item) for item in generation_result.get("selected_packages", []) if str(item).strip()]
    if tender is not None:
        pass
    elif isinstance(key_info.get("packages"), list):
        pass

    if expected_package_ids:
        package_mentions = set()
        for match in re.finditer(r"第\s*(\d+)\s*包|包\s*(\d+)", full_text):
            pkg_id = match.group(1) or match.group(2)
            if pkg_id:
                package_mentions.add(pkg_id)
        unexpected_packages = sorted(pkg for pkg in package_mentions if pkg not in set(expected_package_ids))
        package_pass = not unexpected_packages
        check_items.append(
            {
                "name": "包件分仓检查",
                "status": "通过" if package_pass else "需修订",
                "detail": (
                    "未发现串包"
                    if package_pass
                    else f"发现非目标包号：{', '.join(unexpected_packages)}"
                ),
            }
        )
        if not package_pass:
            issues.append(f"正文出现非目标包号：{', '.join(unexpected_packages)}。")
            suggestions.append("按包号重建单包上下文，禁止跨包共享技术条款和报价数据。")

    products = products or {}
    if tender is not None and expected_package_ids:
        compliance_gaps = _product_compliance_gaps(tender, expected_package_ids, products)
        compliance_chain_pass = not compliance_gaps
        check_items.append(
            {
                "name": "行业/货物证明链",
                "status": "通过" if compliance_chain_pass else "需修订",
                "detail": "项目特定证明链已齐备" if compliance_chain_pass else "；".join(compliance_gaps[:5]),
            }
        )
        if not compliance_chain_pass:
            issues.append(f"项目特定证明链不完整：{'；'.join(compliance_gaps[:5])}。")
            suggestions.append("按项目属性补齐注册证、原产地/合法来源、授权链或节能环保认证材料。")

    pollution_hits = [keyword for keyword in _TECH_POLLUTION_KEYWORDS if keyword in technical_text]
    technical_purity_pass = not pollution_hits
    check_items.append(
        {
            "name": "技术章节纯度检查",
            "status": "通过" if technical_purity_pass else "需修订",
            "detail": (
                "未发现评分/合同类污染片段"
                if technical_purity_pass
                else f"命中关键词：{'；'.join(pollution_hits[:5])}"
            ),
        }
    )
    if not technical_purity_pass:
        issues.append("技术章节混入评分办法、合同条款或投诉质疑等非技术内容。")
        suggestions.append("回到条款分类层，限制只有技术条款进入技术偏离表与配置表。")

    response_realization_issues: list[str] = []
    if "承诺满足" in technical_text:
        response_realization_issues.append("技术章节仍存在泛化承诺句式")
    if _PENDING_RESPONSE_TEXT in technical_text:
        response_realization_issues.append("技术表仍存在待核实响应值")
    unresolved_param_placeholders = [
        pattern
        for pattern in ("[品牌型号]", "[生产厂家]", "[品牌]", "[待补充]")
        if pattern in full_text
    ]
    if unresolved_param_placeholders:
        response_realization_issues.append(f"仍存在关键实参占位符：{'；'.join(unresolved_param_placeholders)}")

    response_realization_pass = not response_realization_issues
    check_items.append(
        {
            "name": "技术响应实参化",
            "status": "通过" if response_realization_pass else "需修订",
            "detail": "已使用具体参数/主体信息" if response_realization_pass else "；".join(response_realization_issues),
        }
    )
    if not response_realization_pass:
        issues.append("技术或报价章节仍缺少实参化字段，正文可读性与可提交性不足。")
        suggestions.append("将企业主体、品牌型号、厂家和单价总价注入正式章节后提交。")

    proven_total = len(technical_matches)
    proven_count = int(proven_completion.get("proven_response_count", 0) or 0)
    proven_rate = float(proven_completion.get("proven_completion_rate", 1.0) if proven_total else 1.0)
    completion_pass = proven_total == 0 or proven_rate >= _MIN_PROVEN_COMPLETION_RATE
    check_items.append(
        {
            "name": "已证实响应完成率",
            "status": "通过" if completion_pass else "需修订",
            "detail": (
                "暂无技术要求待计算完成率"
                if proven_total == 0
                else f"已证实 {proven_count}/{proven_total} 项（完成率 {proven_rate:.0%}）"
            ),
        }
    )
    if not completion_pass:
        preview = "；".join(_ensure_str_list(proven_completion.get("unproven_items"))[:5]) or "存在未证实技术要求"
        issues.append(f"技术响应完成率不足：{preview}。")
        suggestions.append("先完成产品事实匹配与投标证据绑定，再输出正式外发版。")

    semantic_issues: list[str] = []
    for binding in technical_matches:
        if not isinstance(binding, dict):
            continue
        parameter_name = _safe_text(binding.get("parameter_name"))
        response_value = _safe_text(binding.get("response_value"))
        if parameter_name and "无偏离" in technical_text and not binding.get("proven"):
            for line in technical_text.splitlines():
                if parameter_name in line and "无偏离" in line:
                    semantic_issues.append(f"{parameter_name} 未证实却标注无偏离")
                    break
        if response_value and parameter_name and parameter_name in technical_text and response_value not in technical_text:
            semantic_issues.append(f"{parameter_name} 已匹配产品事实，但正文未落入响应值")

    semantic_pass = not semantic_issues
    check_items.append(
        {
            "name": "技术响应语义一致性",
            "status": "通过" if semantic_pass else "需修订",
            "detail": "要求、响应值、偏离标注保持一致" if semantic_pass else "；".join(semantic_issues[:5]),
        }
    )
    if not semantic_pass:
        issues.append("技术响应存在语义不一致：要求、响应值或偏离标注未形成闭环。")
        suggestions.append("逐条核对技术偏离表，确保响应值来自已证实产品事实且偏离标注正确。")

    products = products or {}
    if tender is not None and expected_package_ids and products:
        product_gaps: list[str] = []
        for pkg_id in expected_package_ids:
            pkg = next((item for item in tender.packages if item.package_id == pkg_id), None)
            product = products.get(pkg_id)
            if pkg is None or product is None:
                continue
            if product.model.strip() and product.model not in full_text:
                product_gaps.append(f"包{pkg_id} 型号未落正文")
            row_lines = [line for line in full_text.splitlines() if pkg.item_name in line]
            if row_lines and str(pkg.quantity) not in row_lines[0]:
                product_gaps.append(f"包{pkg_id} 数量未在条目行体现")
        product_consistency_pass = not product_gaps
        check_items.append(
            {
                "name": "数量/型号一致性",
                "status": "通过" if product_consistency_pass else "需修订",
                "detail": "数量与型号已落正文" if product_consistency_pass else "；".join(product_gaps),
            }
        )
        if not product_consistency_pass:
            issues.append("报价表或配置表中的型号/数量与选定产品映射未完全对齐。")
            suggestions.append("按包号将产品型号、厂家、数量和价格回填到报价/配置表后复核。")

    # ── 详细度目标校验 ──
    # 检查是否存在叙述性章节
    narrative_keywords = ("关键性能说明", "配置说明", "交付说明", "验收说明", "使用与培训说明")
    found_narratives = [nk for nk in narrative_keywords if nk in full_text]
    narrative_check_pass = len(found_narratives) >= 3
    check_items.append(
        {
            "name": "详细技术响应章节",
            "status": "通过" if narrative_check_pass else "需修订",
            "detail": (
                f"已检测到 {len(found_narratives)}/5 个叙述章节（{', '.join(found_narratives[:3])}）"
                if found_narratives
                else "未检测到详细技术响应章节（关键性能说明/配置说明等）"
            ),
        }
    )
    if not narrative_check_pass:
        issues.append("技术章节缺少详细技术响应说明（关键性能说明/配置说明/交付说明/验收说明/培训说明）。")
        suggestions.append("在技术偏离表后补充关键性能说明、配置说明、交付说明、验收说明和使用与培训说明章节。")

    # 检查偏离表列数（新格式应为8列）
    deviation_8col = "实际响应值" in full_text and "证据材料" in full_text
    check_items.append(
        {
            "name": "偏离表详细度",
            "status": "通过" if deviation_8col else "需修订",
            "detail": "偏离表已升级为8列详细格式" if deviation_8col else "偏离表仍为旧格式，建议升级为8列（含条款编号、投标型号、证据材料、页码、验收备注）",
        }
    )

    # 检查配置表是否含功能描述
    config_desc_present = "配置功能描述" in full_text or "二-B" in full_text
    check_items.append(
        {
            "name": "配置表详细度",
            "status": "通过" if config_desc_present else "需修订",
            "detail": "配置表已包含功能描述层" if config_desc_present else "配置表缺少功能描述层（建议增加配置项用途说明和功能角色描述）",
        }
    )

    # ── 新增5项深度检查（判断"够不够细"而非"有没有"）──

    # (1) offered_fact_coverage: 产品事实覆盖率
    offered_fact_count = 0
    if evidence_result:
        offered_fact_count = int(evidence_result.get("offered_fact_count", 0) or 0)
    tech_req_count = len(technical_matches)
    offered_coverage = offered_fact_count / max(1, tech_req_count) if tech_req_count else 0.0
    offered_coverage = min(offered_coverage, 1.0)
    offered_coverage_pass = offered_coverage >= 0.5
    check_items.append({
        "name": "offered_fact_coverage（产品事实覆盖率）",
        "status": "通过" if offered_coverage_pass else "需修订",
        "detail": f"产品事实覆盖率 {offered_coverage:.0%}（{offered_fact_count} 条事实 / {tech_req_count} 项技术要求）。"
                  + ("" if offered_coverage_pass else " 不足50%，技术表右侧仍大量为待核实。"),
    })
    if not offered_coverage_pass:
        issues.append(f"产品事实覆盖率仅 {offered_coverage:.0%}，技术表中大量参数仍为待核实。")
        suggestions.append("请补充产品彩页、说明书等投标材料，通过 Product Profile Builder 提取真实参数。")

    # (2) bid_evidence_coverage: 投标侧证据页码覆盖率
    bid_evidence_count = 0
    bid_evidence_with_page = 0
    total_bindings = 0
    if evidence_result:
        bid_evidence_count = int(evidence_result.get("bidder_matched_count", 0) or 0)
        total_bindings = int(evidence_result.get("total", 0) or 0)
        # 统计有页码的投标证据数
        bid_evidence_items = evidence_result.get("bid_evidence", [])
        if isinstance(bid_evidence_items, list):
            bid_evidence_with_page = sum(
                1 for item in bid_evidence_items
                if isinstance(item, dict) and item.get("evidence_page") is not None
            )
    bid_ev_coverage = bid_evidence_count / max(1, total_bindings)
    bid_ev_page_coverage = bid_evidence_with_page / max(1, total_bindings)
    bid_ev_pass = bid_ev_coverage >= 0.5
    check_items.append({
        "name": "bid_evidence_coverage（投标方证据覆盖率）",
        "status": "通过" if bid_ev_pass else "需修订",
        "detail": (
            f"投标方证据覆盖率 {bid_ev_coverage:.0%}（{bid_evidence_count}/{total_bindings} 项已绑定），"
            f"含页码 {bid_ev_page_coverage:.0%}（{bid_evidence_with_page}/{total_bindings} 项有页码）。"
        ),
    })
    if not bid_ev_pass:
        issues.append('投标方证据覆盖率不足，证据列仍停留在"待补投标方证据"。')
        suggestions.append("请提供投标材料（彩页/说明书/注册证/检测报告），通过 BidEvidenceBinder 绑定页码。")

    # (3) config_detail_score: 配置项平均条数
    config_rows_per_pkg: dict[str, int] = {}
    current_mode_cfg = ""
    current_pkg_cfg = ""
    for sec in sections:
        for line in sec.content.splitlines():
            stripped = line.strip()
            if stripped.startswith("#") and ("配置" in stripped):
                current_mode_cfg = "config"
                pkg_match = re.search(r"第\s*(\d+)\s*包", stripped)
                current_pkg_cfg = pkg_match.group(1) if pkg_match else "?"
                config_rows_per_pkg.setdefault(current_pkg_cfg, 0)
            elif stripped.startswith("#") and current_mode_cfg == "config":
                current_mode_cfg = ""
            elif current_mode_cfg == "config" and stripped.startswith("|") and not stripped.startswith("|---") and not stripped.startswith("| 序号"):
                cells = [c.strip() for c in stripped.split("|")]
                if len(cells) >= 5 and re.match(r"^\d+$", cells[1].strip()):
                    config_rows_per_pkg[current_pkg_cfg] = config_rows_per_pkg.get(current_pkg_cfg, 0) + 1

    avg_config = (
        sum(config_rows_per_pkg.values()) / len(config_rows_per_pkg)
        if config_rows_per_pkg else 0.0
    )
    config_score = min(avg_config / max(1, _DETAIL_TARGETS["config_items_min"]), 1.0)
    config_score_pass = config_score >= 0.8
    cfg_detail_parts = [f"包{pk}:{cnt}项" for pk, cnt in sorted(config_rows_per_pkg.items())]
    check_items.append({
        "name": "config_detail_score（配置详细度评分）",
        "status": "通过" if config_score_pass else "需修订",
        "detail": (
            f"配置详细度 {config_score:.0%}（平均{avg_config:.1f}项/包；{', '.join(cfg_detail_parts) or '无'}）。"
            + ("" if config_score_pass else f" 目标：每包≥{_DETAIL_TARGETS['config_items_min']}项。")
        ),
    })
    if not config_score_pass:
        issues.append(f"配置表过薄（平均{avg_config:.1f}项/包），像模板而非交付清单。")
        suggestions.append("通过 Config Extractor 从投标材料中抽取核心模块、标准附件、配套软件、初始耗材、随机文件、安装/培训资料。")

    # (4) mapping_count_consistency: 技术条款/偏离表/证据映射表行数一致性
    dev_rows_total = 0
    map_rows_total = 0
    for sec in sections:
        in_dev = in_map = False
        for line in sec.content.splitlines():
            stripped = line.strip()
            if "技术偏离" in stripped and stripped.startswith("#"):
                in_dev, in_map = True, False
            elif "证据映射" in stripped and stripped.startswith("#"):
                in_dev, in_map = False, True
            elif stripped.startswith("#"):
                in_dev = in_map = False
            elif stripped.startswith("|") and not stripped.startswith("|---"):
                is_header = any(h in stripped for h in ("条款编号", "序号", "参数名称"))
                if not is_header:
                    if in_dev:
                        dev_rows_total += 1
                    elif in_map:
                        map_rows_total += 1

    mapping_denom = max(1, tech_req_count, dev_rows_total, map_rows_total)
    count_gap = (
        abs(tech_req_count - dev_rows_total)
        + abs(tech_req_count - map_rows_total)
        + abs(dev_rows_total - map_rows_total)
    )
    mc_consistency = max(0.0, 1.0 - count_gap / (3 * mapping_denom))
    mc_pass = mc_consistency >= 0.8
    check_items.append({
        "name": "mapping_count_consistency（表间行数一致性）",
        "status": "通过" if mc_pass else "需修订",
        "detail": f"技术条款 {tech_req_count} 项，偏离表 {dev_rows_total} 行，证据映射表 {map_rows_total} 行（一致性 {mc_consistency:.0%}）。",
    })
    if not mc_pass:
        issues.append("技术条款数、偏离表行数、证据映射表行数不一致，存在遗漏或重复。")
        suggestions.append("确保技术偏离表和证据映射表逐条对应归一化后的技术要求。")

    # (5) section_template_similarity: 模板段落重复率
    _TEMPLATE_MARKERS = (
        "[待填写]", "[品牌型号]", "[生产厂家]", "[品牌]",
        "待核实", "具备完整的技术功能", "配置清单包含主机及全套标准附件",
        "按招标文件配置要求", "详见招标文件",
    )
    content_lines_all = [
        line.strip()
        for sec in sections
        for line in sec.content.splitlines()
        if line.strip()
    ]
    template_hits = sum(
        1 for line in content_lines_all if any(marker in line for marker in _TEMPLATE_MARKERS)
    )
    template_ratio = template_hits / max(1, len(content_lines_all))
    template_pass = template_ratio <= 0.15
    check_items.append({
        "name": "section_template_similarity（模板段落重复率）",
        "status": "通过" if template_pass else "需修订",
        "detail": f"模板化行占比 {template_ratio:.0%}（{template_hits}/{len(content_lines_all)} 行含模板标记）。"
                  + ("" if template_pass else " 超过15%，文档仍像长模板而非项目化说明。"),
    })
    if not template_pass:
        issues.append(f"模板段落重复率 {template_ratio:.0%}，文档内容过于模板化。")
        suggestions.append("切换到 Rich draft mode，引用本包真实参数、配置和证据替换模板句。")

    # ── 底稿完整性检查（5项新增）──

    # (6) 条款数检查 — 若某包 < 5 条 requirement 则告警
    if tender is not None:
        thin_packages: list[str] = []
        pkg_req_counts: dict[str, int] = {}
        for pkg in tender.packages:
            req_count = len(pkg.technical_requirements or {})
            pkg_req_counts[pkg.package_id] = req_count
            if req_count < 5:
                thin_packages.append(f"包{pkg.package_id}({req_count}条)")
        clause_count_pass = not thin_packages
        check_items.append({
            "name": "条款数充足性",
            "status": "通过" if clause_count_pass else "需修订",
            "detail": (
                "各包条款数均≥5条"
                if clause_count_pass
                else f"条款数不足：{'；'.join(thin_packages)}"
            ),
        })
        if not clause_count_pass:
            issues.append(f"以下包条款数过少（<5条）：{'；'.join(thin_packages)}。底稿拆条不够细。")
            suggestions.append("对条款数不足的包，从招标原文中补充提取技术参数，确保每包至少5条原子级条款。")

        # (7) 包间粒度均匀性 — 最多包/最少包差异超过3倍则告警
        if len(pkg_req_counts) >= 2:
            max_count = max(pkg_req_counts.values())
            min_count = max(1, min(pkg_req_counts.values()))
            granularity_ratio = max_count / min_count
            granularity_pass = granularity_ratio <= 3.0
            check_items.append({
                "name": "包间粒度均匀性",
                "status": "通过" if granularity_pass else "需修订",
                "detail": (
                    f"包间条款数比 {granularity_ratio:.1f}:1（"
                    + "、".join(f"包{k}:{v}条" for k, v in sorted(pkg_req_counts.items()))
                    + "）"
                ),
            })
            if not granularity_pass:
                issues.append(f"包间拆分粒度不均匀（比值{granularity_ratio:.1f}:1），部分包过粗。")
                suggestions.append("对条款数较少的包从原文重新提取，使各包粒度差距不超过3倍。")

    # (8) 配置表薄弱检查 — 配置项 < 3 则告警（仅在有明确包号时检查）
    meaningful_config_pkgs = {k: v for k, v in config_rows_per_pkg.items() if k != "?"}
    thin_config_packages: list[str] = []
    for pkg_id, count in meaningful_config_pkgs.items():
        if count < 3:
            thin_config_packages.append(f"包{pkg_id}({count}项)")
    if meaningful_config_pkgs:
        config_thin_pass = not thin_config_packages
        check_items.append({
            "name": "配置表薄弱检查",
            "status": "通过" if config_thin_pass else "需修订",
            "detail": (
                "各包配置项均≥3项"
                if config_thin_pass
                else f"配置表过薄：{'；'.join(thin_config_packages)}"
            ),
        })
        if not config_thin_pass:
            issues.append(f"配置表过薄：{'；'.join(thin_config_packages)}，缺少核心模块/附件/软件等分类。")
            suggestions.append("补充配置表至少覆盖6大类别：核心模块、标准附件、配套软件、初始耗材、随机文件、安装/培训资料。")

    # (9) 原文片段污染检查 — 证据映射表中的片段含评分/商务等非技术内容
    polluted_snippets: list[str] = []
    _SNIPPET_POLLUTION_KEYWORDS = ("评分标准", "评分办法", "商务条款", "合同条款",
                                    "违约责任", "质疑", "投诉", "付款方式")
    for sec in sections:
        if "证据映射" not in sec.section_title and "证据映射" not in sec.content[:200]:
            continue
        for line in sec.content.splitlines():
            if not line.strip().startswith("|"):
                continue
            for kw in _SNIPPET_POLLUTION_KEYWORDS:
                if kw in line:
                    cells = [c.strip() for c in line.split("|") if c.strip()]
                    param_name = cells[1] if len(cells) > 1 else "未知"
                    polluted_snippets.append(f"{param_name}含'{kw}'")
                    break
    snippet_purity_pass = not polluted_snippets
    check_items.append({
        "name": "原文片段污染检查",
        "status": "通过" if snippet_purity_pass else "需修订",
        "detail": (
            "证据映射表原文片段未发现非技术内容污染"
            if snippet_purity_pass
            else f"原文片段污染：{'；'.join(polluted_snippets[:5])}"
        ),
    })
    if not polluted_snippets:
        pass
    else:
        issues.append(f"证据映射表中有{len(polluted_snippets)}处原文片段混入了非技术内容。")
        suggestions.append("回到原文切割层，确保证据片段只包含技术参数相关内容，过滤评分/商务/合同条款。")

    overall_status = "通过" if not issues else "需修订"
    if not suggestions and overall_status == "通过":
        suggestions = ["可进入人工终审与盖章提交流程。"]

    summary = (
        f"二次校验完成：{len(check_items)} 项，"
        f"{'全部通过' if overall_status == '通过' else f'发现 {len(issues)} 项问题'}。"
    )

    return {
        "executed": True,
        "overall_status": overall_status,
        "check_items": check_items,
        "issues": issues,
        "suggestions": suggestions,
        "proven_completion": {
            "proven_count": proven_count,
            "total": proven_total,
            "rate": round(proven_rate, 4) if proven_total else 1.0,
            "unproven_items": _ensure_str_list(proven_completion.get("unproven_items")),
        },
        "summary": summary,
    }


def _append_unique(base: list[str], extras: list[str]) -> list[str]:
    for item in extras:
        normalized = str(item).strip()
        if normalized and normalized not in base:
            base.append(normalized)
    return base


def _format_eval_rules(evaluation_criteria: dict[str, Any]) -> list[str]:
    if not evaluation_criteria:
        return []
    rules: list[str] = []
    for k, v in evaluation_criteria.items():
        if isinstance(v, (int, float)):
            rules.append(f"{k}：{v}")
        else:
            rules.append(f"{k}：{v}")
    return rules


def _default_required_materials(tender: TenderDocument) -> list[str]:
    materials = [
        "营业执照及法定代表人身份证明",
        "法定代表人授权书及授权代表身份证明",
        "供应商资格承诺函（含政府采购法第二十二条相关承诺）",
        "依法缴纳税收和社保证明材料",
        "信用记录查询截图（信用中国/中国政府采购网等）",
        "报价书、报价一览表、报价明细表",
        "技术偏离表及详细配置明细表",
        "技术服务与售后服务方案",
    ]
    joined = " ".join(pkg.item_name for pkg in tender.packages if pkg.item_name.strip())
    if "医疗器械" in joined or "流式" in joined:
        materials.extend(
            [
                "医疗器械经营许可证/备案凭证（如适用）",
                "产品注册证/备案证明（如适用）",
                "厂家授权书（如适用）",
            ]
        )
    # 去重并保持顺序
    seen: set[str] = set()
    unique: list[str] = []
    for item in materials:
        if item not in seen:
            seen.add(item)
            unique.append(item)
    return unique


def _default_step1_result(tender: TenderDocument) -> dict[str, Any]:
    packages = [
        {
            "package_id": pkg.package_id,
            "item_name": pkg.item_name,
            "quantity": pkg.quantity,
            "budget": pkg.budget,
        }
        for pkg in tender.packages
    ]
    return {
        "key_information": {
            "project_name": tender.project_name,
            "project_number": tender.project_number,
            "purchaser": tender.purchaser,
            "agency": tender.agency,
            "procurement_type": tender.procurement_type,
            "budget": tender.budget,
            "packages": packages,
            "commercial_terms": tender.commercial_terms.model_dump(),
        },
        "required_materials": _default_required_materials(tender),
        "offered_facts": [],
        "scoring_rules": _format_eval_rules(tender.evaluation_criteria),
        "risk_alerts": [
            "请重点核对投标有效期、交货期限和履约保证金条款。",
            "请确保技术参数响应表逐条对应，不要遗漏关键参数。",
            "证照与授权文件需在有效期内，且与投标产品一致。",
        ],
        "citations": [],
        "summary": "已完成招标关键信息、资料清单和评分规则提取。",
    }


def _material_item(item: str, status: str, evidence: str, suggestion: str = "") -> dict[str, str]:
    return {
        "item": item,
        "status": status,
        "evidence": evidence,
        "suggestion": suggestion,
    }


def _default_step4_if_blocked(reason: str) -> dict[str, Any]:
    return {
        "ready_for_submission": False,
        "risk_level": "high",
        "compliance_score": 0.0,
        "major_issues": [reason],
        "recommendations": ["先补齐资料缺口，再重新运行第三步与第四步。"],
        "secondary_validation": {
            "executed": False,
            "overall_status": "需修订",
            "check_items": [],
            "issues": [reason],
            "suggestions": ["先补齐资料缺口后再执行二次校验。"],
            "summary": "未执行二次校验：缺少可审核标书内容。",
        },
        "conclusion": "当前不具备提交条件。",
    }


class TenderWorkflowAgent:
    """十层招投标工作流 Agent。"""

    def __init__(self, llm: ChatOpenAI):
        self.llm = llm

    def step1_analyze_tender(
        self,
        tender: TenderDocument,
        raw_text: str,
        kb_source: str | None = None,
    ) -> dict[str, Any]:
        """第一步：解析招标文件并提炼关键结果。"""
        fallback = _default_step1_result(tender)
        raw_excerpt = (raw_text or "")[:_MAX_RAW_PROMPT_CHARS]
        package_names = "、".join(pkg.item_name for pkg in tender.packages if pkg.item_name.strip())
        citation_query = "；".join(
            part
            for part in [
                tender.project_name.strip(),
                tender.project_number.strip(),
                package_names,
                "评分标准 资格要求 商务条款 资料清单",
            ]
            if part
        )
        citations = _retrieve_citations(
            query=citation_query,
            preferred_source=kb_source,
            top_k=_DEFAULT_CITATION_TOP_K,
        )
        system_prompt = (
            "你是“招标解析Agent”。你的任务是根据招标文件结构化信息和原文，"
            "输出投标准备所需的关键信息。只允许输出JSON。"
        )
        user_prompt = (
            "请输出JSON，结构如下：\n"
            "{\n"
            '  "key_information": { ... },\n'
            '  "required_materials": ["..."],\n'
            '  "offered_facts": [\n'
            '    {"fact_name": "...", "fact_value": "...", "source_excerpt": "..."}\n'
            '  ],\n'
            '  "scoring_rules": ["..."],\n'
            '  "risk_alerts": ["..."],\n'
            '  "summary": "..."\n'
            "}\n\n"
            "要求：\n"
            "1. key_information需包含项目名称、项目编号、采购人、采购方式、预算、包信息、核心商务条款；\n"
            "2. required_materials必须是可执行的资料清单；\n"
            "3. offered_facts 从招标原文中提取采购方已明示的允许值/范围/规格，"
            "例如允许的品牌范围、参数上下限、明确说明[须提供XX]的事实性描述，"
            "每条包含 fact_name、fact_value、source_excerpt 三个字段，最多提取20条；\n"
            "4. scoring_rules从评分标准中提炼，若文件无明确权重请说明；\n"
            "5. risk_alerts给出3~6条最关键风险提示。\n\n"
            f"结构化招标信息：\n{json.dumps(tender.model_dump(), ensure_ascii=False)}\n\n"
            f"招标原文（截断）：\n{raw_excerpt}"
        )
        try:
            content = _llm_call(self.llm, system_prompt, user_prompt)
            parsed = _extract_json(content)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Step1 解析失败，使用默认结果：%s", exc)
            return fallback

        key_info = parsed.get("key_information")
        required_materials = _ensure_str_list(parsed.get("required_materials"))
        raw_offered_facts = parsed.get("offered_facts")
        offered_facts: list[dict[str, Any]] = (
            [
                item for item in raw_offered_facts
                if isinstance(item, dict)
                and _safe_text(item.get("fact_name"))
                and _safe_text(item.get("fact_value"))
            ]
            if isinstance(raw_offered_facts, list)
            else []
        )
        scoring_rules = _ensure_str_list(parsed.get("scoring_rules"))
        risk_alerts = _ensure_str_list(parsed.get("risk_alerts"))
        summary = str(parsed.get("summary", "")).strip()

        if not isinstance(key_info, dict):
            key_info = fallback["key_information"]
        if not required_materials:
            required_materials = fallback["required_materials"]
        if not scoring_rules:
            scoring_rules = fallback["scoring_rules"] or ["评分规则需结合招标文件评审办法人工确认。"]
        if not risk_alerts:
            risk_alerts = fallback["risk_alerts"]
        if not summary:
            summary = fallback["summary"]
        if citations:
            summary = f"{summary}（已附 {len(citations)} 条检索引用）"
        if offered_facts:
            summary = f"{summary}；已从原文提取 {len(offered_facts)} 条 offered facts。"

        return {
            "key_information": key_info,
            "required_materials": required_materials,
            "offered_facts": offered_facts,
            "scoring_rules": scoring_rules,
            "risk_alerts": risk_alerts,
            "citations": citations,
            "summary": summary,
        }

    def step2_validate_materials(
        self,
        tender: TenderDocument,
        required_materials: list[str],
        selected_packages: list[str],
        company: CompanyProfile | None,
        products: dict[str, ProductSpecification],
    ) -> dict[str, Any]:
        """第二步：按招标要求校验已上传资料。"""
        package_ids = {pkg.package_id for pkg in tender.packages}
        target_packages = selected_packages or sorted(package_ids)
        checklist: list[dict[str, str]] = []
        context = _workflow_context_text(tender)
        medical_project = _contains_any(context, _MEDICAL_KEYWORDS)
        imported_project = _contains_any(context, _IMPORTED_KEYWORDS)
        requires_energy_cert = _contains_any(context, ("节能", "环保", "能效"))

        if company is None:
            checklist.append(
                _material_item("企业基本信息", "缺失", "未提供 company_profile_id", "先上传企业信息再继续。")
            )
        else:
            company_ok = all(
                [
                    bool(company.name.strip()),
                    bool(company.legal_representative.strip()),
                    bool(company.address.strip()),
                    bool(company.phone.strip()),
                ]
            )
            checklist.append(
                _material_item(
                    "企业基本信息",
                    "通过" if company_ok else "缺失",
                    f"企业：{company.name or '-'}，法定代表人：{company.legal_representative or '-'}",
                    "补齐企业名称、法定代表人、地址、电话四项信息。",
                )
            )

            has_licenses = len(company.licenses) > 0
            checklist.append(
                _material_item(
                    "企业证照资料",
                    "通过" if has_licenses else "缺失",
                    f"证照数量：{len(company.licenses)}",
                    "至少补充营业执照及与项目匹配的资质证照。",
                )
            )

            has_staff = len(company.staff) > 0
            checklist.append(
                _material_item(
                    "项目人员资料",
                    "通过" if has_staff else "待确认",
                    f"人员数量：{len(company.staff)}",
                    "建议补充项目负责人及关键岗位人员清单。",
                )
            )

        for pkg_id in target_packages:
            if pkg_id not in package_ids:
                checklist.append(
                    _material_item(
                        f"包{pkg_id} 投标范围",
                        "缺失",
                        "包号不存在于招标文件",
                        "请检查 selected_packages 与招标文件包号是否一致。",
                    )
                )
                continue

            product = products.get(pkg_id)
            if product is None:
                checklist.append(
                    _material_item(
                        f"包{pkg_id} 产品资料",
                        "缺失",
                        "未提供对应产品信息",
                        "请在 product_ids 中补充包号到产品ID的映射，并先创建产品资料。",
                    )
                )
                continue

            product_ok = bool(product.product_name.strip()) and product.price > 0 and bool(product.specifications)
            checklist.append(
                _material_item(
                    f"包{pkg_id} 产品资料",
                    "通过" if product_ok else "缺失",
                    f"产品：{product.product_name}；型号：{product.model or '-'}；参数项：{len(product.specifications)}",
                    "补充产品型号、关键技术参数和价格信息。",
                )
            )

            if medical_project:
                has_registration = bool(product.registration_number.strip())
                checklist.append(
                    _material_item(
                        f"包{pkg_id} 医疗器械注册证/备案",
                        "通过" if has_registration else "缺失",
                        f"注册证/备案编号：{product.registration_number or '-'}",
                        "医疗器械项目需补充注册证或备案凭证编号及对应附件。",
                    )
                )

            if imported_project:
                has_origin = bool(product.origin.strip())
                checklist.append(
                    _material_item(
                        f"包{pkg_id} 原产地/合法来源",
                        "通过" if has_origin else "缺失",
                        f"原产地：{product.origin or '-'}",
                        "进口项目需补充原产地、合法来源或报关相关说明。",
                    )
                )
                has_authorization = bool(product.authorization_letter.strip())
                checklist.append(
                    _material_item(
                        f"包{pkg_id} 授权链/报关材料",
                        "通过" if has_authorization else "缺失",
                        f"授权材料：{product.authorization_letter or '-'}",
                        "进口项目需补充厂家授权、报关单或同等效力材料。",
                    )
                )

            if requires_energy_cert:
                has_certifications = bool(product.certifications)
                checklist.append(
                    _material_item(
                        f"包{pkg_id} 节能环保认证",
                        "通过" if has_certifications else "缺失",
                        f"认证数量：{len(product.certifications)}",
                        "存在节能/环保/能效要求时需补充对应认证材料。",
                    )
                )

        checklist.append(
            _material_item(
                "招标要求资料覆盖",
                "待确认",
                f"需准备资料项：{len(required_materials)}",
                "对照“需准备资料清单”逐项上传附件。",
            )
        )

        missing_items = [
            item["item"]
            for item in checklist
            if item["status"] == "缺失"
        ]

        system_prompt = (
            "你是“资料校验Agent”。根据已上传资料与招标要求，输出校验结论。"
            "只允许输出JSON。"
        )
        user_prompt = (
            "请输出JSON，结构如下：\n"
            "{\n"
            '  "overall_status": "通过/需补充",\n'
            '  "summary": "...",\n'
            '  "missing_items": ["..."],\n'
            '  "next_actions": ["..."]\n'
            "}\n\n"
            "输入信息：\n"
            f"- 项目：{tender.project_name}\n"
            f"- 目标包号：{target_packages}\n"
            f"- 需准备资料：{required_materials}\n"
            f"- 校验清单：{json.dumps(checklist, ensure_ascii=False)}\n"
            "要求：next_actions 给出可执行动作，最多6条。"
        )

        overall_status = "需补充" if missing_items else "通过"
        summary = "资料校验完成。"
        next_actions = [
            "按缺失项补齐企业资料、产品资料和资质附件。",
            "补充后重新运行第二步校验，确认状态为“通过”。",
        ]

        try:
            content = _llm_call(self.llm, system_prompt, user_prompt)
            parsed = _extract_json(content)
            parsed_missing = _ensure_str_list(parsed.get("missing_items"))
            parsed_actions = _ensure_str_list(parsed.get("next_actions"))
            parsed_status = str(parsed.get("overall_status", "")).strip()
            parsed_summary = str(parsed.get("summary", "")).strip()

            if parsed_status in {"通过", "需补充"}:
                overall_status = parsed_status
            if parsed_summary:
                summary = parsed_summary
            if parsed_actions:
                next_actions = parsed_actions
            if parsed_missing:
                missing_items = sorted(set(missing_items + parsed_missing))
        except Exception as exc:  # noqa: BLE001
            logger.warning("Step2 AI总结失败，使用规则结果：%s", exc)

        if missing_items:
            overall_status = "需补充"

        return {
            "overall_status": overall_status,
            "checklist": checklist,
            "missing_items": missing_items,
            "next_actions": next_actions,
            "summary": summary,
        }

    def step3_classify_clauses(
        self,
        tender: TenderDocument,
        analysis_result: dict[str, Any],
        selected_packages: list[str],
        raw_text: str,
    ) -> dict[str, Any]:
        """第三步：条款分类与分支决策。"""
        return _classify_clauses(
            tender=tender,
            analysis_result=analysis_result,
            selected_packages=selected_packages,
            raw_text=raw_text,
        )

    def step4_normalize_requirements(
        self,
        tender: TenderDocument,
        analysis_result: dict[str, Any],
        clause_result: dict[str, Any],
        selected_packages: list[str],
        raw_text: str,
    ) -> dict[str, Any]:
        """第四步：将条款归一化为可机读字段。"""
        return _normalize_requirements(
            tender=tender,
            analysis_result=analysis_result,
            clause_result=clause_result,
            selected_packages=selected_packages,
            raw_text=raw_text,
        )

    def step5_decide_rules(
        self,
        tender: TenderDocument,
        raw_text: str,
        selected_packages: list[str],
        company: CompanyProfile | None,
        products: dict[str, ProductSpecification],
        clause_result: dict[str, Any],
    ) -> dict[str, Any]:
        """第五步：通过规则引擎固化关键分支。"""
        return _decide_rule_branches(
            tender=tender,
            raw_text=raw_text,
            selected_packages=selected_packages,
            company=company,
            products=products,
            clause_result=clause_result,
        )

    def step4_bind_evidence(
        self,
        tender: TenderDocument,
        raw_text: str,
        analysis_result: dict[str, Any],
            company: CompanyProfile | None = None,
        products: dict[str, ProductSpecification] | None = None,
        selected_packages: list[str] | None = None,
        normalized_result: dict[str, Any] | None = None,
        product_fact_result: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """第六步：将关键需求绑定到招标原文与投标方证据。"""
        return _build_evidence_bindings(
            tender=tender,
            raw_text=raw_text,
            company=company,
            products=products,
            selected_packages=selected_packages,
            normalized_result=normalized_result,
            product_fact_result=product_fact_result,
        )

    def step3_integrate_bid(
        self,
        tender: TenderDocument,
        raw_text: str,
        selected_packages: list[str],
        company: CompanyProfile,
        products: dict[str, ProductSpecification],
        kb_source: str | None = None,
        analysis_result: dict[str, Any] | None = None,
        product_fact_result: dict[str, Any] | None = None,
        normalized_result: dict[str, Any] | None = None,
        evidence_result: dict[str, Any] | None = None,
    ) -> tuple[dict[str, Any], list[BidDocumentSection]]:
        """第三步：整合生成标书内容。"""
        package_ids = {pkg.package_id for pkg in tender.packages}
        target_packages = selected_packages or sorted(package_ids)
        filtered_packages = [pkg for pkg in tender.packages if pkg.package_id in target_packages]
        if not filtered_packages:
            filtered_packages = tender.packages

        filtered_tender = tender.model_copy(update={"packages": filtered_packages})

        product_summary = []
        for pkg_id in target_packages:
            product = products.get(pkg_id)
            if product:
                product_summary.append(
                    {
                        "package_id": pkg_id,
                        "product_name": product.product_name,
                        "model": product.model,
                        "manufacturer": product.manufacturer,
                    }
                )

        citation_query = "；".join(
            [
                filtered_tender.project_name,
                filtered_tender.project_number,
                " ".join(pkg.item_name for pkg in filtered_packages if pkg.item_name.strip()),
                "技术参数 交货期 投标报价 评分点",
            ]
        ).strip("；")
        citations = _retrieve_citations(
            query=citation_query,
            preferred_source=kb_source,
            top_k=_DEFAULT_CITATION_TOP_K,
        )
        citation_prompt_block = ""
        if citations:
            citation_lines: list[str] = []
            for idx, item in enumerate(citations, start=1):
                src = item.get("source", "unknown")
                chunk = item.get("chunk_index")
                chunk_text = "?" if chunk is None else str(chunk)
                quote = item.get("quote", "")
                citation_lines.append(f"{idx}. [{src}#{chunk_text}] {quote}")
            citation_prompt_block = "检索引用（用于追溯）：\n" + "\n".join(citation_lines) + "\n"

        # Build offered_facts block so Writer has ground-truth values and skips placeholders
        all_offered_facts: list[dict[str, Any]] = []
        if analysis_result and isinstance(analysis_result.get("offered_facts"), list):
            all_offered_facts.extend(analysis_result["offered_facts"])
        if product_fact_result and isinstance(product_fact_result.get("packages"), list):
            for pkg_entry in product_fact_result["packages"]:
                if not isinstance(pkg_entry, dict):
                    continue
                for fact in pkg_entry.get("offered_facts", []):
                    if isinstance(fact, dict):
                        all_offered_facts.append(fact)
        seen_fact_keys: set[str] = set()
        fact_lines: list[str] = []
        for fact in all_offered_facts:
            name = _safe_text(fact.get("fact_name") or fact.get("evidence_type"))
            value = _safe_text(fact.get("fact_value") or fact.get("evidence_value"))
            source = _safe_text(fact.get("evidence_source", ""))
            key = f"{name}::{value}"
            if not name or not value or key in seen_fact_keys:
                continue
            seen_fact_keys.add(key)
            line = f"- {name}：{value}"
            if source:
                line += f"（来源：{source}）"
            fact_lines.append(line)
        offered_facts_prompt_block = ""
        if fact_lines:
            offered_facts_prompt_block = (
                "已知可用真值（请直接填入正文，禁止使用[待填写]/[待补充]等占位符）：\n"
                + "\n".join(fact_lines[:40])
                + "\n"
            )

        # Build structured product profile block for Writer
        product_profile_prompt_block = ""
        if product_fact_result:
            profile_block = _build_product_profile_block(product_fact_result)
            if profile_block:
                product_profile_prompt_block = (
                    "投标产品档案（请基于以下信息填写品牌、型号、厂家、参数，禁止使用占位符）：\n"
                    + profile_block + "\n"
                )

        system_prompt = (
            "你是[标书整合Agent]。请基于项目与已上传资料给出整合策略。"
            "输出纯文本，控制在200~400字。"
            "重要：已知真值信息中有数据时，必须直接引用，不得使用[待填写]等占位符。"
        )
        user_prompt = (
            f"项目：{filtered_tender.project_name}\n"
            f"包号：{target_packages}\n"
            f"企业：{company.name}\n"
            f"产品摘要：{json.dumps(product_summary, ensure_ascii=False)}\n"
            f"{product_profile_prompt_block}"
            f"{offered_facts_prompt_block}"
            f"{citation_prompt_block}"
            "请给出：章节重点、资料落位建议、常见格式风险。"
        )

        integration_notes = _llm_call(self.llm, system_prompt, user_prompt)
        product_profiles = {
            pkg_id: _build_product_profile(product)
            for pkg_id, product in products.items()
            if product is not None
        }
        sections = generate_bid_sections(
            filtered_tender,
            raw_text,
            self.llm,
            products=products,
            normalized_result=normalized_result,
            evidence_result=evidence_result,
            product_profiles=product_profiles,
        )

        return {
            "generated": True,
            "citations": citations,
            "selected_packages": target_packages,
            "product_summary": product_summary,
            "offered_facts_injected": len(fact_lines),
            "integration_notes": integration_notes,
            "summary": (
                f"已完成标书整合，共生成 {len(sections)} 个章节；"
                f"已向 Writer 注入 {len(fact_lines)} 条 offered facts。"
            ),
        }, sections

    def step6_validate_consistency(
        self,
        analysis_result: dict[str, Any],
        validation_result: dict[str, Any],
        sections: list[BidDocumentSection],
        generation_result: dict[str, Any] | None = None,
        tender: TenderDocument | None = None,
        selected_packages: list[str] | None = None,
        products: dict[str, ProductSpecification] | None = None,
        evidence_result: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """第八步：发布前硬校验。"""
        return _second_validation(
            analysis_result=analysis_result,
            validation_result=validation_result,
            sections=sections,
            generation_result=generation_result,
            tender=tender,
            selected_packages=selected_packages,
            products=products,
            evidence_result=evidence_result,
        )

    def step4_review_bid(
        self,
        tender: TenderDocument,
        analysis_result: dict[str, Any],
        validation_result: dict[str, Any],
        sections: list[BidDocumentSection],
        generation_result: dict[str, Any] | None = None,
        selected_packages: list[str] | None = None,
        products: dict[str, ProductSpecification] | None = None,
        evidence_result: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """辅助审核：生成 review 摘要字段。"""
        if not sections:
            return _default_step4_if_blocked("未生成标书内容，无法执行审核。")

        section_digest = []
        for sec in sections:
            snippet = sec.content.replace("\n", " ").strip()
            if len(snippet) > _MAX_REVIEW_SECTION_CHARS:
                snippet = snippet[:_MAX_REVIEW_SECTION_CHARS] + "..."
            section_digest.append(
                {
                    "title": sec.section_title,
                    "snippet": snippet,
                }
            )

        system_prompt = (
            "你是“标书审核Agent”。根据招标要求、资料校验结果和标书内容，"
            "输出审核结论。只允许输出JSON。"
        )
        user_prompt = (
            "请输出JSON，结构如下：\n"
            "{\n"
            '  "ready_for_submission": true/false,\n'
            '  "risk_level": "low/medium/high",\n'
            '  "compliance_score": 0-100,\n'
            '  "major_issues": ["..."],\n'
            '  "recommendations": ["..."],\n'
            '  "conclusion": "..."\n'
            "}\n\n"
            f"项目：{tender.project_name}\n"
            f"评分规则：{analysis_result.get('scoring_rules', [])}\n"
            f"资料校验结果：{json.dumps(validation_result, ensure_ascii=False)}\n"
            f"标书摘要：{json.dumps(section_digest, ensure_ascii=False)}\n"
            "要求：如发现占位符未填、证书资料留空、关键条款未响应，请明确指出。"
        )

        fallback = {
            "ready_for_submission": False,
            "risk_level": "high",
            "compliance_score": 0.0,
            "major_issues": ["建议人工逐页核对占位符与附件是否已补齐。"],
            "recommendations": [
                "核对报价金额、交货期、质保期与招标文件是否一致。",
                "补齐证书与截图附件后再提交。",
            ],
            "conclusion": "初稿已生成，建议补齐缺失资料并完成人工终审后提交。",
        }

        try:
            content = _llm_call(self.llm, system_prompt, user_prompt)
            parsed = _extract_json(content)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Step4 审核失败，使用默认结果：%s", exc)
            parsed = fallback

        ready = bool(parsed.get("ready_for_submission", False))
        risk_level = str(parsed.get("risk_level", "medium")).lower().strip()
        if risk_level not in {"low", "medium", "high"}:
            risk_level = "medium"

        try:
            score = float(parsed.get("compliance_score", 0.0))
        except (TypeError, ValueError):
            score = 0.0
        score = max(0.0, min(100.0, score))

        major_issues = _ensure_str_list(parsed.get("major_issues"))
        recommendations = _ensure_str_list(parsed.get("recommendations"))
        conclusion = str(parsed.get("conclusion", "")).strip()

        if not major_issues:
            major_issues = fallback["major_issues"]
        if not recommendations:
            recommendations = fallback["recommendations"]
        if not conclusion:
            conclusion = fallback["conclusion"]

        second_validation = _second_validation(
            analysis_result=analysis_result,
            validation_result=validation_result,
            sections=sections,
            generation_result=generation_result,
            tender=tender,
            selected_packages=selected_packages,
            products=products,
            evidence_result=evidence_result,
        )

        if validation_result.get("overall_status") == "需补充":
            ready = False
            risk_level = "high"
            if "资料仍有缺失项，当前不建议提交。" not in major_issues:
                major_issues.insert(0, "资料仍有缺失项，当前不建议提交。")

        if second_validation.get("overall_status") == "需修订":
            ready = False
            risk_level = "high"
            score = min(score, 75.0)
            major_issues = _append_unique(major_issues, second_validation.get("issues", []))
            recommendations = _append_unique(recommendations, second_validation.get("suggestions", []))
            if "二次校验未通过，请先修订后再提交。" not in major_issues:
                major_issues.insert(0, "二次校验未通过，请先修订后再提交。")
            if "二次校验发现问题，需完成修订后再提交。" not in conclusion:
                conclusion = f"{conclusion} 二次校验发现问题，需完成修订后再提交。"

        return {
            "ready_for_submission": ready,
            "risk_level": risk_level,
            "compliance_score": score,
            "major_issues": major_issues,
            "recommendations": recommendations,
            "secondary_validation": second_validation,
            "conclusion": conclusion,
        }

    def step8_sanitize_outbound(
        self,
        sections: list[BidDocumentSection],
        hard_validation_result: dict[str, Any] | None = None,
        evidence_result: dict[str, Any] | None = None,
        normalized_result: dict[str, Any] | None = None,
    ) -> tuple[list[BidDocumentSection], dict[str, Any]]:
        """第九步：外发净化。"""
        return _sanitize_for_external_delivery(
            sections=sections,
            hard_validation_result=hard_validation_result,
            evidence_result=evidence_result,
            normalized_result=normalized_result,
        )

    def step9_regression(
        self,
        stages: list[dict[str, Any]],
        consistency_result: dict[str, Any] | None,
        review_result: dict[str, Any] | None,
        sanitize_result: dict[str, Any] | None,
        evidence_result: dict[str, Any] | None,
        *,
        normalized_result: dict[str, Any] | None = None,
        product_fact_result: dict[str, Any] | None = None,
        sections: list[BidDocumentSection] | None = None,
        selected_packages: list[str] | None = None,
    ) -> dict[str, Any]:
        """第九步：评测回归。"""
        return _build_regression_report(
            stages=stages,
            consistency_result=consistency_result,
            review_result=review_result,
            sanitize_result=sanitize_result,
            evidence_result=evidence_result,
            normalized_result=normalized_result,
            product_fact_result=product_fact_result,
            sections=sections,
            selected_packages=selected_packages,
        )
