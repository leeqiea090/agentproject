"""投标底稿修复脚本 v2 — 针对「投标底稿_检验科购置全自动电泳仪等设备.docx」

修复清单:
1. 删除 6 条脏段落（"说明：当前已结构化生成…"）
2. 技术偏离表: 补齐缺失行 / 修正断裂文字 / 去除重复行
3. 配置清单: 按招标参数推导标配明细
4. 审查表 / 评审表: 预填"是否满足"→"待复核" + 指引备注 + 补全对应章节
5. 清洗响应列冗余提示文字
"""

from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Row

# ── 路径 ──────────────────────────────────────────────────────────
ORIGINAL = Path("/Users/lq/Downloads/投标底稿_检验科购置全自动电泳仪等设备.docx")
REVISED  = ORIGINAL.with_name("投标底稿_检验科购置全自动电泳仪等设备_修复版.docx")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  技术偏离表补丁数据（按包号）
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# 每个 patch 条目: (insert_after_seq, param_name, requirement, response_hint)
# insert_after_seq = 在第 N 行后插入; 0 = 在表头后插入
TECH_PATCHES: dict[int, list[tuple[int, str, str, str]]] = {
    # T3 – 包2 特种蛋白分析仪: 缺「比色杯」
    3: [
        (10, "比色杯", "可清洗", "【待填写：实际配置】"),
    ],
    # T5 – 包3 荧光操作仪: row4 文字断裂 → 不新增，在 fix 里修正
    5: [],
    # T7 – 包4 荧光显微镜: 缺 目镜/观察镜筒/物镜转换器/透射光源类型
    7: [
        (2, "目镜", "10X宽视野目镜，视场数≥20", "【待填写：实际配置】"),
        (3, "观察镜筒", "三目镜筒，视场数≥20", "【待填写：实际配置】"),
        (4, "物镜转换器", "★手动，≥4档，可同时装载≥4个物镜", "【待填写：实际配置】"),
        # 透射光源类型 放到最后
        (99, "透射光源类型", "卤素灯", "【待填写：实际配置】"),
    ],
    # T11 – 包6 流式细胞: 缺 方法/六色淋巴试剂盒; row7 缺 PE; rows 9-10/16-17 重复
    11: [
        (1, "方法", "流式细胞术", "【待填写：实际配置】"),
        (2, "六色淋巴细胞亚群试剂盒", "★具备原厂配套的六色淋巴细胞亚群试剂盒，并提供三类医疗器械注册证", "【待填写：实际配置及注册证编号】"),
    ],
}

# 需要修正的单元格 (table_index, row_index, col_index, old_substring, new_value)
CELL_FIXES = [
    # T5 row4 "、加样针" → 正确参数
    (5, 4, 1, "、加样针", "加样针"),
    (5, 4, 2, "、加样针", "≥3针"),
    # T11 row7 只有 FITC, 缺 PE
    (11, 7, 1, "FITC", "荧光检测灵敏度"),
    (11, 7, 2, "≤90MESF", "FITC≤90MESF、PE≤22MESF"),
]

# T11 需要删除的重复行序号 (1-based data row index)
T11_DUPLICATE_ROWS = [10, 16, 17]  # "自动加样工作站" 重复行
# 同时修正 row9 的文字使之覆盖全部规格
T11_ROW9_FIX = {
    1: "自动加样工作站",
    2: "★具备自动加样工作站，可兼容40管流式管、96孔板、384孔板上样",
}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  配置清单补充数据（按 table_index）
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# (配置名称, 单位, 数量, 是否标配, 用途说明, 备注)
CONFIG_ADDITIONS: dict[int, list[tuple[str, str, str, str, str, str]]] = {
    # T4 – 包2 特种蛋白分析仪 (目前仅1行)
    4: [
        ("主机", "台", "3", "是", "检测主机", "按采购数量配备"),
        ("试剂盘", "个", "3", "是", "试剂存放", "试剂位≥24个"),
        ("样品盘", "个", "3", "是", "样品存放", "样品位≥45个"),
        ("比色杯", "套", "3", "是", "检测耗材", "可清洗，标配"),
        ("稀释模块", "套", "3", "是", "自动稀释", "具有样本稀释功能"),
        ("操作软件", "套", "3", "是", "数据采集/分析", "含LIS对接模块"),
        ("电源线及附件", "套", "3", "是", "配套附件", "标准附件"),
        ("操作说明书", "本", "3", "是", "随机资料", "含中文说明书"),
        ("出厂检验报告", "份", "3", "是", "随机资料", "每台1份"),
        ("合格证", "份", "3", "是", "随机资料", "每台1份"),
    ],
    # T6 – 包3 荧光操作仪 (目前仅1行)
    6: [
        ("主机", "台", "2", "是", "检测主机", "按采购数量配备"),
        ("加样针", "支", "6", "是", "自动加样", "≥3针/台"),
        ("试剂位模块", "套", "2", "是", "试剂存放", "≥12个试剂位"),
        ("样本位模块", "套", "2", "是", "样本存放", "荧光样本位≥160个"),
        ("载片架", "套", "2", "是", "载片装载", "最大加载≥20张/次"),
        ("稀释模块", "套", "2", "是", "样本稀释", "可自定义多种稀释"),
        ("电容液面探测模块", "套", "2", "是", "液面探测", "电容探测技术"),
        ("操作软件", "套", "2", "是", "数据采集/分析", "支持间接免疫荧光法和酶联免疫法"),
        ("电源线及附件", "套", "2", "是", "配套附件", "标准附件"),
        ("操作说明书", "本", "2", "是", "随机资料", "含中文说明书"),
        ("合格证", "份", "2", "是", "随机资料", "每台1份"),
    ],
    # T8 – 包4 荧光显微镜 (目前3行，补充)
    8: [
        ("目镜", "对", "1", "是", "观察部件", "10X宽视野目镜，视场数≥20"),
        ("三目镜筒", "个", "1", "是", "观察部件", "视场数≥20"),
        ("物镜转换器", "个", "1", "是", "光学部件", "手动≥4档"),
        ("载物台", "个", "1", "是", "承载部件", "陶瓷表面"),
        ("阿贝聚光器", "个", "1", "是", "光学部件", "标准配置"),
        ("荧光光源模块", "套", "1", "是", "荧光激发", "无衰减输出≥50000小时"),
        ("激发滤光片", "片", "1", "是", "光学部件", "470~490nm"),
        ("透射光源", "个", "1", "是", "照明部件", "卤素灯"),
        ("电源线及附件", "套", "1", "是", "配套附件", "标准附件"),
        ("操作说明书", "本", "1", "是", "随机资料", "含中文说明书"),
    ],
    # T10 – 包5 化学发光 (目前4行，补充)
    10: [
        ("主机", "台", "1", "是", "检测主机", "按采购数量配备"),
        ("样品仓", "个", "1", "是", "样品存放", "单机≥250样本"),
        ("试剂仓", "个", "1", "是", "试剂存放", "单机≥30个，具有冷藏功能"),
        ("一次性Tip头架", "套", "1", "是", "耗材配件", "一次性Tip加样或双重清洗"),
        ("反应单元", "套", "1", "是", "核心部件", "温控37℃±0.3℃"),
        ("质控品", "套", "1", "是", "配套试剂", "原厂室内质控品"),
        ("操作说明书", "本", "1", "是", "随机资料", "含中文说明书"),
        ("出厂检验报告", "份", "1", "是", "随机资料", "标准附件"),
        ("合格证", "份", "1", "是", "随机资料", "标准附件"),
    ],
    # T12 – 包6 流式细胞 (目前4行，补充)
    12: [
        ("主机", "台", "2", "是", "检测主机", "按采购数量配备"),
        ("蓝色激光器（490±5nm）", "个", "2", "是", "光学部件", "标准配置"),
        ("红色激光器（640±5nm）", "个", "2", "是", "光学部件", "标准配置"),
        ("紫色激光器（405±5nm）", "个", "2", "是", "光学部件", "标准配置"),
        ("PMT检测器", "套", "2", "是", "检测部件", "≥12色荧光检测通道"),
        ("六色淋巴细胞亚群试剂盒", "盒", "2", "是", "配套试剂", "三类医疗器械注册证"),
        ("96孔板适配器", "个", "2", "是", "配套附件", "自动加样工作站配件"),
        ("384孔板适配器", "个", "2", "是", "配套附件", "自动加样工作站配件"),
        ("40管流式管架", "个", "2", "是", "配套附件", "自动加样工作站配件"),
        ("操作说明书", "本", "2", "是", "随机资料", "含中文说明书"),
        ("合格证", "份", "2", "是", "随机资料", "每台1份"),
    ],
}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  审查表 / 评审表 备注数据
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

# T13-T18 资格性审查: (col4: 是否满足, col5: 备注)
QUALIFICATION_REMARKS = [
    "核对营业执照/主体资格证明、资格承诺函及附随证明材料是否齐全并完成签章。",
    "核对资格承诺函中关于关联关系的承诺是否按实际签章确认。",
    "核对资格承诺函中关于信用记录的承诺是否按实际签章确认。",
    "核对围标串标承诺函是否已签字盖章。",
    "如非事业单位/社会团体可在最终稿注明'不适用'；如适用需补承诺并签章。",
    "法定代表人本人参加则本项不适用；授权代理参加需附完整授权及身份证明。",
    "按所投产品医疗器械管理类别后附经营/生产许可、备案凭证或注册证。",
    "核对围标串标承诺函是否已签字盖章。",
]

# T19-T24 符合性审查
COMPLIANCE_COL3_OVERRIDES = {
    "主要商务条款": "二、报价书；六、技术服务和售后服务的内容及措施",
    "联合体投标": "二、报价书或联合体协议（如适用）",
    "其他要求": "十一、投标人关联单位的说明；附四、投标无效情形汇总及自检表；围标串标承诺函",
}
COMPLIANCE_REMARKS = [
    "提交前复核报价唯一性、完整性及是否超预算。",
    "提交前复核签字盖章、目录、页码和格式规范性。",
    "核对主要商务条款承诺书与报价书内容一致，并完成签章。",
    "非联合体可在最终稿标注'不适用'；如联合体需附联合体协议。",
    "仍需结合所投品牌型号逐条补齐技术响应、配置和证明材料。",
    "核对关联关系说明、围标串标承诺函和无效投标自检表。",
]

# T25-T30 评审表: 自评说明 col4 补全
SCORING_SELF_EVAL = {
    1: "逐条填写品牌、型号、规格、配置、响应值及偏离情况，不得仅写'响应/完全响应'。",
    2: "已在技术偏离及详细配置明细表中预设装箱配置、随机附件、随机资料及标配/选配清单填报框架。",
    6: "已将采购文件售后服务要求逐项写入承诺，并补充质保、维护、备件和升级保障安排。",
    7: "报价书及报价一览表对应本包报价，提交前仍需复核分项与总价一致性及预算控制。",
}

# T31 自检表备注
SELF_CHECK_REMARKS = [
    "核对远程开标安排、账号和网络环境。",
    "核对CA证书、解密环境和时限要求。",
    "核对数字证书有效期和绑定关系。",
    "核对本机环境、浏览器/客户端与CA驱动。",
    "对照技术偏离表逐条核查★条款是否实质性响应。",
    "统计非★负偏离条数，单项产品不得达到5条及以上。",
    "已预设逐条响应框架，仍需按所投品牌型号填写响应值和证据。",
    "仍需补齐品牌、型号、规格、配置等实质信息。",
    "复核单项报价不得超单项预算。",
    "如涉及软件或计算机配置，补充正版授权或序列号证明；不适用时注明。",
    "如涉及强制节能产品，补充清单内证明；不适用时注明。",
    "重大偏离需提前自检，必要时调整选型或响应方案。",
    "提交前核对所有需签字、盖章页面。",
    "严禁提供虚假材料，提交前复核真实性。",
    "响应参数需与证明文件保持一致。",
    "核对使用成本和现场条件是否能被采购人接受。",
    "核对授权书签字盖章完整性；法定代表人本人参加则按规则处理。",
    "复核近三年重大违法记录承诺及查询情况。",
    "全程避免与采购人、代理机构、评审专家不当接触。",
    "核对股权、管理关系，避免同一控制下重复投标。",
    "核对不存在围标串标或被视为串通的情形。",
    "复核其他法律法规导致无效投标的情形。",
    "谈判依据以响应文件为准，不再依赖文件外材料。",
    "资格性、符合性表需结合签章和附件在提交前改为最终结论。",
    "上传投标客户端前复核应答点标记与索引位置。",
]

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  工具函数
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def _clone_row(table, source_row_idx: int) -> _Row:
    """复制表格行的 XML 结构并追加到表格末尾，返回新行。"""
    tbl_element = table._tbl
    src_tr = table.rows[source_row_idx]._tr
    new_tr = copy.deepcopy(src_tr)
    tbl_element.append(new_tr)
    return _Row(new_tr, table)


def _delete_row(table, row_idx: int) -> None:
    """删除指定索引的行。"""
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)


def _set_cell(row, col: int, text: str) -> None:
    """设置单元格纯文本（保留格式）。"""
    cell = row.cells[col]
    # 清空所有段落
    for p in cell.paragraphs:
        p.text = ""
    cell.paragraphs[0].text = text


def _clean_response_text(text: str) -> str:
    """清洗技术偏离表「响应文件响应情况」列中的冗余占位提示。"""
    text = re.sub(r"【待填写：品牌/型号/规格/配置及逐条响应】[；;]?\s*", "", text)
    text = re.sub(r"请填写[^。；\n]{0,30}[。；]?\s*", "", text)
    text = text.strip()
    if not text:
        return "【待填写：实际参数值】"
    return text


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  修复步骤
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def step1_remove_dirty_paragraphs(doc: Document) -> int:
    """删除脏段落（"说明：当前已结构化生成…"）以及紧邻的提示段落。"""
    removed = 0
    for p in list(doc.paragraphs):
        if "当前已结构化生成" in p.text and "剩余条款必须继续" in p.text:
            p._element.getparent().remove(p._element)
            removed += 1
    # 也删除 "注：当前仅依据采购文件展开技术条款；未接入投标产品事实/证据时" 提示
    for p in list(doc.paragraphs):
        if "当前仅依据采购文件展开技术条款" in p.text:
            p._element.getparent().remove(p._element)
            removed += 1
    print(f"  [Step 1] 已删除 {removed} 条脏/提示段落")
    return removed


def step2_fix_tech_tables(doc: Document) -> None:
    """修复技术偏离表: 补行、修文字、去重复。"""
    tables = doc.tables

    # 2a. 修正单元格文字
    for ti, ri, ci, old_sub, new_val in CELL_FIXES:
        cell_text = tables[ti].rows[ri].cells[ci].text
        if old_sub in cell_text:
            _set_cell(tables[ti].rows[ri], ci, new_val)
            print(f"  [Step 2a] T{ti} R{ri} C{ci}: '{old_sub}' → '{new_val}'")

    # 2b. T11 合并自动加样工作站重复行: 修正 row9 后删除 rows 10,16,17
    t11 = tables[11]
    # 修正 row9 覆盖完整规格
    for ci, val in T11_ROW9_FIX.items():
        _set_cell(t11.rows[9], ci, val)
    # 从大到小删除避免索引偏移
    for ri in sorted(T11_DUPLICATE_ROWS, reverse=True):
        _delete_row(t11, ri)
    print(f"  [Step 2b] T11: 修正 row9 + 删除 {len(T11_DUPLICATE_ROWS)} 重复行")

    # 2c. 补缺失行
    for ti, patches in TECH_PATCHES.items():
        if not patches:
            continue
        tbl = tables[ti]
        template_row_idx = 1  # 用第1数据行做模板
        for _, param, req, resp in patches:
            new_row = _clone_row(tbl, template_row_idx)
            n_rows = len(tbl.rows)
            _set_cell(new_row, 0, str(n_rows - 1))
            _set_cell(new_row, 1, param)
            _set_cell(new_row, 2, req)
            _set_cell(new_row, 3, resp)
            _set_cell(new_row, 4, "【待填写：无偏离/正偏离/负偏离】")
            print(f"  [Step 2c] T{ti}: 新增行 '{param}'")

    # 2d. 重新编号所有技术偏离表
    for ti in [1, 3, 5, 7, 9, 11]:
        tbl = tables[ti]
        for idx, row in enumerate(tbl.rows[1:], 1):
            _set_cell(row, 0, str(idx))

    # 2e. 清洗响应列文字
    cleaned = 0
    for ti in [1, 3, 5, 7, 9, 11]:
        tbl = tables[ti]
        for row in tbl.rows[1:]:
            old = row.cells[3].text
            new = _clean_response_text(old)
            if new != old:
                _set_cell(row, 3, new)
                cleaned += 1
    print(f"  [Step 2e] 清洗 {cleaned} 个响应列单元格")


def step3_fill_config_tables(doc: Document) -> None:
    """填充配置清单表。"""
    tables = doc.tables
    for ti, additions in CONFIG_ADDITIONS.items():
        tbl = tables[ti]
        existing = len(tbl.rows) - 1  # 不含表头
        template_row_idx = 1
        for i, (name, unit, qty, std, purpose, note) in enumerate(additions):
            new_row = _clone_row(tbl, template_row_idx)
            seq = existing + i + 1
            _set_cell(new_row, 0, str(seq))
            _set_cell(new_row, 1, name)
            _set_cell(new_row, 2, unit)
            _set_cell(new_row, 3, qty)
            _set_cell(new_row, 4, std)
            _set_cell(new_row, 5, purpose)
            _set_cell(new_row, 6, note)
        print(f"  [Step 3] T{ti}: 新增 {len(additions)} 行配置")


def step4_fill_review_tables(doc: Document) -> None:
    """预填审查表 / 评审表。"""
    tables = doc.tables

    # 4a. T13-T18 资格性审查
    for ti in range(13, 19):
        tbl = tables[ti]
        for idx, row in enumerate(tbl.rows[1:]):
            _set_cell(row, 4, "待复核")
            if idx < len(QUALIFICATION_REMARKS):
                _set_cell(row, 5, QUALIFICATION_REMARKS[idx])
    print(f"  [Step 4a] T13-T18 资格性审查: 已预填 '待复核' + 备注")

    # 4b. T19-T24 符合性审查
    for ti in range(19, 25):
        tbl = tables[ti]
        for idx, row in enumerate(tbl.rows[1:]):
            label = row.cells[1].text.strip()
            # 补全对应章节
            if label in COMPLIANCE_COL3_OVERRIDES:
                _set_cell(row, 3, COMPLIANCE_COL3_OVERRIDES[label])
            _set_cell(row, 4, "待复核")
            if idx < len(COMPLIANCE_REMARKS):
                _set_cell(row, 5, COMPLIANCE_REMARKS[idx])
    print(f"  [Step 4b] T19-T24 符合性审查: 已预填 '待复核' + 备注")

    # 4c. T25-T30 评审表: 补全自评说明
    for ti in range(25, 31):
        tbl = tables[ti]
        for row_1based, text in SCORING_SELF_EVAL.items():
            if row_1based < len(tbl.rows):
                _set_cell(tbl.rows[row_1based], 4, text)
    print(f"  [Step 4c] T25-T30 评审表: 已补全自评说明")

    # 4d. T31 自检表
    tbl = tables[31]
    for idx, row in enumerate(tbl.rows[1:]):
        _set_cell(row, 2, "待复核")
        if idx < len(SELF_CHECK_REMARKS):
            _set_cell(row, 3, SELF_CHECK_REMARKS[idx])
    print(f"  [Step 4d] T31 自检表: 已预填 '待复核' + 备注")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  主流程
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def main() -> None:
    print(f"读取: {ORIGINAL}")
    doc = Document(str(ORIGINAL))

    step1_remove_dirty_paragraphs(doc)
    step2_fix_tech_tables(doc)
    step3_fill_config_tables(doc)
    step4_fill_review_tables(doc)

    doc.save(str(REVISED))
    print(f"\n✅ 已保存修复版: {REVISED}")


if __name__ == "__main__":
    main()
