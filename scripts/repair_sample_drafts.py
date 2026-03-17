from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
import sys
from types import SimpleNamespace

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.services.one_click_generator.format_driven_sections.common import (
    _build_affiliated_units_statement_template,
    _build_disabled_unit_declaration_template,
    _build_manufacturer_authorization_template,
    _build_service_acceptance_points,
    _build_service_after_sales_points,
    _build_service_commitment_points,
    _build_service_fee_commitment_template,
    _build_service_installation_points,
    _build_service_milestone_points,
    _build_service_packaging_points,
    _build_small_enterprise_declaration_template,
    _build_service_supply_points,
    _build_service_training_points,
)


TEXTFILE_DIR = Path(__file__).resolve().parents[1] / "textfile"
TARGETS = [
    "投标底稿_检验科购置全自动电泳仪等设备 (3).docx",
    "投标底稿_手术用头架、X射线血液辐照设备(二次).docx",
    "投标底稿_吉林大学中日联谊医院流式细胞仪采购项目.docx",
]

TENDER_OVERRIDES = {
    "投标底稿_检验科购置全自动电泳仪等设备 (3).docx": {
        "purchaser": "哈尔滨医科大学附属第二医院",
        "agency": "方大国际工程咨询股份有限公司",
    },
    "投标底稿_手术用头架、X射线血液辐照设备(二次).docx": {
        "purchaser": "哈尔滨医科大学附属第一医院",
        "agency": "黑龙江伟达项目管理有限公司",
    },
    "投标底稿_吉林大学中日联谊医院流式细胞仪采购项目.docx": {
        "purchaser": "吉林大学中日联谊医院",
        "agency": "北京典方建设工程咨询有限公司",
    },
}

REMOVE_EXACT = {
    "请按招标文件原格式填写本节内容。",
    "注：当前仅依据采购文件展开技术条款；未接入投标产品事实/证据时，响应值与偏离结论不得预填。",
}

NEW_COMPANY_EXPLANATION = (
    "如供应商属于新成立企业，或依法在本阶段不需提供对应缴纳证明，请另附书面说明，"
    "写明成立时间、适用依据及不能提供相关证明材料的原因，并同步附营业执照、主管部门说明或其他佐证材料。"
)

TP_SECTION_GROUPS = [
    (
        "第一章、资格性证明文件",
        [
            "四、资格承诺函",
            "七、法定代表人/单位负责人授权书",
            "八、法定代表人/单位负责人和授权代表身份证明",
        ],
    ),
    (
        "第二章、符合性承诺",
        [
            "十一、投标人关联单位的说明",
            "九、小微企业声明函",
            "十、残疾人福利性单位声明函",
        ],
    ),
    (
        "第三章、商务及技术部分",
        [
            "二、报价书",
            "三、报价一览表",
            "六、技术服务和售后服务的内容及措施",
        ],
    ),
    (
        "第四章、报价书附件",
        [
            "五、技术偏离及详细配置明细表",
            "附一、资格性审查响应对照表",
            "附二、符合性审查响应对照表",
            "附三、详细评审响应对照表",
            "附四、投标无效情形汇总及自检表",
        ],
    ),
]

CS_SECTION_GROUPS = [
    (
        "第一章、资格性证明文件",
        [
            "十一、资格承诺函",
            "六、法定代表人/单位负责人授权书",
            "七、法定代表人/单位负责人和授权代表身份证明",
        ],
    ),
    (
        "第二章、符合性承诺",
        [
            "十、投标人关联单位的说明",
            "八、小微企业声明函",
            "九、残疾人福利性单位声明函",
        ],
    ),
    (
        "第三章、商务及技术部分",
        [
            "二、首轮报价表",
            "三、分项报价表",
            "五、技术服务和售后服务的内容及措施",
        ],
    ),
    (
        "第四章、报价书附件",
        [
            "四、技术偏离及详细配置明细表",
            "附一、资格性审查响应对照表",
            "附二、符合性审查响应对照表",
            "附三、详细评审响应对照表",
            "附四、投标无效情形汇总及自检表",
        ],
    ),
]

SECTION_GROUPS_BY_FILE = {
    "投标底稿_检验科购置全自动电泳仪等设备 (3).docx": TP_SECTION_GROUPS,
    "投标底稿_手术用头架、X射线血液辐照设备(二次).docx": CS_SECTION_GROUPS,
}


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _insert_paragraph_before(paragraph, text: str):
    return paragraph.insert_paragraph_before(text)


def _set_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def _dedupe_repeated_text(text: str) -> str:
    stripped = (text or "").strip()
    if not stripped or len(stripped) % 2:
        return stripped
    half = len(stripped) // 2
    if stripped[:half] == stripped[half:]:
        return stripped[:half]
    return stripped


def _repair_document(path: Path) -> dict[str, int]:
    doc = Document(path)
    stats = {
        "deleted": 0,
        "deduped": 0,
        "replaced": 0,
        "blocks": 0,
        "reordered": 0,
    }

    paragraphs = list(doc.paragraphs)
    awaiting_date_line = False
    skip_month_day = False

    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        deduped = _dedupe_repeated_text(text)
        if deduped != text:
            _set_text(paragraph, deduped)
            text = deduped
            stats["deduped"] += 1

        if text in REMOVE_EXACT:
            _delete_paragraph(paragraph)
            stats["deleted"] += 1
            continue

        if text.startswith("承诺人(") and "签章" in text and "待填写" not in text:
            _set_text(paragraph, "承诺人（供应商或自然人CA签章）：【待填写：投标人名称】")
            awaiting_date_line = True
            stats["replaced"] += 1
            continue

        if awaiting_date_line and text == "年":
            _set_text(paragraph, "日期：【待填写：年 月 日】")
            awaiting_date_line = False
            skip_month_day = True
            stats["replaced"] += 1
            continue

        if skip_month_day and text in {"月", "日"}:
            _delete_paragraph(paragraph)
            stats["deleted"] += 1
            continue

        if skip_month_day and text == "日附件":
            _set_text(paragraph, "附件")
            skip_month_day = False
            stats["replaced"] += 1
            continue

        if text == "缴纳社会保障资金的证明材料清单缴纳社会保障资金的证明材料清单":
            _set_text(paragraph, "缴纳社会保障资金的证明材料清单")
            stats["replaced"] += 1
            continue

        if text.startswith("二、新成立的企业或在法规范围内不需提供"):
            next_paragraph = paragraphs[idx + 1] if idx + 1 < len(paragraphs) else None
            if next_paragraph is not None and next_paragraph.text.strip() == "请按招标文件原格式填写本节内容。":
                _set_text(next_paragraph, NEW_COMPANY_EXPLANATION)
                stats["replaced"] += 1

    tender = _extract_tender_stub(doc, path.name)
    package_names = _extract_package_names(doc)

    if "检验科购置全自动电泳仪等设备" in path.name:
        service_packages = _extract_service_packages(
            doc,
            "六、技术服务和售后服务的内容及措施",
            "七、法定代表人/单位负责人授权书",
        )
        if service_packages:
            _replace_section_body(
                doc,
                "六、技术服务和售后服务的内容及措施",
                "七、法定代表人/单位负责人授权书",
                _build_sample_service_section_lines(service_packages),
            )
            stats["blocks"] += 1
        _replace_section_body(
            doc,
            "九、小微企业声明函",
            "十、残疾人福利性单位声明函",
            _build_small_enterprise_declaration_template(tender, package_names).splitlines(),
        )
        _replace_section_body(
            doc,
            "十、残疾人福利性单位声明函",
            "十一、投标人关联单位的说明",
            _build_disabled_unit_declaration_template(tender, package_names).splitlines(),
        )
        _replace_section_body(
            doc,
            "十一、投标人关联单位的说明",
            "附一、资格性审查响应对照表",
            _build_affiliated_units_statement_template(tender).splitlines(),
        )
        stats["blocks"] += 3

    if "手术用头架、X射线血液辐照设备" in path.name:
        service_packages = _extract_service_packages(
            doc,
            "五、技术服务和售后服务的内容及措施",
            "六、法定代表人/单位负责人授权书",
        )
        if service_packages:
            _replace_section_body(
                doc,
                "五、技术服务和售后服务的内容及措施",
                "六、法定代表人/单位负责人授权书",
                _build_sample_service_section_lines(service_packages),
            )
            stats["blocks"] += 1
        _replace_section_body(
            doc,
            "八、小微企业声明函",
            "九、残疾人福利性单位声明函",
            _build_small_enterprise_declaration_template(tender, package_names).splitlines(),
        )
        _replace_section_body(
            doc,
            "九、残疾人福利性单位声明函",
            "十、投标人关联单位的说明",
            _build_disabled_unit_declaration_template(tender, package_names).splitlines(),
        )
        _replace_section_body(
            doc,
            "十、投标人关联单位的说明",
            "十一、资格承诺函",
            _build_affiliated_units_statement_template(tender).splitlines(),
        )
        stats["blocks"] += 3

    if "吉林大学中日联谊医院流式细胞仪采购项目" in path.name:
        _replace_section_body(
            doc,
            "7.12制造商授权书(格式自拟)",
            "格式8.采购需求响应及偏离表(格式)",
            _build_manufacturer_authorization_template(tender, package_names).splitlines(),
        )
        _replace_section_body(
            doc,
            "三、招标代理服务费承诺",
            "附件：",
            _build_service_fee_commitment_template(tender).splitlines(),
        )
        stats["blocks"] += 2

    section_groups = SECTION_GROUPS_BY_FILE.get(path.name)
    if section_groups and _reorder_document_sections(doc, section_groups):
        stats["reordered"] += 1

    doc.save(path)
    return stats


def _extract_tender_stub(doc: Document, file_name: str | None = None):
    project_name = ""
    project_number = ""
    purchaser = ""
    agency = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if not project_name and text.startswith("项目名称："):
            tail = text.split("：", 1)[1].strip()
            tail = re.split(r"\s+(?:项目编号|招标编号)\s*[：:]", tail, maxsplit=1)[0].strip()
            project_name = tail
        if not project_number and ("项目编号：" in text or "招标编号：" in text):
            project_number = re.split(r"[：:]", text, maxsplit=1)[1].strip()
            project_number = project_number.split()[0]
        if text.startswith("致："):
            unit = text.split("：", 1)[1].strip()
            if "咨询" in unit or "招标" in unit or "代理" in unit:
                if not agency:
                    agency = unit
            elif not purchaser:
                purchaser = unit
        if "北京典方建设工程咨询有限公司" in text:
            agency = "北京典方建设工程咨询有限公司"

    override = TENDER_OVERRIDES.get(file_name or "", {})

    return SimpleNamespace(
        project_name=project_name or "【待填写：项目名称】",
        project_number=project_number or "【待填写：项目编号】",
        purchaser=override.get("purchaser") or purchaser or "【待填写：采购人】",
        agency=override.get("agency") or agency or "【待填写：代理机构】",
    )


def _extract_package_names(doc: Document):
    names: list[SimpleNamespace] = []
    seen: set[str] = set()
    pattern = re.compile(r"(?:合同包|包)\s*\d+\s*[:：]\s*(.+)")
    fallback_pattern = re.compile(r"对应货物[:：]\s*(.+)")

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        match = pattern.search(text)
        if match:
            raw_names = [match.group(1).strip()]
        else:
            fallback = fallback_pattern.search(text)
            raw_names = fallback.group(1).strip().split("、") if fallback else []
        for name in raw_names:
            name = name.strip()
            if not name or name in seen:
                continue
            seen.add(name)
            names.append(SimpleNamespace(item_name=name))
    return names


def _normalize_doc_line(text: str) -> str:
    return (text or "").strip().lstrip("-• ").strip()


def _element_text(element) -> str:
    texts = [node.text for node in element.iter() if node.tag.endswith("}t") and node.text]
    return _normalize_doc_line("".join(texts))


def _is_paragraph_element(element) -> bool:
    return element.tag.endswith("}p")


def _is_section_properties_element(element) -> bool:
    return element.tag.endswith("}sectPr")


def _reorder_document_sections(doc: Document, section_groups: list[tuple[str, list[str]]]) -> bool:
    body = doc.element.body
    children = list(body.iterchildren())
    if not children:
        return False

    sect_pr = deepcopy(children[-1]) if _is_section_properties_element(children[-1]) else None
    content_children = children[:-1] if sect_pr is not None else children
    chapter_titles = {title for title, _ in section_groups}
    ordered_titles = [title for _, titles in section_groups for title in titles]
    known_titles = set(ordered_titles)

    block_starts: list[tuple[int, str]] = []
    seen_titles: set[str] = set()
    for idx, child in enumerate(content_children):
        if not _is_paragraph_element(child):
            continue
        text = _element_text(child)
        if text in known_titles and text not in seen_titles:
            block_starts.append((idx, text))
            seen_titles.add(text)

    if not block_starts:
        return False

    block_starts.sort()
    prefix = [
        deepcopy(child)
        for child in content_children[:block_starts[0][0]]
        if _element_text(child) not in chapter_titles
    ]
    blocks: dict[str, list] = {}

    for position, (start_idx, title) in enumerate(block_starts):
        end_idx = block_starts[position + 1][0] if position + 1 < len(block_starts) else len(content_children)
        blocks[title] = [
            deepcopy(child)
            for child in content_children[start_idx:end_idx]
            if _element_text(child) not in chapter_titles
        ]

    if not any(title in blocks for title in ordered_titles):
        return False

    for child in list(body.iterchildren()):
        body.remove(child)

    for child in prefix:
        body.append(child)

    for chapter_title, titles in section_groups:
        chapter_sections = [title for title in titles if title in blocks]
        if not chapter_sections:
            continue
        paragraph = doc.add_paragraph(chapter_title)
        if paragraph.runs:
            paragraph.runs[0].bold = True
        for title in chapter_sections:
            for child in blocks[title]:
                body.append(child)

    if sect_pr is not None:
        body.append(sect_pr)
    return True


def _extract_service_packages(doc: Document, start_title: str, end_title: str) -> list[SimpleNamespace]:
    packages: list[SimpleNamespace] = []
    current: dict[str, str] | None = None
    in_section = False

    for paragraph in doc.paragraphs:
        text = _normalize_doc_line(paragraph.text)
        if not text:
            continue
        if not in_section:
            if text == start_title:
                in_section = True
            continue
        if text == end_title:
            break

        pkg_match = re.match(r"^(?:合同包|包)\s*(\d+)\s*[:：]\s*(.+)$", text)
        if pkg_match:
            if current:
                packages.append(SimpleNamespace(**current))
            current = {
                "package_id": pkg_match.group(1),
                "item_name": pkg_match.group(2).strip(),
                "product_identity": "",
                "delivery_time": "按采购文件要求",
                "delivery_place": "采购人指定地点",
            }
            continue

        if current is None:
            continue

        if "拟投产品" in text:
            parts = re.split(r"[：:]", text, maxsplit=1)
            current["product_identity"] = parts[1].strip() if len(parts) == 2 else current["product_identity"]
        elif text.startswith("交货期"):
            parts = re.split(r"[：:]", text, maxsplit=1)
            current["delivery_time"] = parts[1].strip() if len(parts) == 2 else current["delivery_time"]
        elif text.startswith("交货地点"):
            parts = re.split(r"[：:]", text, maxsplit=1)
            current["delivery_place"] = parts[1].strip() if len(parts) == 2 else current["delivery_place"]

    if current:
        packages.append(SimpleNamespace(**current))
    return packages


def _build_sample_service_section_lines(packages: list[SimpleNamespace]) -> list[str]:
    tech_lines: list[str] = ["（一）技术服务"]
    service_lines: list[str] = ["（二）售后服务"]
    for pkg in packages:
        item_name = getattr(pkg, "item_name", "") or "【待填写：货物名称】"
        product_identity = getattr(pkg, "product_identity", "") or item_name
        delivery_time = getattr(pkg, "delivery_time", "") or "按采购文件要求"
        delivery_place = getattr(pkg, "delivery_place", "") or "采购人指定地点"

        tech_lines.extend([
            f"包{getattr(pkg, 'package_id', '')}：{item_name}",
            f"拟投产品：{product_identity}",
            f"交货期：{delivery_time}",
            f"交货地点：{delivery_place}",
            "1. 供货组织与进度安排",
            *_build_service_supply_points(
                item_name,
                delivery_time,
                delivery_place,
                product_identity=product_identity,
            ),
            "2. 包装运输与到货保护",
            *_build_service_packaging_points(
                item_name,
                product_identity=product_identity,
            ),
            "3. 安装调试与场地联动",
            *_build_service_installation_points(
                item_name,
                delivery_place=delivery_place,
            ),
            "4. 培训实施",
            *_build_service_training_points(item_name),
            "5. 项目实施里程碑",
            *_build_service_milestone_points(
                item_name,
                delivery_time,
                delivery_place=delivery_place,
            ),
            "6. 验收与资料移交",
            *_build_service_acceptance_points(item_name),
        ])
        service_lines.extend([
            f"包{getattr(pkg, 'package_id', '')}：{item_name}",
            f"拟投产品：{product_identity}",
            f"交货期：{delivery_time}",
            f"交货地点：{delivery_place}",
            "1. 售后与维保安排",
            *_build_service_after_sales_points(item_name),
            "2. 服务保障承诺",
            *_build_service_commitment_points(
                item_name,
                delivery_time,
                delivery_place=delivery_place,
            ),
        ])

    lines = tech_lines + [""] + service_lines
    lines.extend([
        "供应商全称：【待填写：投标人名称】",
        "日期：【待填写：日期】",
    ])
    return lines


def _replace_section_body(doc: Document, start_title: str, end_title: str, lines: list[str]) -> None:
    paragraphs = list(doc.paragraphs)
    start_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip() == start_title), None)
    if start_idx is None:
        return

    end_idx = next((i for i in range(start_idx + 1, len(paragraphs)) if paragraphs[i].text.strip() == end_title), None)
    if end_idx is None:
        return

    anchor = paragraphs[end_idx]
    for paragraph in paragraphs[start_idx + 1:end_idx]:
        _delete_paragraph(paragraph)

    clean_lines = [line.rstrip() for line in lines if line.strip()]
    for line in clean_lines:
        _insert_paragraph_before(anchor, line)


def main() -> None:
    for name in TARGETS:
        path = TEXTFILE_DIR / name
        if not path.exists():
            print(f"skip {name}: file not found")
            continue
        stats = _repair_document(path)
        print(f"{name}: {stats}")


if __name__ == "__main__":
    main()
