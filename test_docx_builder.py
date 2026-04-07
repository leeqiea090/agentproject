from pathlib import Path

from docx import Document

from app.schemas import BidDocumentSection, CompanyProfile, ProcurementPackage, TenderDocument
from app.services.bid_preferences import strip_section_number_prefix
from app.services.docx_builder import (
    _assert_new_structure_only,
    _backfill_required_sections,
    build_bid_docx,
    _ensure_toc_subheadings,
)


def _section_key(title: str) -> str:
    return strip_section_number_prefix(title).strip()


def _build_tp_tender() -> TenderDocument:
    return TenderDocument(
        project_name="示例项目",
        project_number="ABC-001",
        budget=0,
        purchaser="示例采购人",
        procurement_type="竞争性谈判",
        packages=[
            ProcurementPackage(
                package_id="1",
                item_name="示例设备",
                quantity=1,
                budget=0,
            )
        ],
    )


def _build_cs_tender() -> TenderDocument:
    return TenderDocument(
        project_name="示例项目",
        project_number="ABC-001",
        budget=0,
        purchaser="示例采购人",
        procurement_type="竞争性磋商",
        packages=[
            ProcurementPackage(
                package_id="1",
                item_name="示例设备",
                quantity=1,
                budget=0,
            )
        ],
    )


def test_backfill_required_sections_reuses_unnumbered_existing_sections():
    tender = _build_tp_tender()
    sections = [
        BidDocumentSection(section_title="响应文件封面格式", content="封面正文"),
        BidDocumentSection(section_title="报价书", content="报价书正文"),
        BidDocumentSection(section_title="资格承诺函", content="资格承诺正文"),
    ]

    filled = _backfill_required_sections(sections, tender=tender)
    title_keys = [_section_key(section.section_title) for section in filled]

    assert title_keys.count("响应文件封面格式") == 1
    assert title_keys.count("报价书") == 1
    assert title_keys.count("资格承诺函") == 1
    assert filled[0].section_title == "响应文件封面格式"
    assert next(section for section in filled if _section_key(section.section_title) == "资格承诺函").content == "资格承诺正文"


def test_assert_new_structure_only_accepts_unnumbered_tp_titles():
    tp_titles = [
        "响应文件封面格式",
        "报价书",
        "报价一览表",
        "资格承诺函",
        "技术偏离及详细配置明细表",
        "技术服务和售后服务的内容及措施",
        "法定代表人/单位负责人授权书",
        "法定代表人/单位负责人和授权代表身份证明",
        "小微企业声明函",
        "残疾人福利性单位声明函",
        "投标人关联单位的说明",
    ]
    sections = [
        BidDocumentSection(section_title=title, content=f"{title}正文")
        for title in tp_titles
    ]

    _assert_new_structure_only(sections, tender=_build_tp_tender())


def test_assert_new_structure_only_accepts_actual_tp_appendix_titles():
    sections = [
        BidDocumentSection(section_title=title, content=f"{title}正文")
        for title in [
            "一、响应文件封面格式",
            "第一章、资格性证明文件",
            "四、资格承诺函",
            "七、法定代表人/单位负责人授权书",
            "八、法定代表人/单位负责人和授权代表身份证明",
            "第二章、符合性承诺",
            "十一、投标人关联单位的说明",
            "九、小微企业声明函",
            "十、残疾人福利性单位声明函",
            "第三章、商务及技术部分",
            "二、报价书",
            "三、报价一览表",
            "六、技术服务和售后服务的内容及措施",
            "第四章、报价书附件",
            "五、技术偏离及详细配置明细表",
            "附一、资格性审查响应对照表",
            "附二、符合性审查响应对照表",
            "附三、详细评审响应对照表",
            "附四、投标无效情形汇总及自检表",
        ]
    ]

    _assert_new_structure_only(sections, tender=_build_tp_tender())


def test_assert_new_structure_only_accepts_actual_cs_appendix_titles():
    sections = [
        BidDocumentSection(section_title=title, content=f"{title}正文")
        for title in [
            "一、响应文件封面格式",
            "第一章、资格性证明文件",
            "十一、资格承诺函",
            "六、法定代表人/单位负责人授权书",
            "七、法定代表人/单位负责人和授权代表身份证明",
            "第二章、符合性承诺",
            "十、投标人关联单位的说明",
            "八、小微企业声明函",
            "九、残疾人福利性单位声明函",
            "第三章、商务及技术部分",
            "二、首轮报价表",
            "三、分项报价表",
            "五、技术服务和售后服务的内容及措施",
            "第四章、报价书附件",
            "四、技术偏离及详细配置明细表",
            "附一、资格性审查响应对照表",
            "附二、符合性审查响应对照表",
            "附三、详细评审响应对照表",
            "附四、投标无效情形汇总及自检表",
        ]
    ]

    _assert_new_structure_only(sections, tender=_build_cs_tender())


def test_assert_new_structure_only_still_rejects_legacy_tp_outline_titles():
    sections = [
        BidDocumentSection(section_title=title, content=f"{title}正文")
        for title in [
            "一、响应文件封面格式",
            "二、报价书",
            "三、报价一览表",
            "四、资格承诺函",
            "五、技术偏离及详细配置明细表",
            "六、技术服务和售后服务的内容及措施",
            "七、法定代表人/单位负责人授权书",
            "八、法定代表人/单位负责人和授权代表身份证明",
            "九、小微企业声明函",
            "十、残疾人福利性单位声明函",
            "十一、投标人关联单位的说明",
            "七、报价书附件",
        ]
    ]

    try:
        _assert_new_structure_only(sections, tender=_build_tp_tender())
    except RuntimeError as exc:
        assert "报价书附件" in str(exc)
    else:
        raise AssertionError("expected legacy tp outline title to be rejected")


def test_ensure_toc_subheadings_supports_unnumbered_titles():
    content = _ensure_toc_subheadings("资格承诺函", "资格承诺正文")

    assert content.startswith("## 1. 资格承诺正文")
    assert content.endswith("资格承诺正文")


def test_build_bid_docx_keeps_semantic_sections_unique_for_unnumbered_titles(tmp_path: Path):
    tender = _build_tp_tender()
    sections = [
        BidDocumentSection(section_title="响应文件封面格式", content="政府采购\n响应文件\n示例项目"),
        BidDocumentSection(section_title="报价书", content="报价书正文"),
        BidDocumentSection(section_title="资格承诺函", content="资格承诺正文"),
    ]
    output_path = tmp_path / "sample.docx"

    build_bid_docx(
        sections,
        tender,
        company=CompanyProfile(
            name="示例公司",
            legal_representative="张三",
            address="示例地址",
            phone="123456",
        ),
        output_path=output_path,
    )

    doc = Document(output_path)
    headings = [
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.text.strip() and paragraph.style.name.startswith("Heading")
    ]

    assert sum(text.endswith("报价书") for text in headings) == 1
    assert sum(text.endswith("资格承诺函") for text in headings) == 1
    assert headings.count("1. 资格承诺正文") == 1
